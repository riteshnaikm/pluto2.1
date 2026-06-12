import os
import shutil
import subprocess
import tempfile
import pdfplumber
import logging
import hashlib
import json
import nltk
import sqlite3
from nltk.tokenize import sent_tokenize
from flask import Flask, request, jsonify, render_template, Response, stream_with_context, make_response, session, redirect, url_for
from dotenv import load_dotenv
from authlib.integrations.flask_client import OAuth
from functools import wraps

# Windows compatibility fix: Handle missing readline module
import sys
if sys.platform == 'win32':
    try:
        import readline
    except ImportError:
        try:
            # Try pyreadline3 as Windows alternative
            import pyreadline3 as readline
            sys.modules['readline'] = readline
        except ImportError:
            # Create a minimal dummy readline module if pyreadline3 is not available
            class DummyReadline:
                @staticmethod
                def add_history(*args, **kwargs):
                    pass
                @staticmethod
                def write_history_file(*args, **kwargs):
                    pass
            sys.modules['readline'] = DummyReadline()

from pinecone import Pinecone, ServerlessSpec
from langchain_community.vectorstores import Pinecone as PineconeVectorStore
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from rank_bm25 import BM25Okapi
from functools import lru_cache
import re
import pandas as pd
import warnings
genai = None
from openai import OpenAI
from groq import Groq, APIStatusError
from docx import Document
import uuid
from werkzeug.utils import secure_filename
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
import asyncio
import aiohttp
import queue
import threading
from concurrent.futures import ThreadPoolExecutor
from asgiref.wsgi import WsgiToAsgi
import time
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted, Flowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
import markdown
from html import unescape

from pluto.config import apply_flask_config, is_production, resolve_secret_key, strict_eval_json
from pluto.analytics_overview import fetch_overview_core_metrics
from pluto.scheduled_tasks import start_upload_cleanup_scheduler
from pluto.db import apply_sqlite_pragmas, ensure_indexes
from pluto.errors import register_error_handlers
from pluto.extensions import init_app_extensions, limiter
from pluto.uploads import candidate_display_name, display_upload_filename, store_uploaded_file
from pluto.security_headers import apply_security_headers
from pluto.messages import safe_api_error
from pluto.job_history import fetch_job_centric_history
from pluto.llm_parse import (
    handbook_section_index_from_markdown,
    parse_llm_json_response,
    try_extract_eval_field_previews,
    try_extract_jd_match_preview,
)
from pluto.handbook_intake import (
    build_recruiter_handbook_prompt,
    intake_from_json,
    intake_to_json,
    normalize_intake_payload,
    validate_intake,
)
from pluto.auth_decorators import login_required, role_required
from pluto.blueprints.register import register_blueprints
from pluto.upload_cleanup import cleanup_upload_folder
from pluto.users_db import (
    DATABASE_NAME,
    create_or_update_user,
    filter_data_by_role,
    get_accessible_user_emails,
    get_accessible_users,
    get_user_info,
)


# Suppress specific warnings
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! top_p is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! presence_penalty is not default parameter.")
warnings.filterwarnings("ignore", category=UserWarning, message="WARNING! frequency_penalty is not default parameter.")

# Load environment variables
load_dotenv()

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "hr-knowledge-base"
POLICIES_FOLDER = "HR_docs/"
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
UPLOAD_FOLDER = 'uploads'
# Supported resume upload formats:
# - pdf: normal PDFs (text extraction)
# - docx: Word (python-docx)
# - doc: legacy Word (best-effort via LibreOffice conversion to docx)
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# Google OAuth Configuration
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "").strip().strip('"')
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "").strip().strip('"')
SECRET_KEY = resolve_secret_key()
DEFAULT_ADMIN_EMAIL = os.getenv("DEFAULT_ADMIN_EMAIL", "ritesh.m@peoplelogic.in").strip()

# Model Selection Configuration
MODEL_PROVIDER = os.getenv("MODEL_PROVIDER", "gemini").lower()  # Options: "gemini", "openai", or "groq"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # Options: gpt-4o, gpt-4o-mini, gpt-4-turbo, gpt-3.5-turbo
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "").strip().strip('"')  # For OpenAI-compatible providers (e.g., NVIDIA NIM)
# When Groq/OpenAI fails, call this model on OPENAI_BASE_URL (NVIDIA NIM + Nemotron by default).
FALLBACK_OPENAI_MODEL = os.getenv("FALLBACK_OPENAI_MODEL", "nvidia/nemotron-3-nano-30b-a3b").strip().strip('"')
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")  # Options: gemini-2.5-flash, gemini-2.5-pro, gemini-2.0-flash
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")  # Options: openai/gpt-oss-120b, llama-3.3-70b-versatile, llama-3.1-70b-versatile
GROQ_REASONING_EFFORT = os.getenv("GROQ_REASONING_EFFORT", "high")  # Options: low, medium, high (for reasoning models)
GROQ_MAX_COMPLETION_TOKENS = int(os.getenv("GROQ_MAX_COMPLETION_TOKENS", "2048"))
NVIDIA_REASONING_BUDGET = int(os.getenv("NVIDIA_REASONING_BUDGET", "16384"))
NVIDIA_ENABLE_THINKING = os.getenv("NVIDIA_ENABLE_THINKING", "true").strip().lower() in ("1", "true", "yes", "on")
NVIDIA_TOP_P = float(os.getenv("NVIDIA_TOP_P", "1"))
NVIDIA_TEMPERATURE = float(os.getenv("NVIDIA_TEMPERATURE", "1"))
INFO_BUDDY_MODEL_PROVIDER = os.getenv("INFO_BUDDY_MODEL_PROVIDER", "").strip().strip('"').lower()
INFO_BUDDY_GROQ_MODEL = os.getenv("INFO_BUDDY_GROQ_MODEL", "").strip().strip('"')
INFO_BUDDY_GROQ_REASONING_EFFORT = os.getenv("INFO_BUDDY_GROQ_REASONING_EFFORT", "").strip().strip('"')
HANDBOOK_MODEL_PROVIDER = os.getenv("HANDBOOK_MODEL_PROVIDER", "groq").strip().strip('"').lower()
HANDBOOK_GROQ_MODEL = os.getenv("HANDBOOK_GROQ_MODEL", "").strip().strip('"')
HANDBOOK_GROQ_REASONING_EFFORT = os.getenv("HANDBOOK_GROQ_REASONING_EFFORT", "").strip().strip('"')
HANDBOOK_MAX_COMPLETION_TOKENS = int(os.getenv("HANDBOOK_MAX_COMPLETION_TOKENS", "6144"))
# Cap Groq completion budget on first attempt to reduce TPM "prompt + max_tokens" rejections (overridable).
GROQ_MAX_COMPLETION_TOKENS_CAP = int(os.getenv("GROQ_MAX_COMPLETION_TOKENS_CAP", "4096"))
EVALUATION_MODEL_PROVIDER = os.getenv("EVALUATION_MODEL_PROVIDER", "").strip().strip('"').lower()
EVALUATION_OPENAI_MODEL = os.getenv("EVALUATION_OPENAI_MODEL", "").strip().strip('"')
EVALUATION_GROQ_MODEL = os.getenv("EVALUATION_GROQ_MODEL", "").strip().strip('"')
EVALUATION_GROQ_REASONING_EFFORT = os.getenv("EVALUATION_GROQ_REASONING_EFFORT", "").strip().strip('"')
EVALUATION_MAX_COMPLETION_TOKENS = int(os.getenv("EVALUATION_MAX_COMPLETION_TOKENS", "4096"))
EVALUATION_NVIDIA_REASONING_BUDGET = int(os.getenv("EVALUATION_NVIDIA_REASONING_BUDGET", "4096"))
EVALUATION_NVIDIA_ENABLE_THINKING = os.getenv("EVALUATION_NVIDIA_ENABLE_THINKING", "true").strip().lower() in ("1", "true", "yes", "on")
EVALUATION_NVIDIA_TOP_P = float(os.getenv("EVALUATION_NVIDIA_TOP_P", str(NVIDIA_TOP_P)))
EVALUATION_NVIDIA_TEMPERATURE = float(os.getenv("EVALUATION_NVIDIA_TEMPERATURE", str(NVIDIA_TEMPERATURE)))
EVALUATION_CACHE_ENABLED = os.getenv("EVALUATION_CACHE_ENABLED", "true").strip().lower() in ("1", "true", "yes", "on")
EVALUATION_CACHE_TTL_SECONDS = int(os.getenv("EVALUATION_CACHE_TTL_SECONDS", "1800"))
EVALUATION_CACHE_MAX_ENTRIES = int(os.getenv("EVALUATION_CACHE_MAX_ENTRIES", "200"))
FAST_EVAL_MODE = os.getenv("FAST_EVAL_MODE", "false").strip().lower() in ("1", "true", "yes", "on")

# Add this dictionary after imports
ACRONYM_MAP = {
    "wfh": "work from home policy",
    "pto": "paid time off policy",
    "loa": "leave of absence policy",
    "nda": "non-disclosure agreement",
    "od": "on duty policy",
    "hrbp": "human resources business partner",
    "kra": "KRA Policy - Promoting Transparency",
    "regularization": "Time change Request/ Regularization",
    "regularisation": "Time change Request/ Regularization",
    "posh": "Policy On Prevention of Sexual Harassment",
    "appraisal": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "promotion": "PERFORMANCE APPRAISAL & PROMOTION POLICY",
    "prep": "Performance Review & Enhancement Program",
    "Grade": "GRADE STRUCTURE & FLEXIBILITY",
    "leave": "LEAVE POLICY",
    "nda": "Non Compete and Non Disclosure",
    "Office timings": "Office Timing and Attendance Policy",
    "pet": "pet policy",
    "sprint": "Weekly Sprint Policy",
    "work ethics": "WORK PLACE ETHICS"
}

# Standard behavioral questions
QUICK_CHECKS = [
    "Are you willing to relocate if applicable?",
    "What is your notice period?",
    "Can you provide details about your current organization?",
    "Please describe your current role and responsibilities.",
    "What is your current CTC (Cost to Company)?",
    "What is your expected CTC?",
    "What is your educational background?",
    "Can you describe any significant projects you've worked on?",
    "Are there any specific client requirements you want to discuss?",
    "Do you have references from colleagues who might be interested in opportunities with us?"
]
############################################################################################


# Prompt templates
input_prompt_template = """

Act as a senior recruiter with 30+ years of experience.

You MUST evaluate the resume strictly against the job description (JD).
Your reasoning MUST be based ONLY on explicit evidence written in the resume.

IMPORTANT: This is a recruiter-grade evaluation, NOT an ATS keyword scan.

GENERAL PRINCIPLES (APPLY ALWAYS):
- Do NOT infer, assume, or guess skills.
- Do NOT reward buzzwords without demonstrated usage, responsibility, or outcomes.
- Do NOT hallucinate missing skills or capabilities.
- Use recruiter judgment to distinguish between:
  (a) truly missing skills
  (b) explicitly stated equivalent or functionally identical experience.

CRITICAL DISTINCTION (VERY IMPORTANT):
- Explicitly equivalent enterprise evidence COUNTS as MATCHED.
- Adjacent, implied, or loosely related experience does NOT count.

Examples:
- Kafka usage = VALID for MSK (Kafka) unless JD explicitly requires MSK administration.
- MySQL / Oracle / PostgreSQL on AWS = VALID for RDS unless JD explicitly requires RDS operations.
- CI/CD ownership, Docker, cloud deployments = VALID evidence of DevOps collaboration.
- Architecture design, solution reviews, platform ownership = VALID evidence of technical specifications.
- Leadership roles = VALID evidence of mentoring and code review when explicitly stated.

This evaluation MUST work for ANY role (technical or non-technical) and ANY model size (large or small).
Therefore, follow these STRICT rules:

STRICT EVALUATION RULES (APPLY ALL):
1. If the JD lists a MUST-HAVE skill AND there is NO direct or equivalent evidence → treat it as NOT MATCHED.
2. Equivalent enterprise-grade evidence COUNTS if it serves the same functional purpose.
3. Related, adjacent, or implied experience does NOT count.
   (Example: “MQ development” ≠ “MQ administration”. “Linux exposure” ≠ “Linux admin”.)
4. No assumptions. No optimism. No guessing.
5. If more than 40% of MUST-HAVE skills are NOT MATCHED → final verdict CANNOT be shortlist.
6. For non-technical roles, evaluate ONLY achievements, outcomes, metrics, stakeholder scope, and behavioral competencies.
7. For senior roles, require evidence of leadership, ownership, decision-making, mentoring, and strategic or architectural impact.
8. If the JD is specialized, the resume MUST show direct evidence of that specialization.
9. Apply overqualification analysis ONLY when experience, title, or scope significantly exceeds the JD.
10. Penalize weak resumes aggressively. Do NOT inflate scores for verbosity, repetition, or buzzwords.
11. If evidence is ambiguous → treat as NOT MATCHED.

CERTIFICATIONS:
- If certifications are NOT mentioned in the JD → set "Certification Match" = null AND exclude it from scoring.
- If JD requires certifications:
   • 100 if candidate has ALL required certifications
   • 0 if ANY required certification is missing
- Extra certifications DO NOT increase score beyond 100.

EDUCATION RULE:
- If JD says “Any Postgraduate” → postgraduate is OPTIONAL.
- Penalize education ONLY if the mandatory minimum qualification is missing.

JD MATCH CALCULATION:
- INCLUDE only applicable match factors.
- Weight all applicable factors equally.
- EXCLUDE Certification Match when set to null.
- Overall JD Match must reflect recruiter realism, not keyword completeness.

OUTPUT CONSTRAINTS:
- Your output MUST be strictly valid JSON.
- The structure, keys, and format MUST match exactly.
- Do NOT add, remove, rename, or reorder fields.
- Do NOT include explanations, markdown, or text outside the JSON.

### Output:
Return a valid JSON object ONLY. The JSON object MUST have the following keys:

{{
  "JD Match": "85%",
  "MissingKeywords": [...],
  "Profile Summary": "...",
  "Over/UnderQualification Analysis": "...",
  "Match Factors": {{
    "Skills Match": 0-100,
    "Experience Match": 0-100,
    "Education Match": 0-100,
    "Industry Knowledge": 0-100,
    "Certification Match": number or null
  }},
  "Reasoning": "...",
  "Candidate Fit Analysis": {{
    "Dimension Evaluation": [
      {{
        "Dimension": "...",
        "Evaluation": "✅ Strong / ⚠️ Moderate / ❌ Weak",
        "Recruiter Comments": "..."
      }}
    ],
    "Risk and Gaps": [
      {{
        "Area": "...",
        "Risk": "...",
        "Recruiter Strategy": "..."
      }}
    ] or null,
    "Recommendation": {{
      "Verdict": "❌ Not Recommended / ⚠️ Conditional Shortlist / ✅ Strong Shortlist",
      "Fit Level": "High / Medium / Low",
      "Rationale": "..."
    }},
    "Recruiter Narrative": "..."
  }}
}}

Do NOT include any additional text outside the JSON object.

---
Resume:
{resume_text}

JOB DESCRIPTION:
{job_description}

ADDITIONAL CONTEXT (if any):
{additional_context_block}

"""

interview_questions_prompt = """
You are an experienced recruiter. Generate interview questions based ONLY on information in the resume and job description. 
Do NOT assume skills.
Do NOT ask questions about skills not present in either JD or resume.

**CRITICAL**: Return ONLY valid JSON. No explanations, no reasoning, no additional text. Start with {{ and end with }}

Output ONLY this JSON:

{{
  "TechnicalQuestions": [...10 questions...],
  "NonTechnicalQuestions": [...10 questions...]
}}

Rules:
- Technical questions MUST reference only the technologies, responsibilities, and domains in the resume AND JD.
- Non-technical questions must assess behavior, ownership, teamwork, leadership, or role fit.
- No generic or irrelevant questions.
- No text outside JSON.

RESUME:
{resume_text}

JD:
{job_description}

PROFILE SUMMARY:
{profile_summary}


"""

job_stability_prompt = """
You are an HR analytics expert. Evaluate job stability strictly based on dates and job history written in the resume. 
Do NOT assume missing dates. 
Do NOT infer unstated durations.

Return ONLY this JSON:

{{
  "IsStable": true/false,
  "AverageJobTenure": "...",
  "JobCount": number,
  "StabilityScore": 0-100,
  "ReasoningExplanation": "...",
  "RiskLevel": "Low / Medium / High"
}}

RESUME:
{resume_text}

"""


# Add career progression prompt template after other prompt templates
career_prompt = """
You are an expert HR analyst. Analyze career progression strictly using evidence from the resume. 
Do NOT infer unstated roles or responsibilities.

Return ONLY:

{
  "red_flags": [...],
    "reasoning": [
    "...",
    "...",
    "..."
    ]
}

RESUME:
{resume_text}

"""

#######################################################
# Recruiter Handbook Prompt Template
recruiter_handbook_prompt = """
SYSTEM:
You are an expert technical recruiter and talent evaluator specializing in AI, analytics, product, and consulting roles. 
Your task is to generate a detailed recruiter-style evaluation and fit analysis based on a Job Description (JD) and a Candidate Resume. 
You must think and write like a senior talent partner at a top-tier consulting firm (Fractal, Deloitte, BCG, etc.) — structured, insightful, data-driven, and nuanced. 
Your output will form a Recruiter Handbook that helps internal recruiters, interviewers, and hiring managers make data-driven shortlisting decisions.

---

INSTRUCTIONS:

You will be provided with:
- JOB_DESCRIPTION_TEXT: {job_description}
- CANDIDATE_RESUME_TEXT: {resume_text}

Generate the following sections in professional markdown format:

### 1️⃣ JD Snapshot
Summarize the role in 5–6 lines:
- Role title, level, domain/vertical
- Key skills required
- Nature of role (hands-on, consulting-led, AI-driven, etc.)
- Success indicators

---

### 2️⃣ Candidate Summary
Summarize the candidate's profile in 5–6 lines:
- Education and total experience
- Functional & technical focus areas
- Domain experience (industries)
- Career trajectory highlights
- Distinguishing achievements

---

### 3️⃣ Fit Matrix (Comparison Table)
Create a markdown table with columns:  
**Dimension | JD Expectation | Candidate Evidence | Rating (1–5) | Comment**

Include these dimensions:
- Domain Expertise  
- Technical / AI / Cloud Engineering Depth  
- Consulting Gravitas / CXO Advisory Experience  
- Innovation & IP / Emerging Tech Thought Leadership  
- Hands-on IC Credibility  
- Project & Delivery Leadership  
- Business Acumen & Commercial Awareness  
- Culture / Communication Fit  

Rating legend: 5=Excellent | 3=Average | 1=Weak

---

### 4️⃣ Detailed Fit Commentary
Write 5–8 paragraphs of nuanced recruiter commentary:
- Where the candidate aligns strongly  
- Where they are weak or untested  
- How their consulting/technical balance fits  
- Domain alignment and value they could bring  
- Potential role re-alignment if not an exact fit  
Tone: confident, analytical, and evidence-driven.

---

### 5️⃣ Red Flags & Mitigation
List 3–5 potential red flags (if any) and ways to mitigate or probe during interview.

---

### 6️⃣ Interview Focus Areas
List 8–10 recommended interview questions, grouped under:
- Technical/Engineering depth  
- Consulting & stakeholder management  
- Domain expertise  
- AI/GenAI awareness  
- Leadership & delivery

---

### 7️⃣ Recruiter Summary & Pitch
Write 2 short paragraphs (recruiter voice) that summarize:
- Why this candidate could be compelling to the client
- How to position them internally
Use persuasive, professional tone for internal submission.

---

### 8️⃣ Final Verdict
Provide:
- **Fit Verdict:** Strong Shortlist / Conditional / Reconsider / Reject  
- **Fit Score (0–100%):**  
- **Confidence (High / Medium / Low):**  
- **Best-fit Domain (Retail / FSI / TMT / HLS / Other):**

---

**Output Formatting:**
- Write the recruiter handbook in professional markdown (for rendering in a web UI).  
- Keep tone polished, confident, and analytical — like a top-tier recruiter brief.  
- Be specific, never generic. Use evidence-based phrases like:
  - "Shows strong multi-client consulting maturity."
  - "Demonstrates architectural thought leadership but lacks GenAI delivery exposure."
  - "Consulting gravitas evident from CIO-level advisory roles."
- Do NOT include markdown code block markers in your response.
"""

def _gemini_api_key():
    """Google AI Studio / Vertex-style keys (support both env names used in the wild)."""
    return (os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or "").strip()


def ensure_gemini_model():
    """
    Return a live GenerativeModel instance, initializing on demand when possible.
    Startup skips Gemini when using OpenAI/Groq only; lazy init enables a real fallback.
    """
    global gemini_model, genai
    if gemini_model is not None:
        return gemini_model
    key = _gemini_api_key()
    if not key:
        return None
    try:
        import google.generativeai as genai_mod  # type: ignore

        genai = genai_mod
        genai_mod.configure(api_key=key)
        gemini_model = genai_mod.GenerativeModel(GEMINI_MODEL)
        logging.info(f"Lazily initialized Gemini model: {GEMINI_MODEL}")
        return gemini_model
    except Exception as e:
        logging.warning(f"Lazy Gemini initialization failed: {e}")
        return None


def _groq_error_looks_like_token_quota(err):
    """True for Groq 413 / TPM-style rate limits where shrinking max_completion_tokens may help."""
    if not isinstance(err, APIStatusError):
        return False
    code = getattr(err, "status_code", None)
    body = str(getattr(err, "body", "") or "")
    msg = str(err).lower()
    if code == 413:
        return True
    if "rate_limit_exceeded" in body.lower() or "rate_limit_exceeded" in msg:
        return True
    if "tokens per minute" in body.lower() or "tokens per minute" in msg:
        return True
    return False


def _groq_chat_create_with_token_backoff(params, token_param_name="max_completion_tokens", max_attempts=4):
    """
    Call Groq chat.completions.create; on TPM-style errors shrink max_completion_tokens and retry.
    Mutates params[token_param_name] between attempts.
    """
    mt = params.get(token_param_name)
    if mt is None:
        raise ValueError("Groq params missing max completion token field")
    last_exc = None
    for attempt in range(max_attempts):
        try:
            return groq_client.chat.completions.create(**params)
        except APIStatusError as g_err:
            last_exc = g_err
            if _groq_error_looks_like_token_quota(g_err) and attempt < max_attempts - 1:
                new_mt = max(768, int(mt * 0.63))
                if new_mt >= mt:
                    new_mt = max(768, mt - 1024)
                mt = new_mt
                params[token_param_name] = mt
                logging.warning(
                    "Groq TPM/token-quota error; retrying with %s=%s (attempt %s/%s): %s",
                    token_param_name,
                    mt,
                    attempt + 2,
                    max_attempts,
                    g_err,
                )
                time.sleep(1.0 * (attempt + 1))
                continue
            raise
    if last_exc:
        raise last_exc
    raise RuntimeError("Groq chat create failed without a raised exception")


# Initialize Gemini model only when needed (avoids deprecated package warning on non-Gemini runs)
gemini_model = None
if MODEL_PROVIDER == "gemini" or not (OPENAI_API_KEY or GROQ_API_KEY):
    try:
        import google.generativeai as genai  # type: ignore

        _gk = _gemini_api_key()
        if _gk:
            genai.configure(api_key=_gk)
            gemini_model = genai.GenerativeModel(GEMINI_MODEL)
    except Exception as e:
        logging.warning(f"Gemini client initialization skipped/failed: {e}")

# Initialize OpenAI client
openai_client = None
if OPENAI_API_KEY:
    try:
        openai_kwargs = {"api_key": OPENAI_API_KEY}
        if OPENAI_BASE_URL:
            openai_kwargs["base_url"] = OPENAI_BASE_URL
        openai_client = OpenAI(**openai_kwargs)
    except Exception as e:
        print(f"⚠️  WARNING: Failed to initialize OpenAI client: {e}")
        if MODEL_PROVIDER == "openai":
            print("⚠️  Falling back to Gemini")
elif MODEL_PROVIDER == "openai" and not OPENAI_API_KEY:
    print("⚠️  WARNING: MODEL_PROVIDER is set to 'openai' but OPENAI_API_KEY is missing!")
    print("⚠️  Using Gemini instead. Add OPENAI_API_KEY to your .env file to use OpenAI.")

# Initialize Groq client
groq_client = None
if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
    except Exception as e:
        print(f"⚠️  WARNING: Failed to initialize Groq client: {e}")
        if MODEL_PROVIDER == "groq":
            print("⚠️  Falling back to Gemini")
elif MODEL_PROVIDER == "groq" and not GROQ_API_KEY:
    print("⚠️  WARNING: MODEL_PROVIDER is set to 'groq' but GROQ_API_KEY is missing!")
    print("⚠️  Using Gemini instead. Add GROQ_API_KEY to your .env file to use Groq.")

# Legacy variable for backward compatibility
model = gemini_model


def _nemotron_fallback_is_configured():
    """True when OpenAI SDK can reach NVIDIA NIM (needs OPENAI_API_KEY + OPENAI_BASE_URL)."""
    return bool(openai_client) and bool((OPENAI_BASE_URL or "").strip())


# ========================================
# UNIFIED MODEL ABSTRACTION LAYER
# ========================================

class UnifiedModelResponse:
    """Wrapper class to standardize responses from different models"""
    def __init__(self, text):
        self.text = text
        self.output_text = text  # For compatibility with your sample code

def get_info_buddy_provider_settings():
    """Resolve Info Buddy provider/model overrides from env."""
    provider = (INFO_BUDDY_MODEL_PROVIDER or MODEL_PROVIDER).lower()
    groq_model = INFO_BUDDY_GROQ_MODEL or GROQ_MODEL
    groq_reasoning_effort = INFO_BUDDY_GROQ_REASONING_EFFORT or GROQ_REASONING_EFFORT
    return provider, groq_model, groq_reasoning_effort

def get_handbook_provider_settings():
    """Resolve handbook provider/model overrides from env."""
    provider = (HANDBOOK_MODEL_PROVIDER or MODEL_PROVIDER).lower()
    groq_model = HANDBOOK_GROQ_MODEL or GROQ_MODEL
    groq_reasoning_effort = HANDBOOK_GROQ_REASONING_EFFORT or GROQ_REASONING_EFFORT
    return provider, groq_model, groq_reasoning_effort

def get_evaluation_provider_settings():
    """Resolve resume-evaluation provider/model overrides from env."""
    provider = (EVALUATION_MODEL_PROVIDER or MODEL_PROVIDER).lower()
    if provider == "groq":
        model = EVALUATION_GROQ_MODEL or GROQ_MODEL
        groq_reasoning_effort = EVALUATION_GROQ_REASONING_EFFORT or GROQ_REASONING_EFFORT
    elif provider == "openai":
        model = EVALUATION_OPENAI_MODEL or OPENAI_MODEL
        groq_reasoning_effort = GROQ_REASONING_EFFORT
    else:
        model = None
        groq_reasoning_effort = GROQ_REASONING_EFFORT
    return provider, model, groq_reasoning_effort

def get_evaluation_generation_overrides():
    """Build generation overrides for resume evaluation path."""
    provider, model, groq_reasoning_effort = get_evaluation_provider_settings()
    overrides = {
        "provider_override": provider,
        "model_override": model,
        "groq_reasoning_effort_override": groq_reasoning_effort,
        "max_completion_tokens_override": EVALUATION_MAX_COMPLETION_TOKENS
    }
    if provider == "openai" and model and "nemotron" in model.lower():
        overrides["temperature_override"] = EVALUATION_NVIDIA_TEMPERATURE
        overrides["top_p_override"] = EVALUATION_NVIDIA_TOP_P
        overrides["extra_body_override"] = {
            "reasoning_budget": EVALUATION_NVIDIA_REASONING_BUDGET,
            "chat_template_kwargs": {"enable_thinking": EVALUATION_NVIDIA_ENABLE_THINKING}
        }
    return overrides

_EVAL_RESPONSE_CACHE = {}

def _cache_get(key):
    if not EVALUATION_CACHE_ENABLED:
        return None
    hit = _EVAL_RESPONSE_CACHE.get(key)
    if not hit:
        return None
    ts, value = hit
    if (time.time() - ts) > EVALUATION_CACHE_TTL_SECONDS:
        _EVAL_RESPONSE_CACHE.pop(key, None)
        return None
    return value

def _cache_set(key, value):
    if not EVALUATION_CACHE_ENABLED:
        return
    _EVAL_RESPONSE_CACHE[key] = (time.time(), value)
    while len(_EVAL_RESPONSE_CACHE) > EVALUATION_CACHE_MAX_ENTRIES:
        oldest_key = next(iter(_EVAL_RESPONSE_CACHE))
        _EVAL_RESPONSE_CACHE.pop(oldest_key, None)


def generate_content_unified(
    prompt,
    stream=False,
    provider_override=None,
    model_override=None,
    groq_reasoning_effort_override=None,
    max_completion_tokens_override=None,
    temperature_override=None,
    top_p_override=None,
    extra_body_override=None,
    _nemotron_fallback_active=False,
):
    """
    Unified function to generate content from Gemini, OpenAI, or Groq.

    On Groq/OpenAI failure: tries NVIDIA Nemotron on NIM (OPENAI_BASE_URL + FALLBACK_OPENAI_MODEL),
    then Gemini if configured.

    Args:
        prompt (str): The prompt to send to the model
        stream (bool): Whether to stream the response (for real-time output)
        _nemotron_fallback_active (bool): Internal guard to avoid retry loops (do not set from app code).

    Returns:
        UnifiedModelResponse or generator: Response object with .text attribute
    """
    try:
        selected_provider = (provider_override or MODEL_PROVIDER).lower()
        selected_openai_model = model_override if selected_provider == "openai" and model_override else OPENAI_MODEL
        selected_groq_model = model_override if selected_provider == "groq" and model_override else GROQ_MODEL
        selected_groq_reasoning_effort = groq_reasoning_effort_override or GROQ_REASONING_EFFORT

        requested_max_tokens = max_completion_tokens_override or 16384

        if selected_provider == "openai" and openai_client:
            # OpenAI API call
            # Use max_completion_tokens for newer models (gpt-4o, gpt-4o-mini, gpt-5, o1, o3, etc.)
            # Use max_tokens for older models (gpt-3.5-turbo, gpt-4-turbo)
            # Newer models: gpt-4o*, gpt-5*, o1*, o3*
            uses_new_api = any(x in selected_openai_model.lower() for x in ["gpt-4o", "gpt-5", "o1-", "o3-"])
            token_param = "max_completion_tokens" if uses_new_api else "max_tokens"
            
            # Some newer models (gpt-5*, o1*, o3*) don't support custom temperature
            # Only use temperature for models that support it (gpt-4o and older support it)
            supports_temperature = not any(x in selected_openai_model.lower() for x in ["gpt-5", "o1-", "o3-"])
            is_nemotron = "nemotron" in selected_openai_model.lower()
            
            if stream:
                # Streaming response for OpenAI
                params = {
                    "model": selected_openai_model,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    "stream": True,
                    token_param: requested_max_tokens
                }
                if is_nemotron:
                    params["temperature"] = temperature_override if temperature_override is not None else NVIDIA_TEMPERATURE
                    params["top_p"] = top_p_override if top_p_override is not None else NVIDIA_TOP_P
                    params["extra_body"] = extra_body_override if extra_body_override is not None else {
                        "reasoning_budget": NVIDIA_REASONING_BUDGET,
                        "chat_template_kwargs": {"enable_thinking": NVIDIA_ENABLE_THINKING}
                    }
                elif supports_temperature:
                    params["temperature"] = temperature_override if temperature_override is not None else 0.7
                if top_p_override is not None and not is_nemotron:
                    params["top_p"] = top_p_override
                    
                    
                response = openai_client.chat.completions.create(**params)
                
                def stream_generator():
                    for chunk in response:
                        choices = getattr(chunk, "choices", None) or []
                        if not choices:
                            continue
                        delta = getattr(choices[0], "delta", None)
                        content = getattr(delta, "content", None) if delta is not None else None
                        if content:
                            yield content
                
                return stream_generator()
            else:
                # Non-streaming response for OpenAI
                params = {
                    "model": selected_openai_model,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    token_param: requested_max_tokens
                }
                if is_nemotron:
                    params["temperature"] = temperature_override if temperature_override is not None else NVIDIA_TEMPERATURE
                    params["top_p"] = top_p_override if top_p_override is not None else NVIDIA_TOP_P
                    params["extra_body"] = extra_body_override if extra_body_override is not None else {
                        "reasoning_budget": NVIDIA_REASONING_BUDGET,
                        "chat_template_kwargs": {"enable_thinking": NVIDIA_ENABLE_THINKING}
                    }
                elif supports_temperature:
                    params["temperature"] = temperature_override if temperature_override is not None else 0.7
                if top_p_override is not None and not is_nemotron:
                    params["top_p"] = top_p_override
                    
                response = openai_client.chat.completions.create(**params)
                logging.info(f"OpenAI response model: {response.model}, choices count: {len(response.choices)}")
                
                # Check response structure
                if not response.choices:
                    logging.error("OpenAI response has no choices!")
                    raise Exception("OpenAI returned empty choices")
                
                message = response.choices[0].message
                logging.info(f"Message role: {message.role}, has content: {message.content is not None}")
                
                # Some models might have refusal or other fields
                if hasattr(message, 'refusal') and message.refusal:
                    logging.error(f"OpenAI refused: {message.refusal}")
                    raise Exception(f"OpenAI refused to respond: {message.refusal}")
                
                content = message.content
                logging.info(f"Content length: {len(content) if content else 0}, type: {type(content)}")
                
                if not content:
                    logging.error(f"OpenAI returned empty content!")
                    logging.error(f"Full message: {message}")
                    logging.error(f"Finish reason: {response.choices[0].finish_reason}")
                    raise Exception(f"OpenAI returned empty content. Finish reason: {response.choices[0].finish_reason}")
                
                return UnifiedModelResponse(content)
        
        elif selected_provider == "groq" and groq_client:
            # Groq API call (same structure as OpenAI, reasoning models supported)
            # Groq uses max_completion_tokens for all models
            token_param = "max_completion_tokens"
            
            # Reasoning models (gpt-oss, o1, o3) don't support custom temperature
            is_reasoning_model = any(x in selected_groq_model.lower() for x in ["gpt-oss", "o1-", "o3-"])
            supports_temperature = not is_reasoning_model

            groq_requested = max_completion_tokens_override or GROQ_MAX_COMPLETION_TOKENS
            _eff_cap = max(256, GROQ_MAX_COMPLETION_TOKENS_CAP)
            groq_mt = min(groq_requested, _eff_cap)
            if groq_mt < groq_requested:
                logging.info(
                    "Capping Groq %s %s -> %s (effective cap=%s from GROQ_MAX_COMPLETION_TOKENS_CAP)",
                    token_param,
                    groq_requested,
                    groq_mt,
                    _eff_cap,
                )
            
            if stream:
                # Streaming response for Groq
                params = {
                    "model": selected_groq_model,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    "stream": True,
                    token_param: groq_mt
                }
                
                # Add temperature only for non-reasoning models
                if supports_temperature:
                    params["temperature"] = 0.7
                
                # Add reasoning_effort for reasoning models
                if is_reasoning_model:
                    params["reasoning_effort"] = selected_groq_reasoning_effort
                
                response = _groq_chat_create_with_token_backoff(params, token_param)
                
                def stream_generator():
                    for chunk in response:
                        choices = getattr(chunk, "choices", None) or []
                        if not choices:
                            continue
                        delta = getattr(choices[0], "delta", None)
                        content = getattr(delta, "content", None) if delta is not None else None
                        if content:
                            yield content
                
                return stream_generator()
            else:
                # Non-streaming response for Groq
                params = {
                    "model": selected_groq_model,
                    "messages": [
                        {"role": "system", "content": "You are an expert HR analyst and technical recruiter."},
                        {"role": "user", "content": prompt}
                    ],
                    token_param: groq_mt
                }
                
                # Add temperature only for non-reasoning models
                if supports_temperature:
                    params["temperature"] = 0.7
                
                # Add reasoning_effort for reasoning models
                if is_reasoning_model:
                    params["reasoning_effort"] = selected_groq_reasoning_effort
                
                response = _groq_chat_create_with_token_backoff(params, token_param)
                logging.info(f"Groq response model: {response.model}, choices count: {len(response.choices)}")

                # Check response structure
                if not response.choices:
                    logging.error("Groq response has no choices!")
                    raise Exception("Groq returned empty choices")
                
                message = response.choices[0].message
                logging.info(f"Message role: {message.role}, has content: {message.content is not None}")
                
                content = message.content
                logging.info(f"Content length: {len(content) if content else 0}, type: {type(content)}")
                
                if not content:
                    logging.error(f"Groq returned empty content!")
                    logging.error(f"Full message: {message}")
                    logging.error(f"Finish reason: {response.choices[0].finish_reason}")
                    raise Exception(f"Groq returned empty content. Finish reason: {response.choices[0].finish_reason}")
                
                return UnifiedModelResponse(content)
        
        else:
            # Gemini API call (default)
            gm = ensure_gemini_model()
            if gm is None:
                raise RuntimeError(
                    "Gemini is not configured (set GEMINI_API_KEY or GOOGLE_API_KEY) or provider "
                    "is set to gemini without a valid key. OpenAI/Groq are also unavailable for this request."
                )
            if stream:
                # Streaming response for Gemini
                response = gm.generate_content(prompt, stream=True)
                return response  # Gemini already returns a generator
            else:
                # Non-streaming response for Gemini
                response = gm.generate_content(prompt)
                return response  # Gemini response object already has .text
    
    except Exception as e:
        logging.error(f"Error in generate_content_unified: {str(e)}")
        if selected_provider in ["openai", "groq"]:
            fb_model = (FALLBACK_OPENAI_MODEL or "nvidia/nemotron-3-nano-30b-a3b").strip()
            primary_openai_model = (
                (model_override if selected_provider == "openai" and model_override else OPENAI_MODEL) or ""
            ).strip()
            if (
                not _nemotron_fallback_active
                and _nemotron_fallback_is_configured()
                and primary_openai_model.lower() != fb_model.lower()
            ):
                try:
                    logging.warning(
                        "%s failed; falling back to NVIDIA Nemotron (%s on %s)",
                        selected_provider.upper(),
                        fb_model,
                        OPENAI_BASE_URL,
                    )
                    return generate_content_unified(
                        prompt,
                        stream=stream,
                        provider_override="openai",
                        model_override=fb_model,
                        max_completion_tokens_override=max_completion_tokens_override,
                        temperature_override=NVIDIA_TEMPERATURE,
                        top_p_override=NVIDIA_TOP_P,
                        extra_body_override={
                            "reasoning_budget": NVIDIA_REASONING_BUDGET,
                            "chat_template_kwargs": {"enable_thinking": NVIDIA_ENABLE_THINKING},
                        },
                        _nemotron_fallback_active=True,
                    )
                except Exception as ne:
                    logging.error("NVIDIA Nemotron fallback failed: %s", ne, exc_info=True)
            elif not _nemotron_fallback_active and not _nemotron_fallback_is_configured():
                logging.warning(
                    "Nemotron fallback skipped: set OPENAI_API_KEY and OPENAI_BASE_URL (NVIDIA NIM) "
                    "to enable FALLBACK_OPENAI_MODEL (%s).",
                    fb_model,
                )
            gm = ensure_gemini_model()
            if gm is not None:
                logging.warning(f"{selected_provider.upper()} failed, falling back to Gemini")
                try:
                    if stream:
                        return gm.generate_content(prompt, stream=True)
                    else:
                        return gm.generate_content(prompt)
                except Exception as fe:
                    logging.error(f"Gemini fallback failed: {fe}")
                    raise e from fe
            logging.error(
                "No further fallbacks: Nemotron unavailable or failed, and no GEMINI_API_KEY/GOOGLE_API_KEY. "
                "Fix the primary provider error or configure NVIDIA NIM + Gemini."
            )
        raise

# Log which model is being used
print("=" * 60)
print(f"🤖 Model Provider Configuration: {MODEL_PROVIDER.upper()}")
if MODEL_PROVIDER == "openai" and openai_client:
    print(f"✅ ACTUALLY USING: OpenAI Model: {OPENAI_MODEL}")
elif MODEL_PROVIDER == "openai" and not openai_client:
    print(f"⚠️  ACTUALLY USING: Gemini Model: {GEMINI_MODEL} (OpenAI not available)")
elif MODEL_PROVIDER == "groq" and groq_client:
    print(f"✅ ACTUALLY USING: Groq Model: {GROQ_MODEL}")
    if "gpt-oss" in GROQ_MODEL.lower() or "o1-" in GROQ_MODEL.lower() or "o3-" in GROQ_MODEL.lower():
        print(f"   Reasoning Effort: {GROQ_REASONING_EFFORT}")
elif MODEL_PROVIDER == "groq" and not groq_client:
    print(f"⚠️  ACTUALLY USING: Gemini Model: {GEMINI_MODEL} (Groq not available)")
else:
    print(f"✅ ACTUALLY USING: Gemini Model: {GEMINI_MODEL}")
print("=" * 60)

logging.info(f"🤖 Model Provider: {MODEL_PROVIDER.upper()}")
if MODEL_PROVIDER == "openai":
    logging.info(f"📦 Using OpenAI Model: {OPENAI_MODEL}")
elif MODEL_PROVIDER == "groq":
    logging.info(f"📦 Using Groq Model: {GROQ_MODEL}")
else:
    logging.info(f"📦 Using Gemini Model: {GEMINI_MODEL}")

info_buddy_provider, info_buddy_groq_model, info_buddy_groq_reasoning_effort = get_info_buddy_provider_settings()
logging.info(f"🤖 Info Buddy Provider: {info_buddy_provider.upper()}")
if info_buddy_provider == "groq":
    logging.info(f"📦 Info Buddy Groq Model: {info_buddy_groq_model}")
    if any(x in info_buddy_groq_model.lower() for x in ["gpt-oss", "o1-", "o3-"]):
        logging.info(f"🧠 Info Buddy Groq Reasoning Effort: {info_buddy_groq_reasoning_effort}")

eval_provider, eval_model, eval_groq_reasoning_effort = get_evaluation_provider_settings()
if eval_provider == "groq":
    logging.info(f"🤖 Resume Evaluation Provider: GROQ")
    logging.info(f"📦 Resume Evaluation Groq Model: {eval_model}")
    if any(x in (eval_model or "").lower() for x in ["gpt-oss", "o1-", "o3-"]):
        logging.info(f"🧠 Resume Evaluation Groq Reasoning Effort: {eval_groq_reasoning_effort}")
elif eval_provider == "openai":
    logging.info(f"🤖 Resume Evaluation Provider: OPENAI")
    logging.info(f"📦 Resume Evaluation OpenAI Model: {eval_model}")
    if eval_model and "nemotron" in eval_model.lower():
        logging.info(f"🎯 Eval max completion tokens: {EVALUATION_MAX_COMPLETION_TOKENS}")
        logging.info(f"🧠 Eval reasoning budget: {EVALUATION_NVIDIA_REASONING_BUDGET}")
        logging.info(f"💭 Eval thinking enabled: {EVALUATION_NVIDIA_ENABLE_THINKING}")
        logging.info(f"🌡️ Eval temperature/top_p: {EVALUATION_NVIDIA_TEMPERATURE}/{EVALUATION_NVIDIA_TOP_P}")
else:
    logging.info(f"🤖 Resume Evaluation Provider: GEMINI")
    logging.info(f"📦 Resume Evaluation Gemini Model: {GEMINI_MODEL}")
logging.info(
    f"⚡ Eval cache: {'ON' if EVALUATION_CACHE_ENABLED else 'OFF'} "
    f"(ttl={EVALUATION_CACHE_TTL_SECONDS}s, max_entries={EVALUATION_CACHE_MAX_ENTRIES})"
)

# ========================================


# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SECRET_KEY'] = SECRET_KEY
apply_flask_config(app)
init_app_extensions(app)
register_error_handlers(app)
apply_security_headers(app)
_IO_EXECUTOR = ThreadPoolExecutor(max_workers=4, thread_name_prefix="pluto-io")


@app.context_processor
def _inject_dev_asset_version():
    """Bust browser cache for static JS/CSS during local dev (per page load)."""
    if is_production():
        version = os.getenv("PLUTO_ASSET_VERSION", "1")
    else:
        version = str(int(time.time()))
    return {"asset_v": version}


asgi_app = WsgiToAsgi(app)

# Initialize Google OAuth
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    },
    # Explicitly set the callback route to match Google Cloud Console configuration
    authorize_callback='authorize'
)

register_blueprints(app, google=google)


@app.cli.command("cleanup-uploads")
def cleanup_uploads_command():
    """Remove stale upload files (see scripts/cleanup_uploads.py)."""
    days = int(os.getenv("UPLOAD_CLEANUP_MAX_AGE_DAYS", "7"))
    dry = os.getenv("UPLOAD_CLEANUP_DRY_RUN", "").lower() in ("1", "true", "yes")
    folder = app.config.get("UPLOAD_FOLDER", UPLOAD_FOLDER)
    removed, skipped = cleanup_upload_folder(folder, max_age_days=days, dry_run=dry)
    print(f"{'Would remove' if dry else 'Removed'} {len(removed)} file(s); skipped {len(skipped)} in DB.")


# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize NLTK
try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    nltk.download("punkt", quiet=True)


# Initialize Groq LLM
llm = ChatGroq(
    groq_api_key=GROQ_API_KEY,
    # model_name="mixtral-8x7b-32768",  # This generates long text,  max_tokens=4096
    # model_name=   "llama-3.1-8b-instant",#"qwen-2.5-32b", #"deepseek-r1-distill-qwen-32b",
    model_name = "qwen/qwen3-32b",
    temperature=0.377,
    max_tokens=2048,   #4096
    top_p=0.95,
    presence_penalty=0.1,
    frequency_penalty=0.1
)

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)
index_name = PINECONE_INDEX_NAME
if index_name not in pc.list_indexes().names():
    pc.create_index(
        name=index_name,
        dimension=384,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )
index = pc.Index(index_name)

# Initialize embeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Initialize vector store
vectorstore = None
try:
    from langchain_pinecone import PineconeVectorStore as NewPineconeVectorStore
    vectorstore = NewPineconeVectorStore(
        index=index,
        embedding=embeddings,
        text_key="text"
    )
    logging.info("✅ Using new langchain-pinecone vectorstore")
except ImportError:
    # Fallback to old import if new package not available
    try:
        from langchain_community.vectorstores import Pinecone as PineconeVectorStore
        vectorstore = PineconeVectorStore(
            index=index,
            embedding=embeddings,
            text_key="text"
        )
        logging.info("✅ Using old langchain-community vectorstore")
    except Exception as e:
        logging.error(f"❌ Error initializing vectorstore: {e}")
        vectorstore = None
except Exception as e:
    logging.error(f"❌ Error initializing new vectorstore: {e}")
    vectorstore = None

# Initialize database
# DATABASE_NAME lives in pluto.users_db (re-exported above)

def init_db():
    """Initialize database with all required tables"""
    conn = sqlite3.connect(DATABASE_NAME)
    apply_sqlite_pragmas(conn)
    cursor = conn.cursor()
    
    # Create evaluations table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS evaluations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            resume_path TEXT,
            filename TEXT,
            job_title TEXT,
            job_description TEXT,
            match_percentage REAL,
            match_factors TEXT,
            profile_summary TEXT,
            missing_keywords TEXT,
            job_stability TEXT,
            career_progression TEXT,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            oorwin_job_id TEXT,
            candidate_fit_analysis TEXT,
            over_under_qualification TEXT,
            time_taken REAL,
            user_email TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_history table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question TEXT NOT NULL,
            retrieved_docs TEXT,
            final_answer TEXT,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create qa_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS qa_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER,
            rating INTEGER,
            feedback TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(question_id),
            FOREIGN KEY (question_id) REFERENCES qa_history (id)
        )
    ''')
    
    # Create feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            rating INTEGER,
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(evaluation_id),
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create handbook_feedback table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS handbook_feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            handbook_id INTEGER NOT NULL,
            rating INTEGER NOT NULL CHECK(rating >= 1 AND rating <= 5),
            comments TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(handbook_id),
            FOREIGN KEY (handbook_id) REFERENCES recruiter_handbooks (id)
        )
    ''')
    
    # Create interview_questions table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS interview_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            evaluation_id INTEGER,
            technical_questions TEXT,
            nontechnical_questions TEXT,
            behavioral_questions TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (evaluation_id) REFERENCES evaluations (id)
        )
    ''')
    
    # Create recruiter_handbooks table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recruiter_handbooks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            oorwin_job_id TEXT,
            job_title TEXT,
            job_description TEXT,
            additional_context TEXT,
            markdown_content TEXT,
            time_taken REAL,
            user_email TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create users table for authentication and role management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            name TEXT,
            role TEXT NOT NULL DEFAULT 'Recruiter',
            team TEXT,
            manager_email TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (manager_email) REFERENCES users(email)
        )
    ''')
    
    # Handle schema updates for existing tables
    try:
        # Check if evaluations table has new columns
        cursor.execute("PRAGMA table_info(evaluations)")
        columns = [col[1] for col in cursor.fetchall()]
        
        if 'candidate_fit_analysis' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN candidate_fit_analysis TEXT')
            print("Added candidate_fit_analysis column to evaluations")
        
        if 'over_under_qualification' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN over_under_qualification TEXT')
            print("Added over_under_qualification column to evaluations")
        
        if 'user_email' not in columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN user_email TEXT')
            print("Added user_email column to evaluations")
    except Exception as e:
        print(f"Note: Schema update check: {e}")
    
    # Add user_email to recruiter_handbooks if not exists
    try:
        cursor.execute("PRAGMA table_info(recruiter_handbooks)")
        handbook_columns = [col[1] for col in cursor.fetchall()]
        if 'user_email' not in handbook_columns:
            cursor.execute('ALTER TABLE recruiter_handbooks ADD COLUMN user_email TEXT')
            print("Added user_email column to recruiter_handbooks")
    except Exception as e:
        print(f"Note: Schema update check for handbooks: {e}")
    
    # Add time_taken column to evaluations table
    try:
        cursor.execute("PRAGMA table_info(evaluations)")
        eval_columns = [col[1] for col in cursor.fetchall()]
        if 'time_taken' not in eval_columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN time_taken REAL')
            print("Added time_taken column to evaluations")
        if 'evaluation_mode' not in eval_columns:
            cursor.execute("ALTER TABLE evaluations ADD COLUMN evaluation_mode TEXT DEFAULT 'single'")
            print("Added evaluation_mode column to evaluations")
        if 'batch_group_id' not in eval_columns:
            cursor.execute('ALTER TABLE evaluations ADD COLUMN batch_group_id TEXT')
            print("Added batch_group_id column to evaluations")
    except Exception as e:
        print(f"Note: Schema update check for evaluations time_taken: {e}")
    
    # Add time_taken column to recruiter_handbooks table
    try:
        cursor.execute("PRAGMA table_info(recruiter_handbooks)")
        handbook_columns = [col[1] for col in cursor.fetchall()]
        if 'time_taken' not in handbook_columns:
            cursor.execute('ALTER TABLE recruiter_handbooks ADD COLUMN time_taken REAL')
            print("Added time_taken column to recruiter_handbooks")
    except Exception as e:
        print(f"Note: Schema update check for handbooks time_taken: {e}")
    
    # Add intake_json column to recruiter_handbooks
    try:
        cursor.execute("PRAGMA table_info(recruiter_handbooks)")
        handbook_columns = [col[1] for col in cursor.fetchall()]
        if 'intake_json' not in handbook_columns:
            cursor.execute('ALTER TABLE recruiter_handbooks ADD COLUMN intake_json TEXT')
            print("Added intake_json column to recruiter_handbooks")
    except Exception as e:
        print(f"Note: Schema update check for handbooks intake_json: {e}")
    
    try:
        from pluto.voxpro.db import init_voxpro_tables

        init_voxpro_tables(cursor)
    except Exception as e:
        print(f"Note: VoxPro table init: {e}")

    ensure_indexes(cursor)

    # Initialize default admin user (override via DEFAULT_ADMIN_EMAIL in .env)
    try:
        if DEFAULT_ADMIN_EMAIL:
            admin_name = DEFAULT_ADMIN_EMAIL.split('@')[0].replace('.', ' ').title()
            cursor.execute('''
                INSERT OR IGNORE INTO users (email, name, role, team)
                VALUES (?, ?, ?, ?)
            ''', (DEFAULT_ADMIN_EMAIL, admin_name, 'Admin', 'Core'))
            conn.commit()
            logging.info("Initialized default admin user: %s", DEFAULT_ADMIN_EMAIL)
    except Exception as e:
        logging.warning("User initialization: %s", e)
    
    conn.commit()
    conn.close()

# Initialize database at startup
init_db()
start_upload_cleanup_scheduler(app)

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_gemini_response(input_prompt):
    """Get response from selected model (Gemini or OpenAI) and clean it up."""
    try:
        response = generate_content_unified(input_prompt)
        response_text = response.text.strip()
        
        # Remove markdown code block markers if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        # Clean up any extra whitespace and newlines
        response_text = response_text.strip()
        
        # Try to parse as JSON to validate
        try:
            parsed_json = json.loads(response_text)
            return json.dumps(parsed_json)  # Return properly formatted JSON string
        except json.JSONDecodeError:
            # If not valid JSON, try to extract JSON using regex
            match = re.search(r"\{.*\}", response_text, re.DOTALL)
            if match:
                try:
                    parsed_json = json.loads(match.group(0))
                    return json.dumps(parsed_json)  # Return properly formatted JSON string
                except json.JSONDecodeError:
                    raise ValueError("Invalid JSON structure in response")
            else:
                raise ValueError("No valid JSON found in response")
                
    except Exception as e:
        logging.error(f"Error in get_gemini_response: {str(e)}")
        return json.dumps({})  # Return valid empty JSON object as fallback

def sanitize_resume_text(text):
    """
    Remove email addresses and phone numbers from resume text for privacy protection.
    
    Args:
        text: Raw resume text
        
    Returns:
        Sanitized text with email and phone numbers removed
    """
    if not text:
        return text
    
    # Remove email addresses (various formats)
    # Pattern matches: user@domain.com, user.name@domain.co.uk, user+tag@domain.com, etc.
    # Note: Case-insensitive matching for better coverage
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    text = re.sub(email_pattern, '[EMAIL_REMOVED]', text, flags=re.IGNORECASE)
    
    # Remove phone numbers (various formats)
    # IMPORTANT: Patterns are designed to NOT match years (1900-2099) or date ranges
    # Strategy: Use negative lookahead/lookbehind to exclude common date contexts
    
    # US/Standard format: (123) 456-7890 or 123-456-7890 or 123.456.7890 or 123 456 7890
    # Exclude if it looks like a year (starts with 19xx or 20xx)
    phone_pattern2 = r'(?<!\d)(?!(?:19|20)\d{2})\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?!\d)'
    
    # Indian mobile format: +91-98765-43210 or 91-98765-43210 or 98765-43210 or 9876543210
    # Indian mobile numbers start with 6, 7, 8, or 9 (safe, won't match years)
    phone_pattern3 = r'(\+?91[-.\s]?)?[6-9]\d{4}[-.\s]?\d{5}'
    
    # Standalone 10-digit numbers starting with 3-9 (excluding 1xxx and 20xx which could be years)
    # Indian mobile: 6-9, US: 3-5,7-9 (excluding 1xxx and 20xx to avoid years)
    phone_pattern4 = r'\b([3-9]\d{9})\b(?!\s*(?:to|–|-)\s*(?:19|20)\d{2})'  # Exclude if followed by date range
    
    # International format with country code: +[1-9] followed by 7+ digits
    # More restrictive: require country code and sufficient digits to avoid matching dates
    phone_pattern1 = r'\+[1-9]\d{0,2}[-.\s]?\d{4,}[-.\s]?\d{3,}(?!\d)'
    
    # Apply all patterns in order
    text = re.sub(phone_pattern1, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern2, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern3, '[PHONE_REMOVED]', text)
    text = re.sub(phone_pattern4, '[PHONE_REMOVED]', text)
    
    # Clean up multiple consecutive replacements
    text = re.sub(r'\[EMAIL_REMOVED\](?:\s*\[EMAIL_REMOVED\])+', '[EMAIL_REMOVED]', text)
    text = re.sub(r'\[PHONE_REMOVED\](?:\s*\[PHONE_REMOVED\])+', '[PHONE_REMOVED]', text)
    
    return text

def extract_contact_info_from_text(text):
    """Extract basic contact info from raw resume text before sanitization."""
    if not text:
        return {"name": "", "email": "", "phone": "", "linkedin": ""}

    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    linkedin_pattern = r'((?:https?://)?(?:www\.)?linkedin\.com/[^\s)\]>]+)'
    phone_pattern = r'(\+?\d[\d\-\s().]{8,}\d)'

    emails = re.findall(email_pattern, text, flags=re.IGNORECASE)
    linkedins = re.findall(linkedin_pattern, text, flags=re.IGNORECASE)
    phones_raw = re.findall(phone_pattern, text)

    # Clean and keep plausible phone numbers
    phones = []
    for p in phones_raw:
        digits = re.sub(r'\D', '', p)
        if 10 <= len(digits) <= 14:
            phones.append(p.strip())

    # Deduplicate while preserving order
    def uniq(items):
        out, seen = [], set()
        for x in items:
            key = x.lower().strip()
            if key and key not in seen:
                seen.add(key)
                out.append(x.strip())
        return out

    emails = uniq(emails)
    linkedins = uniq(linkedins)
    phones = uniq(phones)

    linkedin = linkedins[0] if linkedins else ""
    if linkedin and not linkedin.lower().startswith("http"):
        linkedin = "https://" + linkedin

    # Candidate name heuristic:
    # pick the first strong alphabetic line near the top that is not a section/header keyword.
    candidate_name = ""
    blocked_tokens = {
        "resume", "curriculum vitae", "cv", "profile", "summary", "objective",
        "experience", "work experience", "education", "skills", "projects",
        "certifications", "contact", "about", "phone", "email", "linkedin"
    }
    lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
    def normalize_name(words):
        def normalize_token(token):
            if len(token) == 1:
                return token.upper()
            parts = re.split(r"([\-'])", token)
            out = []
            for part in parts:
                if part in ("-", "'"):
                    out.append(part)
                elif part:
                    out.append(part[0].upper() + part[1:].lower())
            return "".join(out)
        return " ".join(normalize_token(w) for w in words)

    for line in lines[:18]:
        cleaned = re.sub(r'[\|\(\)\[\]{}:;,/\\]+', ' ', line).strip()
        if len(cleaned) < 3 or len(cleaned) > 60:
            continue
        lower_cleaned = cleaned.lower()
        if any(tok in lower_cleaned for tok in blocked_tokens):
            continue
        # Allow letters, spaces, apostrophes, dots, hyphens; disallow heavy numeric/symbol lines.
        if not re.match(r"^[A-Za-z][A-Za-z .'\-]{1,58}$", cleaned):
            continue
        words = [w for w in cleaned.split() if w]
        if not (2 <= len(words) <= 4):
            continue
        if all(len(w) == 1 for w in words):
            continue
        candidate_name = normalize_name(words)
        break

    return {
        "name": candidate_name,
        "email": emails[0] if emails else "",
        "phone": phones[0] if phones else "",
        "linkedin": linkedin
    }

def extract_contact_info_from_file(file_path):
    """Best-effort contact extraction from raw file text (not sanitized)."""
    if not file_path or not os.path.exists(file_path):
        return {"name": "", "email": "", "phone": "", "linkedin": ""}

    ext = file_path.rsplit(".", 1)[1].lower() if "." in file_path else ""
    raw_text = ""
    try:
        if ext == "pdf":
            with pdfplumber.open(file_path) as pdf:
                parts = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t:
                        parts.append(t)
                raw_text = "\n".join(parts)
        elif ext == "docx":
            doc = Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs if p.text])
        else:
            # .doc and other formats: skip for now
            raw_text = ""
    except Exception as e:
        logging.warning(f"Contact extraction failed for {file_path}: {e}")
        raw_text = ""

    return extract_contact_info_from_text(raw_text)

def extract_text_from_file(file_path, *, return_error: bool = False):
    """
    Extract resume text from supported file types.

    Backwards-compatible behavior:
    - default returns: text (str) OR None
    - if return_error=True: returns (text_or_none, error_message_or_none)
    """

    def _fail(msg: str):
        logging.error(f"File extraction error: {msg}")
        return (None, msg) if return_error else None

    def _ok(text: str):
        return (text, None) if return_error else text

    def _find_soffice_exe() -> str | None:
        # Prefer PATH, then typical Windows locations.
        exe = shutil.which("soffice") or shutil.which("soffice.exe")
        if exe and os.path.exists(exe):
            return exe
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None

    def _convert_doc_to_docx(input_doc_path: str) -> tuple[str | None, str | None, str | None]:
        """
        Convert .doc -> .docx using LibreOffice headless.
        Returns (docx_path, tmp_dir, error_message).
        """
        soffice = _find_soffice_exe()
        if not soffice:
            return None, None, (
                "This looks like a .doc file. To support .doc, please install LibreOffice "
                "and make sure 'soffice' is available (LibreOffice\\program\\soffice.exe)."
            )

        tmp_dir = tempfile.mkdtemp(prefix="pluto_doc_convert_")
        try:
            cmd = [
                soffice,
                "--headless",
                "--nologo",
                "--nodefault",
                "--norestore",
                "--convert-to",
                "docx",
                "--outdir",
                tmp_dir,
                input_doc_path,
            ]
            completed = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
            )
            if completed.returncode != 0:
                stderr = (completed.stderr or "").strip()
                stdout = (completed.stdout or "").strip()
                return None, tmp_dir, f"Failed to convert .doc to .docx (LibreOffice). {stderr or stdout or 'Unknown conversion error'}"

            # LibreOffice outputs <stem>.docx (usually), but sometimes adjusts name.
            stem = os.path.splitext(os.path.basename(input_doc_path))[0]
            expected = os.path.join(tmp_dir, f"{stem}.docx")
            if os.path.exists(expected):
                return expected, tmp_dir, None

            # Fallback: pick the first docx in output dir
            for fn in os.listdir(tmp_dir):
                if fn.lower().endswith(".docx"):
                    return os.path.join(tmp_dir, fn), tmp_dir, None

            return None, tmp_dir, "LibreOffice conversion completed but no .docx output was produced."
        except subprocess.TimeoutExpired:
            return None, tmp_dir, "Timed out converting .doc to .docx. Please try again or convert the file manually."
        except Exception as e:
            return None, tmp_dir, f"Unexpected .doc conversion error: {str(e)}"

    def _extract_text_from_docx(docx_path: str) -> tuple[str | None, str | None]:
        try:
            doc = Document(docx_path)
            chunks = []
            for para in doc.paragraphs:
                if para.text:
                    chunks.append(para.text)
            text = "\n".join(chunks)
            text = sanitize_resume_text(text)
            if not text.strip():
                return None, "No readable text found in the .docx file."
            return text, None
        except Exception as e:
            return None, f"Failed to read .docx: {str(e)}"

    def _extract_text_from_pdf(pdf_path: str) -> tuple[str | None, str | None]:
        # PyMuPDF is typically 2–4× faster than pdfplumber for plain text extraction.
        try:
            import fitz  # PyMuPDF

            parts = []
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    t = page.get_text("text") or ""
                    if t.strip():
                        parts.append(t)
            text = sanitize_resume_text("\n".join(parts))
            if len(text.strip()) >= 50:
                return text, None
        except Exception as pymupdf_err:
            logging.debug("PyMuPDF extract fallback to pdfplumber: %s", pymupdf_err)

        try:
            with pdfplumber.open(pdf_path) as pdf:
                parts = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t:
                        parts.append(t)
            text = "\n".join(parts)
            text = sanitize_resume_text(text)
            return text, None
        except ModuleNotFoundError as e:
            if "PyCryptodome" in str(e) or "Crypto" in str(e):
                return None, "PyCryptodome is required for some encrypted PDFs. Install it with: pip install pycryptodome"
            return None, f"Missing dependency while reading PDF: {str(e)}"
        except Exception as e:
            return None, f"Failed to read PDF: {str(e)}"

    def _bool_env(name: str, default: str = "1") -> bool:
        v = os.getenv(name, default)
        return str(v).strip().lower() in ("1", "true", "yes", "y", "on")

    def _find_tesseract_cmd() -> str | None:
        """
        Find the Tesseract executable.
        - Respects TESSERACT_CMD env var if set
        - Falls back to PATH
        - Falls back to common Windows install locations
        """
        override = os.getenv("TESSERACT_CMD", "").strip().strip('"')
        if override and os.path.exists(override):
            return override

        exe = shutil.which("tesseract") or shutil.which("tesseract.exe")
        if exe and os.path.exists(exe):
            return exe

        candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        # Common winget install location (per-user)
        local_appdata = os.getenv("LOCALAPPDATA", "").strip().strip('"')
        if local_appdata:
            candidates.append(os.path.join(local_appdata, "Programs", "Tesseract-OCR", "tesseract.exe"))
        for p in candidates:
            if os.path.exists(p):
                return p
        return None

    def _ocr_pdf(pdf_path: str, *, max_pages: int = 8, lang: str = "eng") -> tuple[str | None, str | None]:
        """
        OCR a scanned/image PDF (best-effort).

        Requires Python deps:
          - pymupdf, pytesseract, pillow
        And system install:
          - Tesseract OCR (tesseract.exe)
        """
        # Lazy imports so non-OCR flows don't require these deps at runtime.
        try:
            import fitz  # PyMuPDF
        except Exception:
            return None, "OCR requires PyMuPDF. Please install it: pip install pymupdf"
        try:
            import pytesseract
        except Exception:
            return None, "OCR requires pytesseract. Please install it: pip install pytesseract"
        try:
            from PIL import Image
        except Exception:
            return None, "OCR requires Pillow. Please install it: pip install pillow"

        tess_cmd = _find_tesseract_cmd()
        if not tess_cmd:
            return None, (
                "Scanned (image-only) PDF detected. To support scanned PDFs, please install Tesseract OCR "
                "and ensure tesseract.exe is available (PATH or set TESSERACT_CMD)."
            )
        try:
            pytesseract.pytesseract.tesseract_cmd = tess_cmd
        except Exception:
            pass

        try:
            doc = fitz.open(pdf_path)
        except Exception as e:
            return None, f"Failed to open PDF for OCR: {str(e)}"

        try:
            page_count = doc.page_count
            pages_to_ocr = min(max(page_count, 0), max_pages)
            texts = []
            for i in range(pages_to_ocr):
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=200)
                mode = "RGBA" if pix.alpha else "RGB"
                img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                if mode == "RGBA":
                    img = img.convert("RGB")
                txt = pytesseract.image_to_string(img, lang=lang) or ""
                if txt.strip():
                    texts.append(txt)

            combined = "\n".join(texts).strip()
            if not combined:
                return None, "No readable text found via OCR. The scanned PDF may be too low quality."

            combined = sanitize_resume_text(combined)
            if pages_to_ocr < page_count:
                combined += f"\n\n[Note: OCR processed first {pages_to_ocr} of {page_count} pages]"
            return combined, None
        except Exception as e:
            return None, f"OCR failed: {str(e)}"
        finally:
            try:
                doc.close()
            except Exception:
                pass

    # Validate input path
    if not file_path or file_path in ("NULL", "None") or str(file_path).strip() == "":
        return _fail("Invalid file path")
    if not os.path.exists(file_path):
        return _fail(f"File not found: {file_path}")

    ext = file_path.rsplit(".", 1)[1].lower() if "." in file_path else ""
    try:
        if ext == "pdf":
            extracted, err = _extract_text_from_pdf(file_path)
            # Scanned/image PDFs won't have extractable text; fall back to OCR if enabled
            if not extracted or len(re.sub(r"\s+", "", extracted)) < 50:
                if _bool_env("ENABLE_PDF_OCR", "1"):
                    ocr_text, ocr_err = _ocr_pdf(
                        file_path,
                        max_pages=int(os.getenv("OCR_MAX_PAGES", "8")),
                        lang=os.getenv("OCR_LANG", "eng"),
                    )
                    if ocr_text:
                        return _ok(ocr_text)
                    return _fail(ocr_err or "Scanned (image-only) PDF detected, but OCR failed.")

                return _fail(
                    "Scanned (image-only) PDFs aren’t supported right now. "
                    "Please upload a text-based PDF (selectable text) or a DOC/DOCX file."
                )
            return _ok(extracted)

        if ext == "docx":
            text, err = _extract_text_from_docx(file_path)
            if text:
                return _ok(text)
            return _fail(err or "Failed to extract text from .docx")

        if ext == "doc":
            docx_path, tmp_dir, conv_err = _convert_doc_to_docx(file_path)
            try:
                if not docx_path:
                    return _fail(conv_err or "Failed to convert .doc to .docx")
                text, err = _extract_text_from_docx(docx_path)
                if text:
                    return _ok(text)
                return _fail(err or "Converted .doc but could not extract text")
            finally:
                # Cleanup conversion artifacts
                try:
                    if tmp_dir and os.path.exists(tmp_dir):
                        shutil.rmtree(tmp_dir, ignore_errors=True)
                except Exception:
                    pass

        return _fail("Unsupported file format. Please upload a PDF, DOCX, or DOC file.")

    except Exception as e:
        return _fail(str(e))

def hybrid_search(query, k=5):
    """Perform hybrid search using BM25 and vector similarity."""
    try:
        # Get vector search results
        vector_results = vectorstore.similarity_search(query, k=k)
        
        # Get BM25 results
        bm25_results = []
        if os.path.exists(POLICIES_FOLDER):
            for filename in os.listdir(POLICIES_FOLDER):
                if filename.endswith(('.txt', '.md')):
                    with open(os.path.join(POLICIES_FOLDER, filename), 'r', encoding='utf-8') as f:
                        text = f.read()
                        sentences = sent_tokenize(text)  # Use NLTK sentence tokenizer
                        bm25_results.extend(sentences)
        
        # Combine and deduplicate results
        combined_results = []
        seen_texts = set()
        
        # Add vector search results
        for doc in vector_results:
            if doc.page_content not in seen_texts:
                combined_results.append(doc.page_content)
                seen_texts.add(doc.page_content)
        
        # Add BM25 results
        for sentence in bm25_results:
            if sentence not in seen_texts:
                combined_results.append(sentence)
                seen_texts.add(sentence)
        
        # Join results with newlines
        return "\n".join(combined_results)
    
    except Exception as e:
        logging.error(f"Error in hybrid_search: {e}")
        return ""

def save_evaluation(
    eval_id,
    filename,
    job_title,
    rank_score,
    missing_keywords,
    profile_summary,
    match_factors,
    job_stability,
    additional_info=None,
    oorwin_job_id=None,
    candidate_fit_analysis=None,
    over_under_qualification=None,
    user_email=None,
    time_taken=None,
    evaluation_mode='single',
    batch_group_id=None,
    job_description=None,
):
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Convert data to JSON strings if they're not already
        try:
            missing_keywords_json = json.dumps(missing_keywords) if not isinstance(missing_keywords, str) else missing_keywords
        except Exception as e:
            logging.error(f"Error converting missing_keywords to JSON: {e}")
            missing_keywords_json = '[]'
            
        try:
            match_factors_json = json.dumps(match_factors) if not isinstance(match_factors, str) else match_factors
        except Exception as e:
            logging.error(f"Error converting match_factors to JSON: {e}")
            match_factors_json = '{}'
            
        try:
            job_stability_json = json.dumps(job_stability) if not isinstance(job_stability, str) else job_stability
        except Exception as e:
            logging.error(f"Error converting job_stability to JSON: {e}")
            job_stability_json = '{}'
        
        # Ensure all JSON strings are valid
        if not missing_keywords_json or missing_keywords_json == 'null':
            missing_keywords_json = '[]'
        if not match_factors_json or match_factors_json == 'null':
            match_factors_json = '{}'
        if not job_stability_json or job_stability_json == 'null':
            job_stability_json = '{}'
        
        # Convert rank_score to integer if it's a string
        rank_score_int = int(rank_score) if isinstance(rank_score, str) else rank_score
        
        # Ensure all values are strings except rank_score_int and eval_id
        filename_str = str(filename)
        job_title_str = str(job_title)
        profile_summary_str = str(profile_summary)
        
        # Handle oorwin_job_id (can be None or empty string)
        oorwin_job_id_str = str(oorwin_job_id).strip() if oorwin_job_id else None
        if oorwin_job_id_str == '' or oorwin_job_id_str == 'None':
            oorwin_job_id_str = None
        
        # Convert additional_info to JSON string if it's a dict or list
        if isinstance(additional_info, (dict, list)):
            additional_info_str = json.dumps(additional_info)
        else:
            additional_info_str = str(additional_info) if additional_info is not None else ""
        
        # Extract career progression from additional_info
        career_progression = additional_info.get('career_progression', {}) if isinstance(additional_info, dict) else {}
        career_progression_json = json.dumps(career_progression)
        
        # Convert new fields to JSON
        candidate_fit_analysis_json = json.dumps(candidate_fit_analysis) if candidate_fit_analysis else '{}'
        over_under_qualification_str = str(over_under_qualification) if over_under_qualification else ''
        
        # Debug: Log the actual values being inserted
        logging.info(f"Values to insert - eval_id: {eval_id}, filename: {filename_str}, job_title: {job_title_str}")
        logging.info(f"JSON values - missing_keywords_json type: {type(missing_keywords_json)}, value: {missing_keywords_json[:100] if len(missing_keywords_json) > 100 else missing_keywords_json}")
        logging.info(f"JSON values - match_factors_json type: {type(match_factors_json)}")
        logging.info(f"All param types: rank_score_int={type(rank_score_int)}, oorwin_job_id_str={type(oorwin_job_id_str)}, datetime={type(datetime.now())}")
        
        # Convert datetime to string
        timestamp_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Prepare user_email
        user_email_str = user_email if user_email else None

        evaluation_mode_str = (evaluation_mode or 'single').strip() or 'single'
        batch_group_id_str = str(batch_group_id).strip() if batch_group_id else None
        if batch_group_id_str == '' or batch_group_id_str == 'None':
            batch_group_id_str = None
        job_description_str = str(job_description).strip() if job_description else ''
        
        # Prepare all parameters as a tuple (excluding eval_id which is not in the INSERT)
        # Order must match the INSERT statement columns exactly:
        # resume_path, filename, job_title, job_description, match_percentage, 
        # match_factors, profile_summary, missing_keywords, 
        # job_stability, career_progression, technical_questions,
        # nontechnical_questions, behavioral_questions, oorwin_job_id, 
        # candidate_fit_analysis, over_under_qualification, user_email, timestamp
        params = (
            filename_str,                    # resume_path
            filename_str,                    # filename
            job_title_str,                   # job_title
            job_description_str,             # job_description
            rank_score_int,                 # match_percentage
            match_factors_json,             # match_factors
            profile_summary_str,             # profile_summary
            missing_keywords_json,           # missing_keywords
            job_stability_json,              # job_stability
            career_progression_json,         # career_progression
            None,                            # technical_questions
            None,                            # nontechnical_questions
            None,                            # behavioral_questions
            oorwin_job_id_str,               # oorwin_job_id
            candidate_fit_analysis_json,     # candidate_fit_analysis
            over_under_qualification_str,    # over_under_qualification
            time_taken,                      # time_taken (in seconds, can be None)
            user_email_str,                  # user_email
            evaluation_mode_str,             # evaluation_mode: single | batch
            batch_group_id_str,              # batch_group_id (shared UUID for one comparison run)
            timestamp_str                    # timestamp
        )
        
        # Log all parameter types
        logging.info(f"Parameter types: {[type(p).__name__ for p in params]}")
        logging.info(f"Inserting evaluation - user_email: {user_email_str}, timestamp: {timestamp_str}")
        
        cursor.execute(
            """
            INSERT INTO evaluations (
                resume_path, filename, job_title, job_description, match_percentage, 
                match_factors, profile_summary, missing_keywords, 
                job_stability, career_progression, technical_questions,
                nontechnical_questions, behavioral_questions, oorwin_job_id, 
                candidate_fit_analysis, over_under_qualification, time_taken, user_email,
                evaluation_mode, batch_group_id, timestamp
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            params
            )
        
        # Get the auto-generated ID
        db_id = cursor.lastrowid
        
        conn.commit()
        
        # Log successful save
        logging.info(f"✅ EVALUATION SAVED TO DATABASE!")
        logging.info(f"   database_id: {db_id}")
        logging.info(f"   filename: {filename_str}")
        logging.info(f"   job_title: {job_title_str}")
        logging.info(f"   match_percentage: {rank_score_int}")
        logging.info(f"   oorwin_job_id: {oorwin_job_id_str}")
        
        conn.close()
        return db_id  # Return the database ID instead of True
    except Exception as e:
        logging.error(f"Database error in save_evaluation: {str(e)}", exc_info=True)
        logging.error(f"Data being saved - eval_id: {eval_id}, filename: {filename}, job_title: {job_title}")
        logging.error(f"Data types - rank_score: {type(rank_score)}, missing_keywords: {type(missing_keywords)}")
        return False

def save_feedback(evaluation_id, rating, comments):
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO feedback (evaluation_id, rating, comments, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, rating, comments, datetime.now())
        )
        logging.debug(f"Feedback inserted: evaluation_id={evaluation_id}, rating={rating}, comments={comments}")
        conn.commit()
        conn.close()
        return True
    except sqlite3.Error as e:
        logging.error(f"Database error in save_feedback: {str(e)}")
        return False
    except Exception as e:
        logging.error(f"Unexpected error in save_feedback: {str(e)}")
        return False

def save_interview_questions(evaluation_id, technical_questions, nontechnical_questions, behavioral_questions):
    """Save interview questions to database"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        technical_json = json.dumps(technical_questions) if not isinstance(technical_questions, str) else technical_questions
        nontechnical_json = json.dumps(nontechnical_questions) if not isinstance(nontechnical_questions, str) else nontechnical_questions
        behavioral_json = json.dumps(behavioral_questions) if not isinstance(behavioral_questions, str) else behavioral_questions
        
        cursor.execute(
            "INSERT INTO interview_questions (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions, timestamp) VALUES (?, ?, ?, ?, ?)",
            (evaluation_id, technical_json, nontechnical_json, behavioral_json, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
        conn.commit()
        conn.close()
        logging.debug(f"Interview questions saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_interview_questions: {str(e)}")
        return False

def save_recruiter_handbook(evaluation_id, markdown_content, json_summary):
    """Save recruiter handbook to database"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Ensure json_summary is a string
        json_summary_str = json_summary if isinstance(json_summary, str) else json.dumps(json_summary)
        
        cursor.execute(
            "INSERT INTO recruiter_handbooks (evaluation_id, markdown_content, json_summary, timestamp) VALUES (?, ?, ?, ?)",
            (evaluation_id, markdown_content, json_summary_str, datetime.now())
        )
        conn.commit()
        conn.close()
        logging.debug(f"Recruiter handbook saved successfully for: {evaluation_id}")
        return True
    except Exception as e:
        logging.error(f"Database error in save_recruiter_handbook: {str(e)}")
        return False

# Add these constants near the top with other constants
BOT_INFO = {
    "name": "PeopleBot",
    "creator": "PeopleLogic",
    "responsibility": "Help recruiters in HR policies, benefits & with any other questions!",
    "capabilities": "Help recruiters in HR policies, benefits & with any other questions"
}

GREETINGS = ["hello", "hi", "hey", "greetings", "good morning", "good afternoon", "good evening"]
IDENTITY_QUESTIONS = [
    "who are you",
    "who are u",
    "who r yu",
    "who r u",
    "who are you",
    "what is your name",
    "what are you",
    "who built you",
    "what are u",
    "who built u",
    "who created you",
    "what can you do",
    "what do you do",
    "what is your name",
    "tell me about yourself",
    "tell me about you",
     "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do", "can you tell me your name", "can you introduce yourself",  
    "what should I call you", "how do I address you",  
    "could you tell me what you are", "what do people call you",  
    "what is your full name", "how were you created",  
    "give me some details about you", "tell me your background",  "who u", "who dis", "whats ur deal", "whats ur function", "who tf r u",  
    "who's this", "u bot?", "who you be", "ur name?", "who made u",  
    "who's your maker", "who made this bot", "whu r u", "whi r u", "who r yu", "wht is ur name", "whts ur name",  
    "whats ur name", "wat is your name", "wat r u", "whu built u",  
    "who made you", "who designed you", "who programmed you", "what are you capable of", "what skills do you have",  
    "describe yourself", "what is your role", "explain yourself",  
    "what services do you provide", "how do you work",  
    "can you tell me what you do",  

    # Mixed uppercase/lowercase variations  
    "Who Are You", "WHO ARE YOU", "Who r U", "wHo bUiLt YoU",  
    "WHAT IS YOUR NAME", "wHaT dO yOu dO",  

    # Slang & misspellin
    "whoz u", "what u do", "wots ur name", "whut do u do",  
    "wats ur function", "who dun made u",  

    # Phonetic spellings & accents  
    "hoo r u", "wat iz ur name", "hoo maid u", "wat cn u do",  
    "whut ur name", "whachu do",  

    # Extra punctuation variations  
    "who are you?", "who are you!", "who are you??",  
    "what is your name?", "who built you??", "who made you?!",
]

def handle_special_queries(question):
    """Handle greetings and identity-related questions."""
    question_lower = question.lower().strip("?!. ")
    
    # Handle greetings
    if question_lower in GREETINGS:
        return f"Hello! I'm {BOT_INFO['name']}, your HR assistant. How can I help you today?"
    
    # Handle identity questions
    if any(q in question_lower for q in IDENTITY_QUESTIONS):
        if "who" in question_lower or "what is your name" in question_lower:
            return f"I'm {BOT_INFO['name']}, an AI assistant built by {BOT_INFO['creator']}. {BOT_INFO['responsibility']}"
        elif "created" in question_lower or "built" in question_lower:
            return f"I was created by {BOT_INFO['creator']} to {BOT_INFO['responsibility']}"
        elif "can you do" in question_lower or "do you do" in question_lower:
            return f"I can {BOT_INFO['capabilities']}"
        else:
            return f"I'm {BOT_INFO['name']}, an AI assistant created by {BOT_INFO['creator']}. {BOT_INFO['capabilities']}"

    # Handle holiday list queries (static 2025 list provided by HR)
    if "holiday" in question_lower or "holidays" in question_lower:
        year = "2025"
        header = f"## Company Holidays {year}\n\nBelow are the declared holidays for {year}.\n\n"

        india_table = (
            "### India Offices (Bangalore/APAC & EU, Hyderabad, Mumbai, Delhi)\n"
            "| Date | Day | Bangalore/APAC & EU | Hyderabad | Mumbai | Delhi |\n"
            "|------|-----|----------------------|-----------|--------|-------|\n"
            "| 1-Jan-2025 | Wednesday | New Year | New Year | New Year | New Year |\n"
            "| 14-Jan-2025 | Tuesday | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti | Pongal/ Makar Sankranti |\n"
            "| 14-Mar-2025 | Friday | - | Holi | Holi | Holi |\n"
            "| 31-Mar-2025 | Monday | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | Ramzan (Id Ul Fitr) | - |\n"
            "| 18-Apr-2025 | Friday | Good Friday | - | - | Good Friday |\n"
            "| 1-May-2025 | Thursday | May Day | May Day | May Day | May Day |\n"
            "| 15-Aug-2025 | Friday | Independence Day | Independence Day | Independence Day | Independence Day |\n"
            "| 27-Aug-2025 | Wednesday | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi | Ganesh Chaturthi |\n"
            "| 2-Oct-2025 | Thursday | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara | Gandhi Jayanthi/Dasara |\n"
            "| 20-Oct-2025 | Monday | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi | Diwali-Naraka Chaturdashi |\n"
            "| 25-Dec-2025 | Thursday | Christmas | Christmas | Christmas | Christmas |\n"
        )

        us_table = (
            "\n### Global Services - US\n"
            "| Date | Day | Holiday |\n"
            "|------|-----|---------|\n"
            "| 1-Jan-2025 | Wednesday | New Year |\n"
            "| 18-Apr-2025 | Friday | Good Friday |\n"
            "| 26-May-2025 | Monday | Memorial Day |\n"
            "| 4-Jul-2025 | Friday | Independence Day |\n"
            "| 1-Sep-2025 | Monday | Labour Day |\n"
            "| 20-Oct-2025 | Monday | Diwali |\n"
            "| 27-Nov-2025 | Thursday | Thanksgiving |\n"
            "| 28-Nov-2025 | Friday | Day after Thanksgiving |\n"
            "| 24-Dec-2025 | Wednesday | Christmas Eve |\n"
            "| 25-Dec-2025 | Thursday | Christmas Day |\n"
        )

        footnote = "\n> Note: If a holiday falls on a weekend, local HR guidelines on compensatory off apply."
        return header + india_table + us_table + footnote
    
    return None

# HR Assistant routes (other pages / APIs: pluto.blueprints, pluto.routes).
@limiter.limit("30 per minute")
@login_required
def ask_question():
    try:
        data = request.get_json()
        question = data.get('question')
        online_mode = data.get('online_mode', False)

        if not question:
            return jsonify({'error': 'No question provided'}), 400

        def generate():
            complete_response = []  # Store complete response
            try:
                info_buddy_provider, info_buddy_groq_model, info_buddy_groq_reasoning_effort = get_info_buddy_provider_settings()
                info_buddy_model_override = info_buddy_groq_model if info_buddy_provider == "groq" else None
                selected_info_buddy_model = (
                    info_buddy_model_override if info_buddy_provider == "groq"
                    else OPENAI_MODEL if info_buddy_provider == "openai"
                    else GEMINI_MODEL
                )
                logging.info(
                    "Info Buddy request provider=%s model=%s",
                    info_buddy_provider,
                    selected_info_buddy_model
                )
                print(f"ℹ️ Info Buddy request provider={info_buddy_provider.upper()} model={selected_info_buddy_model}")

                # Check for special queries first
                special_response = handle_special_queries(question)
                if special_response:
                    complete_response.append(special_response)
                    yield special_response
                    return

                # Expand acronyms in the question
                expanded_question = expand_acronyms(question)

                if online_mode:
                    # ONLINE MODE: Answer any general question using LLM knowledge
                    # No RAG constraints - can answer anything
                    detailed_prompt = f"""You are an expert AI assistant. Provide a comprehensive and detailed answer to the following question. Your response should be thorough, well-structured, and accurate.

                    Question: {expanded_question}

                    Instructions:
                    1. Use your knowledge to provide a complete answer
                    2. If the question is about HR policies, benefits, or company-specific information, note that you may not have the latest company-specific details
                    3. Format your response with clear sections and bullet points where appropriate
                    4. Include relevant examples and context
                    5. If you're uncertain about specific facts, mention that
                    """
                    
                    response = generate_content_unified(
                        detailed_prompt,
                        stream=True,
                        provider_override=info_buddy_provider,
                        model_override=info_buddy_model_override,
                        groq_reasoning_effort_override=info_buddy_groq_reasoning_effort
                    )
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text
                else:
                    # RAG MODE: Strict retrieval from local documents only
                    # Use hybrid search (BM25 + Vector) for better coverage
                    
                    # Step 1: Hybrid retrieval - combine BM25 and Vector search
                    all_retrieved_docs = []
                    
                    # Vector search (semantic similarity) - increased k for better coverage
                    if vectorstore is not None:
                        vector_docs = vectorstore.similarity_search(expanded_question, k=15)
                        all_retrieved_docs.extend([(doc, 'vector') for doc in vector_docs])
                        logging.info(f"🔍 Vector search retrieved {len(vector_docs)} documents")
                    
                    # BM25 search (keyword matching) - better for exact terms and tables
                    if bm25_index and bm25_corpus:
                        try:
                            query_tokens = expanded_question.lower().split()
                            bm25_scores = bm25_index.get_scores(query_tokens)
                            
                            # Get top BM25 results - increased k for better coverage
                            top_indices = sorted(range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True)[:10]
                            bm25_results = []
                            for idx in top_indices:
                                if bm25_scores[idx] > 0:  # Only include relevant results
                                    text_content = " ".join(bm25_corpus[idx])
                                    # Create a Document-like object with metadata for consistency
                                    from langchain_core.documents import Document as LangchainDocument
                                    # Get metadata for this chunk if available
                                    metadata = bm25_metadata[idx] if idx < len(bm25_metadata) else {}
                                    bm25_doc = LangchainDocument(page_content=text_content, metadata=metadata)
                                    bm25_results.append(bm25_doc)
                                    all_retrieved_docs.append((bm25_doc, 'bm25'))
                            
                            logging.info(f"🔍 BM25 search retrieved {len(bm25_results)} documents")
                        except Exception as e:
                            logging.warning(f"⚠️ BM25 search failed: {e}")
                    
                    # Step 2: Deduplicate and prioritize
                    seen_content = set()
                    unique_docs = []
                    table_docs = []
                    text_docs = []
                    
                    for doc, source in all_retrieved_docs:
                        content_hash = hash(doc.page_content[:100])  # Hash first 100 chars for dedup
                        if content_hash not in seen_content:
                            seen_content.add(content_hash)
                            # Classify as table or text
                            if "[TABLE DATA]" in doc.page_content or ("|" in doc.page_content and doc.page_content.count("|") > 3):
                                table_docs.append(doc)
                            else:
                                text_docs.append(doc)
                            unique_docs.append(doc)
                    
                    # Step 3: Prioritize tables and limit context window
                    # Tables first (they're often most precise), then text chunks
                    prioritized_docs = table_docs + text_docs
                    context_docs = prioritized_docs[:12]  # Take top 12 for context
                    
                    # Step 4: Build context with proper source citations (filename + page)
                    if context_docs:
                        context_parts = []
                        for i, doc in enumerate(context_docs):
                            # Extract actual source filename and page from metadata
                            source_name = doc.metadata.get('source', 'Unknown Document') if hasattr(doc, 'metadata') and doc.metadata else 'Unknown Document'
                            page_num = doc.metadata.get('page', 'N/A') if hasattr(doc, 'metadata') and doc.metadata else 'N/A'
                            
                            # Format citation with actual filename
                            if source_name != 'Unknown Document':
                                citation = f"{source_name}"
                                if page_num != 'N/A':
                                    citation += f", page {page_num}"
                            else:
                                citation = f"Source {i+1}"
                            
                            # Add relevance indicator for tables with proper citation
                            if doc in table_docs:
                                context_parts.append(f"[RELEVANT TABLE DATA - {citation}]\n{doc.page_content}")
                            else:
                                context_parts.append(f"[RELEVANT CONTEXT - {citation}]\n{doc.page_content}")
                        context = "\n\n---\n\n".join(context_parts)
                        
                        # Log retrieval stats
                        logging.info(f"📚 Total unique documents retrieved: {len(unique_docs)} ({len(table_docs)} tables, {len(text_docs)} text)")
                        if table_docs:
                            logging.info(f"📊 Including {len(table_docs)} table chunks in context")
                    else:
                        context = ""
                        logging.warning("⚠️ No documents retrieved from knowledge base")
                    
                    # Step 5: Enhanced RAG prompt with strict enforcement
                    if context:
                        prompt = f"""You are an expert HR Assistant. Answer the question STRICTLY based on the provided context from company policy documents.

                    Question: {expanded_question}
                    
                    Context from Company Documents:
                    {context}

                    STRICT RULES FOR FORMATTING:
                    
                    1. **ONLY use information from the context provided above.** Do not use external knowledge.
                    
                    2. **TABLE FORMATTING (CRITICAL)**:
                       - If the context contains tables (look for markers like [TABLE DATA] ... [END TABLE]), you MUST reproduce the markdown table **verbatim** from the context
                       - Use this EXACT format:
                         ```
                         | Column 1 | Column 2 | Column 3 |
                         |---------|---------|----------|
                         | Data 1  | Data 2  | Data 3   |
                         | Data 4  | Data 5  | Data 6   |
                         ```
                       - Always include the header separator row (|---------|---------|)
                       - Keep table columns aligned and readable
                       - If a table is complex, break it into smaller, clearer tables
                       - Add a brief title above each table (e.g., "### Performance Rating Table")
                    
                    3. **RESPONSE STRUCTURE**:
                       - Start with a brief 1-2 sentence summary
                       - Use clear headings (## for main sections, ### for subsections)
                       - Use bullet points (• or -) for lists, NOT long paragraphs
                       - Add blank lines between sections for readability
                       - Keep paragraphs SHORT (3-4 sentences max)
                       - Use bold (**text**) for key terms and important points
                    
                    4. **SOURCE CITATION**:
                       - **ALWAYS cite sources using actual document names** (e.g., "According to [Leave Policy.pdf, page 5]")
                       - Place citations at the END of sentences or paragraphs, not mid-sentence
                       - Format: `[Document Name.pdf, page X]`
                       - Use actual filename, not generic "Source 1" or "Source 2"
                    
                    5. **READABILITY ENHANCEMENTS**:
                       - Break dense information into digestible chunks
                       - Use numbered lists (1., 2., 3.) for step-by-step processes
                       - Use bullet lists (•) for features, benefits, or items
                       - Add horizontal rules (---) to separate major sections
                       - Use emojis sparingly for visual breaks (✅, 📋, 📊, etc.)
                    
                    6. **EXAMPLE OF GOOD FORMATTING**:
                       ```
                       ## Performance Appraisal Policy
                       
                       The performance appraisal policy outlines how employees are evaluated annually.
                       
                       ### Key Features
                       • Appraisals are conducted yearly (April to March)
                       • Based on previously agreed KRAs
                       • Results can lead to salary hikes or promotions
                       
                       ### Performance Rating Table
                       | Score | Rating | Description |
                       |-------|--------|-------------|
                       | 5     | Exceptional | Targets met at 200% or above |
                       | 4     | Outstanding | Targets exceeded significantly |
                       | 3     | Good | Consistently met expectations |
                       
                       [APPRAISAL & PROMOTION POLICY.pdf, page 1]
                       ```
                    
                    7. **IMPORTANT**: If the question cannot be answered from the provided context, you MUST respond with:
                       - "I'm sorry, but the information about '[topic]' is not available in our company policy documents."
                       - "💡 **Suggestion**: Please enable the **'Go Online'** toggle and try asking your question again."
                    
                    8. DO NOT make up information or use knowledge outside the provided context.
                    
                    Answer:"""
                    else:
                        # No context found - give helpful message
                        prompt = f"""The question "{expanded_question}" could not be answered from the available company policy documents.

                    Please note:
                    - The information may not be in the current knowledge base
                    - The document may need to be updated or added
                    - You can try rephrasing the question or enabling "Go Online" mode for general information
                    
                    Would you like to:
                    1. Try rephrasing your question
                    2. Enable "Go Online" mode for general information
                    3. Contact HR for company-specific policies not yet in the system"""
                        
                        # Still generate response but with this constraint
                        response = generate_content_unified(
                            prompt,
                            stream=True,
                            provider_override=info_buddy_provider,
                            model_override=info_buddy_model_override,
                            groq_reasoning_effort_override=info_buddy_groq_reasoning_effort
                        )
                        for chunk in response:
                            if chunk.text:
                                complete_response.append(chunk.text)
                                yield chunk.text
                        return
                    
                    response = generate_content_unified(
                        prompt,
                        stream=True,
                        provider_override=info_buddy_provider,
                        model_override=info_buddy_model_override,
                        groq_reasoning_effort_override=info_buddy_groq_reasoning_effort
                    )
                    for chunk in response:
                        if chunk.text:
                            complete_response.append(chunk.text)
                            yield chunk.text

                # Store the complete Q&A in history after streaming is done
                final_answer = "".join(complete_response)
                conn = sqlite3.connect(DATABASE_NAME)
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, retrieved_docs, final_answer)
                            VALUES (?, ?, ?)''', (question, None, final_answer))
                conn.commit()
                conn.close()

            except Exception as e:
                error_msg = f"Error: {str(e)}"
                yield error_msg
                # Store error in database
                conn = sqlite3.connect(DATABASE_NAME)
                c = conn.cursor()
                c.execute('''INSERT INTO qa_history (question, final_answer)
                            VALUES (?, ?)''', (question, error_msg))
                conn.commit()
                conn.close()

        return Response(stream_with_context(generate()), mimetype='text/plain')

    except Exception as e:
        print(f"Error in ask_question: {str(e)}")
        return jsonify({'error': str(e)}), 500

@login_required
@role_required('Admin')
def update_index_api():
    """Manually refresh the Pinecone & BM25 index."""
    try:
        # Rebuild BM25 index
        build_bm25_index(POLICIES_FOLDER)
        
        # Repopulate Pinecone
        populate_pinecone_index()
        
        return jsonify({"message": "Indexes updated successfully"}), 200
    except Exception as e:
        logging.error(f"❌ Index Update Error: {e}", exc_info=True)
        return jsonify({"error": "Failed to update indexes"}), 500

# Resume Evaluator Routes
def extract_json_from_text(text):
    """Extract JSON object from text, handling deep nesting properly"""
    # Find the first { and match with its corresponding }
    start_idx = text.find('{')
    if start_idx == -1:
        return None
    
    # Count braces to find matching closing brace
    brace_count = 0
    in_string = False
    escape_next = False
    
    for i in range(start_idx, len(text)):
        char = text[i]
        
        # Handle string detection (to ignore braces inside strings)
        if char == '\\' and not escape_next:
            escape_next = True
            continue
        
        if char == '"' and not escape_next:
            in_string = not in_string
        
        escape_next = False
        
        # Count braces only when not inside a string
        if not in_string:
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    # Found matching closing brace
                    return text[start_idx:i+1]
    
    return None


def _stream_resume_eval_llm(formatted_prompt, eval_overrides, preview_queue):
    """
    Stream main resume-eval LLM call; push match_preview tuples to preview_queue.
    Returns parsed dict (same shape as async_gemini_generate).
    """
    cache_key = hashlib.sha256(
        f"{eval_overrides.get('provider_override')}|{eval_overrides.get('model_override')}|"
        f"{eval_overrides.get('groq_reasoning_effort_override')}|"
        f"{eval_overrides.get('max_completion_tokens_override')}|"
        f"{eval_overrides.get('temperature_override')}|{eval_overrides.get('top_p_override')}|"
        f"{json.dumps(eval_overrides.get('extra_body_override'), sort_keys=True) if eval_overrides.get('extra_body_override') else ''}|"
        f"{formatted_prompt}".encode("utf-8")
    ).hexdigest()
    cached_response = _cache_get(cache_key)
    if cached_response is not None:
        cached = json.loads(json.dumps(cached_response))
        match_str = cached.get("JD Match", "0%")
        try:
            pct = int(str(match_str).strip().strip("%") or 0)
        except (TypeError, ValueError):
            pct = 0
        preview_queue.put(("preview", pct, f"{pct}%"))
        return cached

    stream = generate_content_unified(
        formatted_prompt,
        stream=True,
        provider_override=eval_overrides.get("provider_override"),
        model_override=eval_overrides.get("model_override"),
        groq_reasoning_effort_override=eval_overrides.get("groq_reasoning_effort_override"),
        max_completion_tokens_override=eval_overrides.get("max_completion_tokens_override"),
        temperature_override=eval_overrides.get("temperature_override"),
        top_p_override=eval_overrides.get("top_p_override"),
        extra_body_override=eval_overrides.get("extra_body_override"),
    )

    parts = []
    preview_sent = False
    field_previews_sent = set()
    for piece in stream:
        text_chunk = getattr(piece, "text", piece) if not isinstance(piece, str) else piece
        if not text_chunk:
            continue
        parts.append(text_chunk)
        accumulated = "".join(parts)
        if not preview_sent:
            pct, pct_str = try_extract_jd_match_preview(accumulated)
            if pct is not None:
                preview_sent = True
                preview_queue.put(("preview", pct, pct_str))
        if preview_queue is not None:
            for field_key, snippet in try_extract_eval_field_previews(
                accumulated, field_previews_sent
            ):
                preview_queue.put(("field", field_key, snippet))

    response_text = "".join(parts).strip()
    if not response_text:
        raise ValueError("AI model returned empty evaluation response")
    try:
        parsed = parse_llm_json_response(response_text)
        if "JD Match" in str(formatted_prompt) and "JD Match" not in parsed:
            raise ValueError("parsed evaluation JSON missing JD Match")
        _cache_set(cache_key, parsed)
        return parsed
    except (ValueError, json.JSONDecodeError) as parse_err:
        logging.error("Streamed eval JSON parse failed: %s", parse_err)
        if strict_eval_json():
            raise ValueError(
                "Could not parse AI evaluation response. Please retry."
            ) from parse_err
        return get_default_resume_evaluation()


async def async_gemini_generate(
    prompt,
    provider_override=None,
    model_override=None,
    groq_reasoning_effort_override=None,
    max_completion_tokens_override=None,
    temperature_override=None,
    top_p_override=None,
    extra_body_override=None
):
    """Async wrapper for model generation (Gemini or OpenAI) with improved JSON handling"""
    try:
        cache_key = hashlib.sha256(
            f"{provider_override}|{model_override}|{groq_reasoning_effort_override}|"
            f"{max_completion_tokens_override}|{temperature_override}|{top_p_override}|"
            f"{json.dumps(extra_body_override, sort_keys=True) if extra_body_override else ''}|{prompt}".encode("utf-8")
        ).hexdigest()
        cached_response = _cache_get(cache_key)
        if cached_response is not None:
            return json.loads(json.dumps(cached_response))

        response = generate_content_unified(
            prompt,
            provider_override=provider_override,
            model_override=model_override,
            groq_reasoning_effort_override=groq_reasoning_effort_override,
            max_completion_tokens_override=max_completion_tokens_override,
            temperature_override=temperature_override,
            top_p_override=top_p_override,
            extra_body_override=extra_body_override
        )
        
        # Check if response has text attribute
        if not hasattr(response, 'text'):
            logging.error(f"❌ Response object has no 'text' attribute. Type: {type(response)}")
            logging.error(f"   Response attributes: {dir(response)}")
            return _fallback_resume_eval(prompt, "Model returned no text attribute")
        
        response_text = response.text
        
        # Check if response_text is None or empty
        if response_text is None:
            logging.error("❌ response.text is None")
            if "JD Match" in str(prompt):
                return _fallback_resume_eval(prompt)
            else:
                return get_default_career_analysis()
        
        if not isinstance(response_text, str):
            logging.error(f"❌ response.text is not a string: {type(response_text)}")
            if "JD Match" in str(prompt):
                return _fallback_resume_eval(prompt)
            else:
                return get_default_career_analysis()
        
        response_text = response_text.strip()
        
        # Reduced logging for performance (only log on errors)
        # logging.info(f"✅ Raw response length: {len(response_text)} chars")
        # logging.info(f"📄 Raw response preview (first 300): {response_text[:300]}...")
        # logging.info(f"📄 Raw response preview (last 200): ...{response_text[-200:]}")
        
        # Remove markdown code block markers if present (do this before brace checks)
        response_text = re.sub(r'^```json\s*', '', response_text, flags=re.MULTILINE)
        response_text = re.sub(r'^```\s*', '', response_text, flags=re.MULTILINE)
        response_text = re.sub(r'\s*```$', '', response_text)
        response_text = response_text.strip()

        # If still wrapped in fences, remove them greedily
        if response_text.startswith("```") and response_text.endswith("```"):
            response_text = response_text[3:-3].strip()

        # Check if response looks like it starts with explanatory text
        first_char = response_text.strip()[0] if response_text.strip() else ''
        if first_char != '{':
            logging.warning(f"⚠️ Response doesn't start with '{{'. First 100 chars: {response_text[:100]}")
            # Try to find where JSON actually starts
            json_start = response_text.find('{')
            if json_start > 0:
                logging.info(f"   Found '{{' at position {json_start}, removing {json_start} chars before it")
                response_text = response_text[json_start:]
                logging.info(f"   After trim, first 100 chars: {response_text[:100]}")

        # Final check - must start with {
        if not response_text.startswith('{'):
            logging.error(f"❌ Response still doesn't start with '{{' after cleaning")
            logging.error(f"   First 200 chars: {response_text[:200]}")
            if "JD Match" in str(prompt):
                return _fallback_resume_eval(prompt)
            else:
                return get_default_career_analysis()
        
        # Try to parse as JSON directly first
        try:
            parsed = json.loads(response_text)
            # logging.info("✅ Direct JSON parse successful")  # Reduced logging
            # Final validation - must be a dict
            if not isinstance(parsed, dict):
                logging.error(f"❌ Parsed JSON is not a dict: {type(parsed)}")
                if "JD Match" in str(prompt):
                    return _fallback_resume_eval(prompt)
                else:
                    return get_default_career_analysis()
            _cache_set(cache_key, parsed)
            return parsed
        except json.JSONDecodeError as e:
            # Only log if extraction also fails
            # logging.warning(f"⚠️ Direct JSON parse failed: {str(e)}")
            # logging.warning(f"   Error at position {e.pos}: {response_text[max(0,e.pos-20):e.pos+20]}")
            
            # Try to extract JSON object with proper brace matching
            json_str = extract_json_from_text(response_text)
            if json_str:
                # Validate extracted string looks like JSON
                json_str = json_str.strip()
                if not json_str.startswith('{') or not json_str.endswith('}'):
                    logging.error(f"❌ Extracted string doesn't look like JSON")
                    logging.error(f"   Starts with: {json_str[:50]}")
                    logging.error(f"   Ends with: {json_str[-50:]}")
                    if "JD Match" in str(prompt):
                        return _fallback_resume_eval(prompt)
                    else:
                        return get_default_career_analysis()
                
                try:
                    parsed = json.loads(json_str)
                    # logging.info(f"✅ Extracted JSON successfully (length: {len(json_str)} chars)")  # Reduced logging
                    # Validate it's a dict with expected structure
                    if not isinstance(parsed, dict):
                        logging.error(f"❌ Parsed JSON is not a dict: {type(parsed)}")
                        logging.error(f"   Parsed value: {repr(parsed)[:200]}")
                        if "JD Match" in str(prompt):
                            return _fallback_resume_eval(prompt)
                        else:
                            return get_default_career_analysis()
                    # Double-check it has expected keys for resume evaluation
                    if "JD Match" in str(prompt) and "JD Match" not in parsed:
                        logging.error(f"❌ Parsed dict missing 'JD Match' key")
                        logging.error(f"   Available keys: {list(parsed.keys())}")
                        return _fallback_resume_eval(prompt, "Missing JD Match in parsed JSON")
                    _cache_set(cache_key, parsed)
                    return parsed
                except json.JSONDecodeError as e2:
                    logging.error(f"❌ Failed to parse extracted JSON: {str(e2)}")
                    logging.error(f"   Error position: {e2.pos}")
                    logging.error(f"   Extracted text (first 500 chars): {json_str[:500]}...")
                    logging.error(f"   Extracted text (last 200 chars): ...{json_str[-200:]}")
                    logging.error(f"   Extracted text around error: ...{json_str[max(0,e2.pos-50):e2.pos+50]}...")
                    # Check if this is a resume evaluation prompt (has "JD Match" in prompt)
                    if "JD Match" in str(prompt):
                        return _fallback_resume_eval(prompt)
                    else:
                        return get_default_career_analysis()
            else:
                logging.error(f"❌ No JSON object found in response")
                logging.error(f"   Response text (first 500 chars): {response_text[:500]}...")
                # Check if this is a resume evaluation prompt
                if "JD Match" in str(prompt):
                    return get_default_resume_evaluation()
                else:
                    return get_default_career_analysis()
                
    except Exception as e:
        logging.error(f"❌ Model generation error: {str(e)}", exc_info=True)
        logging.error(f"   Error type: {type(e).__name__}")
        logging.error(f"   Error message: {str(e)}")
        # Check if this is a resume evaluation prompt
        if "JD Match" in str(prompt):
            return _fallback_resume_eval(prompt, str(e))
        else:
            return get_default_career_analysis()

async def async_analyze_stability(resume_text):
    """Async job stability analysis"""
    try:
        stability_prompt = job_stability_prompt.format(resume_text=resume_text)
        response = await async_gemini_generate(stability_prompt)
        
        if not response:
            raise ValueError("Failed to get stability analysis")
            
        # Ensure all required fields exist
        default_data = {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_analyze_stability: {str(e)}")
        return {
            "IsStable": True,
            "AverageJobTenure": "Unknown",
            "JobCount": 0,
            "StabilityScore": 0,
            "ReasoningExplanation": "Could not analyze job stability",
            "RiskLevel": "Unknown"
        }

async def async_generate_questions(resume_text, job_description, profile_summary):
    """Async interview questions generation"""
    try:
        questions_prompt = interview_questions_prompt.format(
            resume_text=resume_text,
            job_description=job_description,
            profile_summary=profile_summary
        )
        response = await async_gemini_generate(questions_prompt)
        
        if not response:
            raise ValueError("Failed to generate interview questions")
            
        # Ensure we have the required fields with proper defaults
        default_data = {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }
        
        # Merge response with defaults
        for key, default_value in default_data.items():
            if key not in response:
                response[key] = default_value
            elif not isinstance(response[key], list):
                response[key] = [str(response[key])] if response[key] else []
                
        return response
        
    except Exception as e:
        logging.error(f"Error in async_generate_questions: {str(e)}")
        return {
            "TechnicalQuestions": [],
            "NonTechnicalQuestions": []
        }

async def async_generate_recruiter_handbook(resume_text, job_description):
    """Async recruiter handbook generation - returns markdown text"""
    try:
        handbook_prompt = recruiter_handbook_prompt.format(
            resume_text=resume_text,
            job_description=job_description
        )
        
        # Use selected model to generate the recruiter handbook (run in thread pool to avoid blocking)
        # This returns markdown text, not JSON
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, lambda: generate_content_unified(handbook_prompt))
        response_text = response.text.strip()
        
        if not response_text:
            raise ValueError("Failed to generate recruiter handbook")
        
        # Return just the markdown content (JSON summary removed as per user request)
        return {
            "markdown_content": response_text,
            "json_summary": None
        }
        
    except Exception as e:
        logging.error(f"Error in async_generate_recruiter_handbook: {str(e)}")
        return {
            "markdown_content": f"## Error\n\nFailed to generate recruiter handbook: {str(e)}",
            "json_summary": None
        }

@limiter.limit("15 per minute")
@login_required
async def evaluate_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        file_path, filename = store_uploaded_file(
            file, app.config['UPLOAD_FOLDER'], ALLOWED_EXTENSIONS
        )

        # Extract contact info from raw resume (before sanitization)
        contact_info = extract_contact_info_from_file(file_path)

        # Extract text off the request thread (pdfplumber / docx can block)
        resume_text, extract_err = _IO_EXECUTOR.submit(
            extract_text_from_file, file_path, return_error=True
        ).result()
        if resume_text is None:
            msg = extract_err or 'Failed to extract text from file'
            # Extraction failures due to unsupported/invalid user uploads should be 400 (not server error)
            lowered = str(msg).lower()
            status = 400 if any(
                k in lowered for k in (
                    'unsupported',
                    'invalid file',
                    'please upload',
                    'scanned',
                    'image-only',
                    'docx',
                    'doc file',
                    'pdf',
                    'convert',
                    'libreoffice',
                )
            ) else 500
            return jsonify({'error': msg}), status

        # Generate evaluation using evaluation-specific provider/model settings
        eval_overrides = get_evaluation_generation_overrides()
        formatted_prompt = input_prompt_template.format(
            resume_text=resume_text,
            job_description=job_description,
            additional_context_block=additional_context_block
        )
        
        try:
            if FAST_EVAL_MODE:
                # Keep one high-quality core call, use fast heuristics for secondary dimensions.
                main_response = await async_gemini_generate(
                    formatted_prompt,
                    **eval_overrides
                )
                stability_data = get_fast_stability_estimate(resume_text)
                career_data = get_fast_career_estimate(resume_text)
            else:
                # Run all analyses concurrently using asyncio.gather
                main_response, stability_data, career_data = await asyncio.gather(
                    async_gemini_generate(
                        formatted_prompt,
                        **eval_overrides
                    ),
                    async_analyze_stability(resume_text),
                    analyze_career_progression(resume_text)  # Now properly awaited
                )
            
            if not main_response:
                raise ValueError("Failed to get main evaluation response")
                
            if not career_data:
                career_data = {
                    "progression_score": 50,
                    "key_observations": ["Failed to analyze career progression"],
                    "career_path": [],
                    "red_flags": ["Analysis error"],
                    "reasoning": "Failed to process career data"
                }
                
        except Exception as e:
            logging.error(f"Error during concurrent analysis: {str(e)}")
            return jsonify({'error': 'Failed to analyze resume'}), 500
        
        # Extract values from main response
        match_percentage_str = main_response.get("JD Match", "0%")
        match_percentage = int(match_percentage_str.strip('%'))
        missing_keywords = main_response.get("MissingKeywords", [])
        profile_summary = main_response.get("Profile Summary", "No summary provided.")
        over_under_qualification = main_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
        match_factors = main_response.get("Match Factors", {})
        candidate_fit_analysis = main_response.get("Candidate Fit Analysis", {})

        # Prepare additional information
        additional_info = {
            "job_stability": stability_data,
            "career_progression": career_data,
            "reasoning": main_response.get("Reasoning", "")
        }

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())

        # Get user email from session
        user_email = session.get('user', {}).get('email') if 'user' in session else None
        
        # Save evaluation to database with additional info
        db_id = save_evaluation(eval_id, filename, job_title, match_percentage, missing_keywords, profile_summary, match_factors, stability_data, additional_info, None, candidate_fit_analysis, over_under_qualification, user_email)
        if db_id:
            # Use backend-default questions (fast path, no additional LLM call)
            technical_questions, nontechnical_questions = get_default_interview_questions(job_title)
            technical_questions = (technical_questions or [])[:5]
            nontechnical_questions = (nontechnical_questions or [])[:5]
            behavioral_questions = (QUICK_CHECKS or [])[:5]

            # Save interview questions with proper JSON encoding (use database ID)
            if save_interview_questions(db_id, 
                                     json.dumps(technical_questions), 
                                     json.dumps(nontechnical_questions), 
                                     json.dumps(behavioral_questions)):
                return jsonify({
                    'id': eval_id,
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords,
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors,
                    'candidate_fit_analysis': candidate_fit_analysis,
                    'job_stability': stability_data,
                    'career_progression': career_data,
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions,
                    'contact_info': contact_info
                })
            else:
                return jsonify({'error': 'Failed to save interview questions'}), 500
        else:
            return jsonify({'error': 'Failed to save evaluation'}), 500

    except Exception as e:
        logging.error(f"Error in evaluate_resume: {str(e)}")
        return jsonify({'error': str(e)}), 500

def _fallback_resume_eval(prompt, detail: str = ""):
    """Strict mode raises on parse failure; otherwise return placeholder eval."""
    if strict_eval_json() and "JD Match" in str(prompt):
        raise ValueError(detail or "AI response could not be parsed. Please retry.")
    return get_default_resume_evaluation()


@limiter.limit("10 per minute")
@login_required
async def evaluate_resume_stream():
    """Streaming resume evaluation (async SSE; overlapping LLM tasks)."""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        oorwin_job_id = request.form.get('oorwin_job_id', '').strip()  # Get JobID from form
        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        if not job_title or not job_description:
            return jsonify({'error': 'Missing job title or description'}), 400

        file_path, filename = store_uploaded_file(
            file, app.config['UPLOAD_FOLDER'], ALLOWED_EXTENSIONS
        )

        # Extract contact info from raw resume (before sanitization)
        contact_info = extract_contact_info_from_file(file_path)

        # Extract text off the event loop (pdfplumber / docx can block)
        resume_text, extract_err = await asyncio.get_event_loop().run_in_executor(
            _IO_EXECUTOR,
            lambda: extract_text_from_file(file_path, return_error=True),
        )
        if resume_text is None:
            msg = extract_err or 'Failed to extract text from file'
            lowered = str(msg).lower()
            status = 400 if any(
                k in lowered for k in (
                    'unsupported',
                    'invalid file',
                    'please upload',
                    'scanned',
                    'image-only',
                    'docx',
                    'doc file',
                    'pdf',
                    'convert',
                    'libreoffice',
                )
            ) else 500
            return jsonify({'error': msg}), status

        # Generate unique ID for evaluation
        eval_id = str(uuid.uuid4())
        
        # Get user email from session
        user_email = session.get('user', {}).get('email') if 'user' in session else None
        
        # Track start time for performance metrics
        import time
        start_time = time.time()
        eval_overrides = get_evaluation_generation_overrides()

        async def generate():
            stability_task = None
            career_task = None

            try:
                # Send initial response
                yield f"data: {json.dumps({'status': 'processing', 'message': 'Analyzing resume...', 'eval_id': eval_id})}\n\n"

                # Kick off the secondary analyses NOW. They run concurrently with
                # the (slower) main JD-match call below. In FAST_EVAL_MODE we use
                # local heuristics and skip scheduling these tasks.
                if not FAST_EVAL_MODE:
                    stability_task = asyncio.create_task(async_analyze_stability(resume_text))
                    career_task = asyncio.create_task(analyze_career_progression(resume_text))

                # Step 1: Main resume analysis (most important - show results immediately)
                yield f"data: {json.dumps({'status': 'step1', 'message': 'Evaluating resume against job requirements...'})}\n\n"
                formatted_prompt = input_prompt_template.format(
                    resume_text=resume_text,
                    job_description=job_description,
                    additional_context_block=additional_context_block
                )
                preview_queue = queue.Queue()
                def _run_main_with_done():
                    try:
                        return _stream_resume_eval_llm(
                            formatted_prompt, eval_overrides, preview_queue
                        )
                    finally:
                        preview_queue.put(("_done", None, None))

                main_task = asyncio.create_task(
                    asyncio.to_thread(_run_main_with_done)
                )

                match_preview_sent = False
                field_previews_sse = set()
                while not main_task.done() or not preview_queue.empty():
                    try:
                        item = await asyncio.to_thread(preview_queue.get, True, 0.12)
                    except queue.Empty:
                        continue
                    kind = item[0]
                    if kind == "preview" and not match_preview_sent:
                        match_preview_sent = True
                        _, pct, pct_str = item
                        yield f"data: {json.dumps({'status': 'match_preview', 'match_percentage': pct, 'match_percentage_str': pct_str})}\n\n"
                    elif kind == "field":
                        _, field_key, snippet = item
                        if field_key not in field_previews_sse:
                            field_previews_sse.add(field_key)
                            yield f"data: {json.dumps({'status': 'eval_field_preview', 'field': field_key, 'snippet': snippet})}\n\n"
                    elif kind == "_done":
                        break

                try:
                    main_response = await asyncio.wait_for(main_task, timeout=600)
                except Exception as gen_error:
                    logging.error(
                        "Error in streamed resume eval: %s: %s",
                        type(gen_error).__name__,
                        gen_error,
                        exc_info=True,
                    )
                    yield f"data: {json.dumps({'status': 'error', 'message': f'AI generation failed: {str(gen_error)}'})}\n\n"
                    return
                
                # Debug logging (reduced for performance)
                # logging.info(f"🔍 Main response type: {type(main_response)}")
                # logging.info(f"🔍 Main response value (first 200 chars): {str(main_response)[:200]}")
                
                if not main_response:
                    logging.error("❌ main_response is empty/None")
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to analyze resume - empty response'})}\n\n"
                    return
                
                # Verify it's a dictionary
                if not isinstance(main_response, dict):
                    logging.error(f"❌ main_response is not a dict!")
                    logging.error(f"   Type: {type(main_response)}")
                    logging.error(f"   Value (repr): {repr(main_response)[:500]}")
                    logging.error(f"   Value (str): {str(main_response)[:500]}")
                    logging.error(f"   This suggests JSON parsing failed and returned a string fragment")
                    yield f"data: {json.dumps({'status': 'error', 'message': f'Invalid response format from AI. Expected dict, got {type(main_response).__name__}'})}\n\n"
                    return
                
                # Verify it has required keys
                required_keys = ["JD Match", "Match Factors", "Profile Summary"]
                missing_keys = [key for key in required_keys if key not in main_response]
                if missing_keys:
                    logging.error(f"❌ Missing required keys in response: {missing_keys}")
                    logging.error(f"   Available keys: {list(main_response.keys())}")
                    yield f"data: {json.dumps({'status': 'error', 'message': f'AI response missing required fields: {missing_keys}'})}\n\n"
                    return
                
                # logging.info(f"✅ Main response validated - has {len(main_response)} keys")  # Reduced logging
                
                # Extract basic values with error handling
                try:
                    # Double-check main_response is a dict before accessing
                    if not isinstance(main_response, dict):
                        error_msg = f"main_response is {type(main_response)}, not dict. Value: {repr(str(main_response)[:200])}"
                        logging.error(f"❌ {error_msg}")
                        raise TypeError(error_msg)
                    
                    # Safely get JD Match
                    if "JD Match" not in main_response:
                        logging.error(f"❌ 'JD Match' key not found in main_response")
                        logging.error(f"   Available keys: {list(main_response.keys())}")
                        logging.error(f"   main_response value: {repr(main_response)[:500]}")
                        raise KeyError("'JD Match' key not found in response")
                    
                    match_percentage_str = main_response.get("JD Match", "0%")
                    if match_percentage_str is None:
                        logging.warning("JD Match is None, using default")
                        match_percentage_str = "0%"
                    
                    if not isinstance(match_percentage_str, str):
                        logging.warning(f"JD Match is not a string: {type(match_percentage_str)}, value: {match_percentage_str}")
                        match_percentage_str = str(match_percentage_str) if match_percentage_str else "0%"
                    
                    # Clean and extract percentage
                    match_percentage_str = match_percentage_str.strip()
                    if not match_percentage_str:
                        match_percentage_str = "0%"
                    match_percentage = int(match_percentage_str.strip('%'))
                except (ValueError, AttributeError, TypeError) as e:
                    logging.error(f"❌ Error parsing match percentage: {e}")
                    logging.error(f"   main_response type: {type(main_response)}")
                    logging.error(f"   main_response value: {repr(str(main_response)[:300])}")
                    logging.error(f"   match_percentage_str: {repr(match_percentage_str) if 'match_percentage_str' in locals() else 'not defined'}")
                    match_percentage = 0
                    match_percentage_str = "0%"
                
                missing_keywords = main_response.get("MissingKeywords", [])
                profile_summary = main_response.get("Profile Summary", "No summary provided.")
                over_under_qualification = main_response.get("Over/UnderQualification Analysis", "No qualification mismatch concerns detected.")
                match_factors = main_response.get("Match Factors", {})
                candidate_fit_analysis = main_response.get("Candidate Fit Analysis", {})
                
                # Send basic results immediately (user sees this in ~5-8 seconds instead of 20)
                basic_results = {
                    'status': 'basic_results',
                    'id': eval_id,
                    'match_percentage': match_percentage,
                    'match_percentage_str': match_percentage_str,
                    'missing_keywords': missing_keywords if isinstance(missing_keywords, list) else [],
                    'profile_summary': profile_summary,
                    'over_under_qualification': over_under_qualification,
                    'match_factors': match_factors if isinstance(match_factors, dict) else {},
                    'candidate_fit_analysis': candidate_fit_analysis if isinstance(candidate_fit_analysis, dict) else {},
                    'contact_info': contact_info
                }
                
                # Test JSON serialization before yielding
                try:
                    json_str = json.dumps(basic_results)
                    yield f"data: {json_str}\n\n"
                except (TypeError, ValueError) as json_err:
                    logging.error(f"JSON serialization error: {json_err}")
                    logging.error(f"Problematic data types: match_factors={type(match_factors)}, candidate_fit_analysis={type(candidate_fit_analysis)}")
                    # Send error with details
                    yield f"data: {json.dumps({'status': 'error', 'message': f'Data serialization error: {str(json_err)}'})}\n\n"
                    return
                
                # Step 2 + Step 3: fast-path mode for sub-10s UX
                if FAST_EVAL_MODE:
                    yield f"data: {json.dumps({'status': 'step2', 'message': 'Running fast stability/career estimation...'})}\n\n"
                    stability_data = get_fast_stability_estimate(resume_text)
                    career_data = get_fast_career_estimate(resume_text)
                else:
                    # PERF FIX: these tasks were scheduled at the top of generate()
                    # and have been running concurrently with the JD-match call.
                    # In practice they're often already done by this point.
                    yield f"data: {json.dumps({'status': 'step2', 'message': 'Finalizing stability and progression analyses...'})}\n\n"
                    try:
                        stability_data, career_data = await asyncio.gather(
                            stability_task, career_task
                        )
                    except Exception as followup_err:
                        logging.error(f"❌ Concurrent follow-up analyses failed: {followup_err}", exc_info=True)
                        stability_data = {
                            "IsStable": True,
                            "AverageJobTenure": "Unknown",
                            "JobCount": 0,
                            "StabilityScore": 0,
                            "ReasoningExplanation": "Stability analysis failed",
                            "RiskLevel": "Unknown",
                        }
                        career_data = None

                if not career_data:
                    career_data = {
                        "progression_score": 50,
                        "key_observations": ["Failed to analyze career progression"],
                        "career_path": [],
                        "red_flags": ["Analysis error"],
                        "reasoning": "Failed to process career data"
                    }

                # Send stability and career data as soon as parallel step completes
                additional_data = {
                    'status': 'additional_data',
                    'job_stability': stability_data,
                    'career_progression': career_data
                }
                yield f"data: {json.dumps(additional_data)}\n\n"

                # Step 3: Use backend-default questions (no extra LLM latency)
                yield f"data: {json.dumps({'status': 'step3', 'message': 'Loading interview question bank...'})}\n\n"
                technical_questions, nontechnical_questions = get_default_interview_questions(job_title)
                technical_questions = (technical_questions or [])[:5]
                nontechnical_questions = (nontechnical_questions or [])[:5]
                behavioral_questions = (QUICK_CHECKS or [])[:5]
                
                # Send questions
                questions_data_response = {
                    'status': 'questions',
                    'technical_questions': technical_questions,
                    'nontechnical_questions': nontechnical_questions,
                    'behavioral_questions': behavioral_questions
                }
                yield f"data: {json.dumps(questions_data_response)}\n\n"
                
                # Step 4: Save to database
                yield f"data: {json.dumps({'status': 'step4', 'message': 'Saving results...'})}\n\n"
                
                additional_info = {
                    "job_stability": stability_data,
                    "career_progression": career_data,
                    "reasoning": main_response.get("Reasoning", "")
                }
                
                # Calculate time taken
                time_taken = round(time.time() - start_time, 2)  # Round to 2 decimal places
                
                # Debug: Log the data being saved
                logging.info(f"Attempting to save evaluation: eval_id={eval_id}, filename={filename}, job_title={job_title}, oorwin_job_id={oorwin_job_id}, user_email={user_email}, time_taken={time_taken}s")
                
                db_id = save_evaluation(eval_id, filename, job_title, match_percentage, missing_keywords, profile_summary, match_factors, stability_data, additional_info, oorwin_job_id, candidate_fit_analysis, over_under_qualification, user_email, time_taken)
                logging.info(f"Save evaluation result: db_id={db_id}, oorwin_job_id={oorwin_job_id}")
                
                if db_id:
                    # Use the database ID (integer) for saving interview questions
                    if save_interview_questions(db_id, 
                                             json.dumps(technical_questions), 
                                             json.dumps(nontechnical_questions), 
                                             json.dumps(behavioral_questions)):
                        # Send the database ID back to frontend for feedback submission
                        yield f"data: {json.dumps({'status': 'complete', 'message': 'Analysis complete!', 'db_id': db_id})}\n\n"
                    else:
                        yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save interview questions'})}\n\n"
                else:
                    yield f"data: {json.dumps({'status': 'error', 'message': 'Failed to save evaluation'})}\n\n"
                    
            except Exception as e:
                error_type = type(e).__name__
                error_msg = str(e)
                logging.error(f"❌ Error in streaming evaluation: {error_type}: {error_msg}")
                logging.error(f"   Error repr: {repr(error_msg)}")
                logging.error(f"   Error type: {error_type}")
                logging.error(f"   Full traceback:", exc_info=True)
                # Make sure error message is JSON-safe
                safe_error_msg = error_msg[:500] if len(error_msg) > 500 else error_msg
                yield f"data: {json.dumps({'status': 'error', 'message': safe_error_msg, 'error_type': error_type})}\n\n"
            finally:
                for _t in (stability_task, career_task):
                    if _t is not None and not _t.done():
                        _t.cancel()

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        logging.error(f"Error in evaluate_resume_stream: {str(e)}")
        return jsonify({'error': str(e)}), 500

@login_required
def get_interview_questions(evaluation_id):
    """Get interview questions for a specific evaluation"""
    conn = None
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # First, get the evaluation details to regenerate questions if needed
        cursor.execute(
            """
            SELECT e.resume_path, e.job_title, e.job_description, e.profile_summary 
            FROM evaluations e 
            WHERE e.id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        # Then get existing questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        result = cursor.fetchone()
        
        # Initialize default values
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        if result:
            try:
                # Parse saved questions with proper error handling
                def parse_json_safely(json_str):
                    if not json_str:
                        return []
                    try:
                        data = json.loads(json_str)
                        if isinstance(data, list):
                            return data
                        elif isinstance(data, str):
                            try:
                                return json.loads(data)
                            except:
                                return [data]
                        else:
                            return [str(data)]
                    except json.JSONDecodeError:
                        try:
                            # Try to clean and parse the string
                            cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                            items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                            return [item for item in items if item]
                        except:
                            return []

                technical_questions = parse_json_safely(result[0])
                nontechnical_questions = parse_json_safely(result[1])
                behavioral_questions = parse_json_safely(result[2])
            except Exception as e:
                logging.error(f"Error parsing interview questions: {str(e)}")
                technical_questions = []
                nontechnical_questions = []
                behavioral_questions = []

        # Only regenerate questions if they are completely missing (not just empty)
        # This prevents regenerating questions when they exist but are empty arrays
        if not result and eval_result:
            logging.info(f"No interview questions found in database for evaluation {evaluation_id}, generating new ones")
            resume_path = eval_result[0]
            if not resume_path or resume_path == 'NULL' or resume_path == 'None':
                logging.error(f"Invalid resume_path for evaluation {evaluation_id}: {resume_path}")
                return jsonify({'error': 'Resume file path is invalid'}), 400
            resume_text = extract_text_from_file(resume_path)
            if resume_text:
                        questions_data = asyncio.run(async_generate_questions(
                            resume_text,
                            eval_result[2],  # job_description
                            eval_result[3]   # profile_summary
                        ))
                        
            technical_questions = questions_data.get("TechnicalQuestions", [])
            nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
            behavioral_questions = QUICK_CHECKS
                        
             # Save regenerated questions
            cursor.execute(
                            """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (evaluation_id,
                     json.dumps(technical_questions), 
                             json.dumps(nontechnical_questions), 
                     json.dumps(behavioral_questions))
                        )
            conn.commit()
            logging.info(f"Generated and saved new questions for evaluation {evaluation_id}")

            return jsonify({
                    'technical_questions': technical_questions or ["No technical questions available"],
                    'nontechnical_questions': nontechnical_questions or ["No non-technical questions available"],
                    'behavioral_questions': behavioral_questions or QUICK_CHECKS
                })

    except Exception as e:
        logging.error(f"Database error in get_interview_questions: {str(e)}")
        return jsonify({
            'technical_questions': ["Error loading technical questions"],
            'nontechnical_questions': ["Error loading non-technical questions"],
            'behavioral_questions': QUICK_CHECKS
        })
    finally:
        if conn:
            conn.close()

# --- Document Processing ---
# Optimized chunking for HR policy documents:
# - Larger chunks (1200) preserve context and complete policy explanations
# - Higher overlap (250) ensures continuity across chunks
# - Better separators prioritize paragraph/sentence boundaries
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1200,        # Increased from 400 for better context retention
    chunk_overlap=250,      # Increased from 50 (20% overlap for continuity)
    separators=["\n\n", "\n", ". ", " ", ""]  # Better paragraph/sentence awareness
)

def dataframe_to_clean_markdown(df: pd.DataFrame) -> str:
    """Convert a pandas DataFrame into a clean GitHub-flavored markdown table.

    - Fills missing/None headers with generic names (Column 1, Column 2, ...)
    - Collapses multi-line cell content into single lines
    - Ensures string-typed cells
    - Removes the index column from the markdown output
    """
    try:
        # Defensive copy
        table_df = df.copy()

        # Normalize headers
        normalized_columns = []
        for i, col in enumerate(table_df.columns):
            col_str = str(col).strip() if col is not None else ""
            if col_str == "" or col_str.lower() == "none":
                col_str = f"Column {i+1}"
            normalized_columns.append(col_str)
        table_df.columns = normalized_columns

        # Drop index-like first column if it looks like 0..N sequence
        if table_df.shape[1] > 0:
            first_col_vals = list(table_df.iloc[:, 0])
            def looks_like_sequential_index(vals):
                try:
                    ints = [int(str(v)) for v in vals]
                    return ints == list(range(len(ints)))
                except Exception:
                    return False
            if looks_like_sequential_index(first_col_vals):
                table_df = table_df.iloc[:, 1:]

        # Drop columns that are entirely empty after stripping
        if table_df.shape[1] > 0:
            non_empty_cols = []
            for c in table_df.columns:
                col_str = table_df[c].astype(str).str.strip()
                if (col_str != "").any():
                    non_empty_cols.append(c)
            if non_empty_cols:
                table_df = table_df[non_empty_cols]

        # Normalize cell content: string type, collapse newlines/tabs, trim
        for c in table_df.columns:
            table_df[c] = (
                table_df[c]
                .astype(str)
                .str.replace("\r\n|\r|\n", " ", regex=True)
                .str.replace("\t", " ", regex=True)
                .str.replace("\\s+", " ", regex=True)
                .str.strip()
            )

        # Render as markdown without index
        return table_df.to_markdown(index=False, tablefmt="pipe")
    except Exception as e:
        logging.warning(f"Failed to render markdown table cleanly, falling back: {e}")
        # Fallback to basic to_markdown if anything goes wrong
        try:
            return df.to_markdown(index=False, tablefmt="pipe")
        except Exception:
            return df.to_markdown()

def process_pdf(pdf_path, documents, table_chunks):
    """Extract text and tables from a PDF file."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Extract and process text
                text = page.extract_text() or ""
                if text:
                    text_chunks = text_splitter.split_text(text)
                    documents.extend(text_chunks)
                
                # Extract and process tables
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:  # Ensure table has headers and data
                        df = pd.DataFrame(table[1:], columns=table[0])
                        table_markdown = dataframe_to_clean_markdown(df)
                        
                        # Enrich table for better BM25 matching (same as Pinecone indexing)
                        column_names = " ".join([str(col).lower() for col in df.columns if col])
                        sample_values = []
                        # Use positional access to avoid issues with non-standard/duplicate column labels
                        num_sample_cols = min(3, len(df.columns))
                        for col_idx in range(num_sample_cols):
                            series = df.iloc[:, col_idx]
                            sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                        sample_context = " ".join(sample_values[:10])
                        enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                        table_chunks.append(enriched_table)
    except Exception as e:
        logging.error(f"❌ Error processing PDF {pdf_path}: {e}")

def populate_pinecone_index():
    """Extract content from PDF documents and populate Pinecone index."""
    try:
        documents = []
        table_chunks = []
        texts = []
        metadatas = []
        
        # Get all PDF files from the policies folder
        if not os.path.exists(POLICIES_FOLDER):
            logging.warning(f"Policies folder {POLICIES_FOLDER} does not exist")
            return
            
        pdf_files = [f for f in os.listdir(POLICIES_FOLDER) if f.endswith('.pdf')]
        if not pdf_files:
            logging.warning(f"No PDF files found in {POLICIES_FOLDER}")
            return
            
        total_files = len(pdf_files)
        logging.info(f"📚 Processing {total_files} PDF files")
        
        for idx, filename in enumerate(pdf_files, 1):
            pdf_path = os.path.join(POLICIES_FOLDER, filename)
            logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        chunks = text_splitter.split_text(text)
                        documents.extend(chunks)
                        for chunk in chunks:
                            texts.append(chunk)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                        logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
                    
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Improve table representation for better retrieval:
                            # 1. Extract column names as descriptive keywords
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            
                            # 2. Extract sample data values as context
                            sample_values = []
                            # Use positional access to avoid issues with non-standard/duplicate column labels
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):  # First up to 3 columns
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            
                            # 3. Create enriched table chunk with context
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            texts.append(enriched_table)
                            metadatas.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
                            logging.info(f"   Page {page_num}: Added table {table_num} (enriched with context)")
    except Exception as e:
        logging.error(f"❌ Error in document processing: {str(e)}")
        raise

    try:
        all_chunks = documents + table_chunks
        total_chunks = len(texts)
        
        if total_chunks == 0:
            raise ValueError("No content extracted from documents")
        
        logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
        # Initialize Pinecone components
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index = pc.Index(PINECONE_INDEX_NAME)
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
        # Insert in batches
        batch_size = 50  # Reduced batch size for better reliability
        for i in range(0, total_chunks, batch_size):
            batch_texts = texts[i:i + batch_size]
            batch_metas = metadatas[i:i + batch_size]
            PineconeVectorStore.from_texts(
                texts=batch_texts,
                embedding=embeddings,
                index_name=PINECONE_INDEX_NAME,
                metadatas=batch_metas
            )
            logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
        # Verify insertion
        stats = index.describe_index_stats()
        vector_count = stats['total_vector_count']
        logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
    except Exception as e:
        logging.error(f"❌ Error in Pinecone operations: {str(e)}")
        raise
#         total_files = len(pdf_files)
#         logging.info(f"📚 Processing {total_files} PDF files")
        
#         for idx, filename in enumerate(pdf_files, 1):
#             pdf_path = os.path.join(POLICIES_FOLDER, filename)
#             logging.info(f"📄 Processing file {idx}/{total_files}: {filename}")
            
#             with pdfplumber.open(pdf_path) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):
#                     text = page.extract_text() or ""
#                     if text:
#                         chunks = text_splitter.split_text(text)
#                         documents.extend(chunks)
#                         logging.info(f"   Page {page_num}: Added {len(chunks)} text chunks")
#                     
#                     tables = page.extract_tables()
#                     for table_num, table in enumerate(tables, 1):
#                         if table and len(table) > 1:
#                             df = pd.DataFrame(table[1:], columns=table[0])
#                             table_chunks.append(df.to_markdown())
#                             logging.info(f"   Page {page_num}: Added table {table_num}")
#     except Exception as e:
#         logging.error(f"❌ Error in document processing: {str(e)}")
#         raise

#     try:
#     all_chunks = documents + table_chunks
#     total_chunks = len(all_chunks)
        
#     if total_chunks == 0:
#             raise ValueError("No content extracted from documents")
        
#     logging.info(f"📊 Preparing to insert {total_chunks} chunks into Pinecone")
        
#         # Initialize Pinecone components
#     pc = Pinecone(api_key=PINECONE_API_KEY)
#     index = pc.Index(PINECONE_INDEX_NAME)
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
#     # Insert in batches
#     batch_size = 50  # Reduced batch size for better reliability
#     for i in range(0, total_chunks, batch_size):
#             batch = all_chunks[i:i + batch_size]
#             PineconeVectorStore.from_texts(
#                 texts=batch,
#                 embedding=embeddings,
#                 index_name=PINECONE_INDEX_NAME
#             )
#             logging.info(f"✅ Inserted batch {i//batch_size + 1}/{(total_chunks-1)//batch_size + 1}")
        
#         # Verify insertion
#             stats = index.describe_index_stats()
#             vector_count = stats['total_vector_count']
#             logging.info(f"🎉 Successfully populated index with {vector_count} vectors")
        
#     except Exception as e:
#         logging.error(f"❌ Error in Pinecone operations: {str(e)}")
#         raise

def initialize_pinecone():
    """Initialize Pinecone. Create and populate index if it doesn't exist."""
    try:
        logging.info("🔧 Initializing Pinecone...")
        
        # Check if index exists
        pc = Pinecone(api_key=PINECONE_API_KEY)
        index_name = PINECONE_INDEX_NAME
        
        if index_name in pc.list_indexes().names():
            logging.info(f"📋 Index '{index_name}' already exists - using existing index")
            # Check if index has data
            index = pc.Index(index_name)
            stats = index.describe_index_stats()
            vector_count = stats['total_vector_count']
            
            if vector_count > 0:
                logging.info(f"✅ Index '{index_name}' has {vector_count} vectors - no need to populate")
                return True
            else:
                logging.info(f"⚠️ Index '{index_name}' exists but is empty - populating...")
                populate_pinecone_index()
                return True
        else:
            # Create new index only if it doesn't exist
            logging.info(f"🆕 Index '{index_name}' doesn't exist - creating new index...")
            pc.create_index(
                name=index_name,
                dimension=384,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
            
            # Wait for index to be ready
            time.sleep(10)
            logging.info(f"✅ Index '{index_name}' created successfully")
            
            # Populate the new index
            logging.info("📚 Populating new Pinecone index...")
            populate_pinecone_index()
            
            return True
        
    except Exception as e:
        logging.error(f"❌ Error initializing Pinecone: {str(e)}")
        return False

# --- BM25 Setup ---
bm25_index = None
bm25_corpus = None
bm25_metadata = []  # Store metadata for each BM25 chunk (filename, page, type)

def build_bm25_index(folder_path):
    """Builds BM25 index from policy documents with metadata tracking."""
    global bm25_index, bm25_corpus, bm25_metadata
    
    all_texts = []
    table_chunks = []
    text_metadata = []  # Track metadata for text chunks
    table_metadata = []  # Track metadata for table chunks
    
    # Process all PDF files with metadata
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            
            # Process PDF with metadata tracking
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text() or ""
                    if text:
                        text_chunks = text_splitter.split_text(text)
                        all_texts.extend(text_chunks)
                        # Add metadata for each text chunk
                        for chunk in text_chunks:
                            text_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "text"
                            })
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table and len(table) > 1:
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_markdown = dataframe_to_clean_markdown(df)
                            
                            # Enrich table for better BM25 matching
                            column_names = " ".join([str(col).lower() for col in df.columns if col])
                            sample_values = []
                            num_sample_cols = min(3, len(df.columns))
                            for col_idx in range(num_sample_cols):
                                series = df.iloc[:, col_idx]
                                sample_values.extend([str(val).lower() for val in series.dropna().head(3).tolist()])
                            sample_context = " ".join(sample_values[:10])
                            enriched_table = f"[TABLE DATA] Topic: {column_names} {sample_context}\n\n{table_markdown}\n\n[END TABLE]"
                            table_chunks.append(enriched_table)
                            # Add metadata for table chunk
                            table_metadata.append({
                                "source": filename,
                                "page": page_num,
                                "type": "table"
                            })
    
    # Combine text and tables for indexing
    all_chunks = all_texts + table_chunks
    bm25_metadata = text_metadata + table_metadata
    
    if all_chunks:
        # Tokenize for BM25
        bm25_corpus = [text.split() for text in all_chunks]
        bm25_index = BM25Okapi(bm25_corpus)
        logging.info(f"✅ BM25 index built with {len(bm25_corpus)} document chunks (with metadata tracking)")
    else:
        logging.warning("⚠️ No content found for BM25 indexing")

def expand_query_with_llm(question, llm):
    """Expands user query using LLM to include synonyms but retains original meaning."""
    expansion_prompt = f"""
    Provide alternative phrasings and related terms for: '{question}', 
    ensuring the original word is always included. Include HR-specific terms if applicable.
    """
    try:
        expanded_query = llm.invoke(expansion_prompt).content
        logging.info(f"🔍 Query Expansion: {expanded_query}")
        return expanded_query
    except Exception as e:
        logging.error(f"❌ Query Expansion Failed: {e}")
        return question  # Fall back to the original question

def hybrid_search(question, llm, retriever):
    """Performs hybrid retrieval using BM25 and Pinecone vectors."""
    global bm25_index, bm25_corpus
    
    # Expand query
    expanded_query = expand_query_with_llm(question, llm)
    
    results = []
    
    # Step 1: BM25 Keyword Search
    if bm25_index and bm25_corpus:
        bm25_results = bm25_index.get_top_n(expanded_query.split(), bm25_corpus, n=5)
        bm25_texts = [" ".join(text) for text in bm25_results]
        results.extend(bm25_texts)
        logging.info(f"🔍 BM25 Retrieved {len(bm25_texts)} results")
    
    # Step 2: Vector Search
    pinecone_results = retriever.invoke(expanded_query)
    pinecone_texts = [doc.page_content for doc in pinecone_results]
    results.extend(pinecone_texts)
    
    # Prioritize table content (tables contain | character in markdown)
    table_texts = [text for text in results if "|" in text]
    non_table_texts = [text for text in results if "|" not in text]
    
    # Combine results: tables first, then other content
    combined_results = table_texts + non_table_texts
    
    # Remove duplicates while preserving order
    unique_results = []
    seen = set()
    for text in combined_results:
        # Use a hash of the text as a unique identifier
        text_hash = hash(text)
        if text_hash not in seen:
            seen.add(text_hash)
            unique_results.append(text)
    
    # Join and truncate to avoid token limits
    final_text = "\n\n".join(unique_results)[:5000]
    
    return final_text

def save_qa_to_db(question, retrieved_docs, final_answer, feedback=None):
    """Stores a Q&A pair in SQLite with optional feedback."""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        logging.info(f"Saving Q&A to DB - Question: {question[:50]}...")  # Debug log
        
        query = """
        INSERT INTO qa_history (question, retrieved_docs, final_answer, feedback) 
        VALUES (?, ?, ?, ?)
        """
        cursor.execute(query, (question, retrieved_docs, final_answer, feedback))
        conn.commit()
        
        question_id = cursor.lastrowid
        logging.info(f"✅ Q&A stored successfully with ID: {question_id}")
        return question_id
    except Exception as e:
        logging.error(f"❌ Error saving Q&A to DB: {e}", exc_info=True)
        return None
    finally:
        conn.close()

def setup_llm_chain():
    """Initialize the LLM and retrieval chain."""
    # Initialize LLM with optimized parameters
    llm = ChatGroq(
        # model_name= "mixtral-8x7b-32768", 
        #  model_name= "llama-3.1-8b-instant",
        model_name =  "qwen/qwen3-32b",     #"qwen-2.5-32b",
        groq_api_key=GROQ_API_KEY,
        temperature=0.377,
        max_tokens=32768,
        top_p=0.95,
        presence_penalty=0.1,
        frequency_penalty=0.1,
        streaming=True
    )
    
    # Initialize retriever only if vectorstore is available
    retriever = None
    if vectorstore is not None:
        try:
            retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
            logging.info("✅ Retriever initialized successfully")
        except Exception as e:
            logging.error(f"❌ Error initializing retriever: {e}")
            retriever = None
    else:
        logging.warning("⚠️ Vectorstore not available, retriever will be None")
    
    return llm, None, retriever  # Return llm, qa_chain (None for now), retriever

def expand_acronyms(question):
    """Expand HR-related acronyms in the question."""
    expanded_question = question.lower()
    for acronym, full_form in ACRONYM_MAP.items():
        expanded_question = expanded_question.replace(acronym.lower(), full_form.lower())
    return expanded_question

async def analyze_career_progression(resume_text):
    """Analyze career progression from resume text using Gemini."""
    try:
        formatted_prompt = f"""You are an expert HR analyst. Analyze this candidate's career progression.
Return ONLY a JSON object with the following structure, no other text:
{{
    "progression_score": <number 0-100>,
    "key_observations": [<list of string observations>],
    "career_path": [
        {{
            "title": "<job title>",
            "company": "<company name>",
            "duration": "<time period>",
            "level": "<Entry/Mid/Senior/Lead/Manager>",
            "progression": "<Promotion/Lateral/Step Back>"
        }}
    ],
    "red_flags": [<list of string concerns>],
    "reasoning": "<analysis explanation>"
}}

Resume text:
{resume_text}"""

        # Get response from Gemini
        response = await async_gemini_generate(formatted_prompt)
        
        # If response is already a dict (from async_gemini_generate)
        if isinstance(response, dict):
            parsed_response = response
        else:
            try:
                parsed_response = json.loads(response) if isinstance(response, str) else {}
            except json.JSONDecodeError:
                logging.error(f"Failed to parse response as JSON: {response}")
                return get_default_career_analysis()

        # Validate and clean the response data
        cleaned_data = {
            "progression_score": validate_progression_score(parsed_response.get("progression_score", 50)),
            "key_observations": validate_list(parsed_response.get("key_observations", [])) or ["No key observations found"],
            "career_path": validate_career_path(parsed_response.get("career_path", [])),
            "red_flags": validate_list(parsed_response.get("red_flags", [])) or ["No red flags identified"],
            "reasoning": str(parsed_response.get("reasoning", "No analysis provided")).strip()
        }

        # Ensure we have valid data
        if cleaned_data["progression_score"] == 50 and not cleaned_data["career_path"]:
            return get_default_career_analysis()
            
        return cleaned_data

    except Exception as e:
        logging.error(f"Career progression analysis error: {str(e)}")
        logging.error(f"Full traceback:", exc_info=True)
        return get_default_career_analysis()

def get_default_career_analysis():
    """Return default career analysis structure"""
    return {
        "progression_score": 50,
        "key_observations": ["Unable to analyze career progression"],
        "career_path": [],
        "red_flags": ["Analysis encountered technical issues"],
        "reasoning": "Analysis failed to complete"
    }

def get_default_resume_evaluation():
    """Return default resume evaluation structure when JSON parsing fails"""
    return {
        "JD Match": "0%",
        "MissingKeywords": [],
        "Profile Summary": "Unable to analyze resume due to technical error. Please try again.",
        "Over/UnderQualification Analysis": "Analysis unavailable",
        "Match Factors": {
            "Skills Match": 0,
            "Experience Match": 0,
            "Education Match": 0,
            "Industry Knowledge": 0,
            "Certification Match": None
        },
        "Reasoning": "JSON parsing failed. The AI response could not be properly parsed.",
        "Candidate Fit Analysis": {
            "Dimension Evaluation": [],
            "Risk and Gaps": None,
            "Recommendation": {
                "Verdict": "❌ Analysis Failed",
                "Fit Level": "Unknown",
                "Rationale": "Technical error prevented proper evaluation"
            },
            "Recruiter Narrative": "Unable to generate evaluation due to technical issues."
        }
    }

def get_fast_stability_estimate(resume_text):
    """Heuristic stability estimate for fast evaluation mode (no LLM)."""
    text = (resume_text or "")
    # Rough signal: count year mentions as a proxy for role transitions.
    year_hits = re.findall(r"\b(?:19|20)\d{2}\b", text)
    estimated_roles = max(1, min(8, len(year_hits) // 2 if year_hits else 2))
    if estimated_roles <= 2:
        score = 82
        risk = "Low"
        tenure = "2.5+ years (estimated)"
    elif estimated_roles <= 4:
        score = 67
        risk = "Moderate"
        tenure = "1.5-2.5 years (estimated)"
    else:
        score = 48
        risk = "High"
        tenure = "<1.5 years (estimated)"

    return {
        "IsStable": score >= 60,
        "AverageJobTenure": tenure,
        "JobCount": estimated_roles,
        "StabilityScore": score,
        "ReasoningExplanation": "Fast heuristic estimate used for low-latency evaluation. Re-run detailed mode for deeper stability analysis.",
        "RiskLevel": risk
    }

def get_fast_career_estimate(resume_text):
    """Heuristic career progression estimate for fast evaluation mode (no LLM)."""
    text = (resume_text or "").lower()
    progression_terms = ["lead", "senior", "architect", "manager", "principal", "head"]
    hit_count = sum(1 for t in progression_terms if t in text)
    score = min(85, 45 + hit_count * 8)
    observations = [
        "Fast estimate from role/title signals in resume text.",
        "Use detailed mode for deeper promotion-path reasoning."
    ]
    if hit_count >= 3:
        observations.insert(0, "Strong leadership/seniority indicators detected.")
    elif hit_count >= 1:
        observations.insert(0, "Some progression indicators detected.")
    else:
        observations.insert(0, "Limited explicit progression indicators detected.")

    return {
        "progression_score": score,
        "key_observations": observations[:5],
        "career_path": [],
        "red_flags": ["Heuristic output - validate during interview screening."],
        "reasoning": "Fast heuristic estimate used to keep response latency low."
    }

def validate_progression_score(score):
    """Validate and normalize progression score"""
    try:
        if isinstance(score, str):
            score = score.strip('%')
        score = float(score)
        return int(max(0, min(100, score)))
    except (ValueError, TypeError):
        return 50

def validate_list(items):
    """Validate and clean list items"""
    if not isinstance(items, list):
        return []
    return [str(item).strip() for item in items if item and str(item).strip()]

def validate_career_path(path):
    """Validate and clean career path entries"""
    if not isinstance(path, list):
        return []
    
    cleaned_path = []
    required_fields = ["title", "company", "duration", "level", "progression"]
    
    for entry in path:
        if not isinstance(entry, dict):
            continue
        
        cleaned_entry = {}
        for field in required_fields:
            cleaned_entry[field] = str(entry.get(field, "Not specified")).strip()
        cleaned_path.append(cleaned_entry)
    
    return cleaned_path

def update_db_schema():
    """Update database schema if needed"""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # Add new columns if they don't exist
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN job_stability TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
        
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN career_progression TEXT;
        ''')
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Add oorwin_job_id column to evaluations table
    try:
        cursor.execute('''
            ALTER TABLE evaluations 
            ADD COLUMN oorwin_job_id TEXT;
        ''')
        logging.info("Added oorwin_job_id column to evaluations table")
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Create recruiter_handbooks table
    try:
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recruiter_handbooks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                oorwin_job_id TEXT,
                job_title TEXT,
                job_description TEXT,
                additional_context TEXT,
                markdown_content TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        logging.info("Created recruiter_handbooks table")
    except sqlite3.OperationalError:
        pass  # Table already exists
    
    # Create index on oorwin_job_id for faster queries
    try:
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_evaluations_job_id 
            ON evaluations(oorwin_job_id)
        ''')
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_handbooks_job_id 
            ON recruiter_handbooks(oorwin_job_id)
        ''')
        logging.info("Created indexes on oorwin_job_id")
    except sqlite3.OperationalError:
        pass  # Index already exists
    
    conn.commit()
    conn.close()

@login_required
def get_evaluation_details(evaluation_id):
    """API endpoint to get evaluation details by ID"""
    conn = None
    try:
        logging.info(f"Fetching evaluation details for ID: {evaluation_id}")
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Helper function for parsing JSON safely
        def parse_json_safely(json_str):
            if not json_str:
                logging.info("Empty JSON string, returning empty list")
                return []
            try:
                data = json.loads(json_str)
                if isinstance(data, list):
                    logging.info(f"Successfully parsed list with {len(data)} items")
                    return data
                elif isinstance(data, str):
                    try:
                        parsed_data = json.loads(data)
                        logging.info(f"Successfully parsed nested JSON string")
                        return parsed_data
                    except:
                        logging.info(f"Failed to parse nested JSON, treating as single item")
                        return [data]
                else:
                    logging.info(f"Non-list data type: {type(data)}, converting to string")
                    return [str(data)]
            except json.JSONDecodeError as e:
                logging.warning(f"JSON decode error: {str(e)}, attempting cleanup")
                try:
                    # Try to clean and parse the string
                    cleaned_str = json_str.strip('[]"\' ').replace('\\', '')
                    items = [item.strip('"\' ') for item in cleaned_str.split(',')]
                    result = [item for item in items if item]
                    logging.info(f"Cleanup successful, extracted {len(result)} items")
                    return result
                except Exception as e2:
                    logging.error(f"Cleanup failed: {str(e2)}")
                    return []
        
        # Get evaluation details first to get job title for default questions
        cursor.execute('''
            SELECT 
                e.id, 
                e.filename, 
                e.job_title, 
                e.match_percentage, 
                e.profile_summary, 
                e.job_stability,
                e.career_progression,
                e.timestamp,
                e.missing_keywords,
                e.behavioral_questions,
                e.technical_questions,
                e.nontechnical_questions
            FROM evaluations e
            WHERE e.id = ?
        ''', (evaluation_id,))
        
        row = cursor.fetchone()
        if not row:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        logging.info(f"Found evaluation with ID: {row[0]}, filename: {row[1]}")
        job_title = row[2]
        
        # Parse JSON fields
        try:
            job_stability = json.loads(row[5]) if row[5] else {}
            logging.info(f"Parsed job_stability: {type(job_stability)}")
        except Exception as e:
            logging.error(f"Error parsing job_stability: {str(e)}")
            job_stability = {}
            
        try:
            career_progression = json.loads(row[6]) if row[6] else {}
            logging.info(f"Parsed career_progression: {type(career_progression)}")
        except Exception as e:
            logging.error(f"Error parsing career_progression: {str(e)}")
            career_progression = {}
        
        # Parse missing keywords with special handling
        try:
            missing_keywords_raw = row[8]
            if missing_keywords_raw:
                try:
                    missing_keywords = json.loads(missing_keywords_raw)
                    logging.info(f"Parsed missing_keywords: {type(missing_keywords)}")
                    # If it's not a list, try to convert it
                    if not isinstance(missing_keywords, list):
                        if isinstance(missing_keywords, str):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords.strip('[]').split(',')]
                        else:
                            missing_keywords = [str(missing_keywords)]
                except Exception as e:
                    logging.error(f"Error parsing missing_keywords JSON: {str(e)}")
                    # If JSON parsing fails, try to extract from string
                    if isinstance(missing_keywords_raw, str):
                        # Check if it looks like a list
                        if missing_keywords_raw.startswith('[') and missing_keywords_raw.endswith(']'):
                            # Remove brackets and split by commas
                            missing_keywords = [k.strip(' "\'') for k in missing_keywords_raw.strip('[]').split(',')]
                        else:
                            missing_keywords = [missing_keywords_raw]
                    else:
                        missing_keywords = []
            else:
                missing_keywords = []
        except Exception as e:
            logging.error(f"Error processing missing_keywords: {str(e)}")
            missing_keywords = []
        
        # Initialize question variables
        technical_questions = []
        nontechnical_questions = []
        behavioral_questions = []
        
        # PRIORITY 1: Try to get interview questions from interview_questions table FIRST
        # This is the dedicated table for storing interview questions
        cursor.execute(
            "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
            (evaluation_id,)
        )
        iq_result = cursor.fetchone()
        
        if not iq_result:
            logging.info(f"No interview questions found with numeric ID, trying string ID")
            # If no results, try with the string representation of the ID
            cursor.execute(
                "SELECT technical_questions, nontechnical_questions, behavioral_questions FROM interview_questions WHERE evaluation_id = ?",
                (str(evaluation_id),)
            )
            iq_result = cursor.fetchone()
        
        if iq_result:
            logging.info(f"Found interview questions in interview_questions table for evaluation {evaluation_id}")
            technical_questions = parse_json_safely(iq_result[0])
            nontechnical_questions = parse_json_safely(iq_result[1])
            behavioral_questions = parse_json_safely(iq_result[2])
            logging.info(f"Retrieved from interview_questions table: {len(technical_questions)} technical, {len(nontechnical_questions)} non-technical, {len(behavioral_questions)} behavioral questions")
        else:
            logging.info(f"No interview questions found in interview_questions table for evaluation {evaluation_id}")
            
            # PRIORITY 2: Fallback to evaluations table if interview_questions table is empty
            logging.info("Falling back to evaluations table for questions")
        
        # Try to get behavioral questions from evaluations
        try:
            behavioral_questions_raw = row[9]
            if behavioral_questions_raw:
                behavioral_questions = parse_json_safely(behavioral_questions_raw)
                logging.info(f"Parsed behavioral_questions from evaluations: {len(behavioral_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing behavioral_questions from evaluations: {str(e)}")
        
        # Try to get technical questions from evaluations
        try:
            if row[10]:
                technical_questions = parse_json_safely(row[10])
                logging.info(f"Parsed technical_questions from evaluations: {len(technical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing technical_questions from evaluations: {str(e)}")
        
        # Try to get non-technical questions from evaluations
        try:
            if row[11]:
                nontechnical_questions = parse_json_safely(row[11])
                logging.info(f"Parsed nontechnical_questions from evaluations: {len(nontechnical_questions)} questions")
        except Exception as e:
            logging.error(f"Error parsing nontechnical_questions from evaluations: {str(e)}")
        
        # If still no behavioral questions, use default QUICK_CHECKS
        if not behavioral_questions:
            logging.info("No behavioral questions found, using QUICK_CHECKS")
            behavioral_questions = QUICK_CHECKS
        
        
        # Only generate default questions if we still don't have any questions at all
        if not technical_questions and not nontechnical_questions:
            logging.info(f"Generating default questions for job title: {job_title}")
            default_technical, default_nontechnical = get_default_interview_questions(job_title)
            
            if not technical_questions:
                technical_questions = default_technical
                logging.info(f"Using default technical questions: {len(technical_questions)} questions")
            
            if not nontechnical_questions:
                nontechnical_questions = default_nontechnical
                logging.info(f"Using default non-technical questions: {len(nontechnical_questions)} questions")
        
        # Create response
        response = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'match_percentage': row[3],
            'profile_summary': row[4] or "No summary available",
            'job_stability': job_stability,
            'career_progression': career_progression,
            'timestamp': row[7],
            'missing_keywords': missing_keywords,
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        }
        
        logging.info(f"Returning response with {len(technical_questions)} technical questions, {len(nontechnical_questions)} non-technical questions, {len(behavioral_questions)} behavioral questions")
        return jsonify(response)
    
    except Exception as e:
        logging.error(f"Error fetching evaluation details: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
    finally:
        if conn:
            conn.close()

@login_required
async def generate_questions_api(evaluation_id):
    """API endpoint to generate interview questions for an evaluation"""
    conn = None
    try:
        logging.info(f"Generating questions for evaluation ID: {evaluation_id}")
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get evaluation details
        cursor.execute(
            """
            SELECT resume_path, job_title, job_description, profile_summary 
            FROM evaluations 
            WHERE id = ?
            """,
            (evaluation_id,)
        )
        eval_result = cursor.fetchone()
        
        if not eval_result:
            logging.warning(f"No evaluation found with ID: {evaluation_id}")
            return jsonify({'error': 'Evaluation not found'}), 404
        
        # Extract resume text
        resume_path = eval_result[0]
        job_description = eval_result[2]
        profile_summary = eval_result[3]
        
        if not resume_path or resume_path == 'NULL' or resume_path == 'None' or str(resume_path).strip() == '':
            logging.error(f"Invalid resume_path for evaluation {evaluation_id}: {resume_path}")
            return jsonify({'error': 'No valid resume path found for this evaluation'}), 400
        
        resume_text = extract_text_from_file(resume_path)
        if not resume_text:
            return jsonify({'error': 'Failed to extract text from resume'}), 400
        
        # Generate questions
        logging.info(f"Generating questions for resume: {resume_path}")
        questions_data = await async_generate_questions(
            resume_text,
            job_description,
            profile_summary
        )
        
        technical_questions = questions_data.get("TechnicalQuestions", [])
        nontechnical_questions = questions_data.get("NonTechnicalQuestions", [])
        behavioral_questions = QUICK_CHECKS
        
        # Save questions to database
        try:
            # First check if there's an existing entry
            cursor.execute(
                "SELECT id FROM interview_questions WHERE evaluation_id = ?",
                (evaluation_id,)
            )
            existing = cursor.fetchone()
            
            if existing:
                # Update existing entry
                cursor.execute(
                    """
                    UPDATE interview_questions 
                    SET technical_questions = ?,
                        nontechnical_questions = ?,
                        behavioral_questions = ?
                    WHERE evaluation_id = ?
                    """,
                    (
                        json.dumps(technical_questions), 
                        json.dumps(nontechnical_questions), 
                        json.dumps(behavioral_questions), 
                        evaluation_id
                    )
                )
            else:
                # Insert new entry
                cursor.execute(
                    """
                    INSERT INTO interview_questions 
                    (evaluation_id, technical_questions, nontechnical_questions, behavioral_questions) 
                    VALUES (?, ?, ?, ?)
                    """,
                    (
                        evaluation_id,
                        json.dumps(technical_questions),
                        json.dumps(nontechnical_questions),
                        json.dumps(behavioral_questions)
                    )
                )
            
            conn.commit()
            logging.info(f"Saved questions for evaluation ID: {evaluation_id}")
        except Exception as e:
            logging.error(f"Error saving questions to database: {str(e)}")
            conn.rollback()
        
        # Return the generated questions
        return jsonify({
            'technical_questions': technical_questions,
            'nontechnical_questions': nontechnical_questions,
            'behavioral_questions': behavioral_questions
        })
        
    except Exception as e:
        logging.error(f"Error generating questions: {str(e)}")
        return jsonify({'error': str(e)}), 500
        
    finally:
        if conn:
            conn.close()

@limiter.limit("8 per minute")
@login_required
def generate_recruiter_handbook():
    """API endpoint to generate a comprehensive recruiter handbook"""
    try:
        data = request.get_json()
        (
            intake,
            errors,
            job_title,
            job_description,
            additional_context,
            oorwin_job_id,
            client_call_transcript,
        ) = _parse_handbook_generation_request(data)

        if errors:
            return jsonify({
                'success': False,
                'message': errors[0]
            }), 400
        
        # Check if handbook already exists for this job ID
        if oorwin_job_id:
            try:
                conn = sqlite3.connect(DATABASE_NAME)
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT id, markdown_content, job_title, timestamp, user_email
                    FROM recruiter_handbooks
                    WHERE oorwin_job_id = ?
                    ORDER BY timestamp DESC
                    LIMIT 1
                ''', (oorwin_job_id,))
                
                existing_handbook = cursor.fetchone()
                conn.close()
                
                if existing_handbook:
                    handbook_id, markdown_content, existing_job_title, timestamp, creator_email = existing_handbook
                    logging.info(f"Handbook already exists for JobID {oorwin_job_id} (ID: {handbook_id}, created: {timestamp})")
                    
                    # Get creator name if available
                    creator_name = None
                    if creator_email:
                        try:
                            conn = sqlite3.connect(DATABASE_NAME)
                            cursor = conn.cursor()
                            cursor.execute('SELECT name FROM users WHERE email = ?', (creator_email,))
                            user_row = cursor.fetchone()
                            if user_row:
                                creator_name = user_row[0] or creator_email
                            else:
                                creator_name = creator_email
                            conn.close()
                        except:
                            creator_name = creator_email
                    
                    return jsonify({
                        'success': True,
                        'existing': True,
                        'message': f'A handbook already exists for Job ID "{oorwin_job_id}". Showing existing handbook.',
                        'handbook_id': handbook_id,
                        'markdown_content': markdown_content,
                        'job_title': existing_job_title or job_title,
                        'oorwin_job_id': oorwin_job_id,
                        'created_at': timestamp,
                        'created_by': creator_name or creator_email or 'Unknown',
                        'created_by_email': creator_email
                    })
            except Exception as check_error:
                logging.error(f"Error checking for existing handbook: {str(check_error)}")
                # Continue with generation if check fails
            
        logging.info(f"Generating new recruiter handbook for JobID: {oorwin_job_id or 'None'}...")
        
        # Track start time for performance metrics
        import time
        handbook_start_time = time.time()
        
        handbook_prompt = _build_recruiter_handbook_prompt(
            job_description=job_description,
            additional_context=additional_context,
            intake=intake,
            client_call_transcript=client_call_transcript,
        )

        # Generate handbook using handbook-specific provider/model override
        try:
            handbook_provider, handbook_groq_model, handbook_groq_reasoning_effort = get_handbook_provider_settings()
            handbook_model_override = (
                handbook_groq_model if handbook_provider == "groq"
                else OPENAI_MODEL if handbook_provider == "openai"
                else None
            )
            response = generate_content_unified(
                handbook_prompt,
                provider_override=handbook_provider,
                model_override=handbook_model_override,
                groq_reasoning_effort_override=handbook_groq_reasoning_effort,
                max_completion_tokens_override=HANDBOOK_MAX_COMPLETION_TOKENS
            )
            logging.info(f"Response received: {type(response)}, has text attr: {hasattr(response, 'text')}")
            
            if not response:
                raise Exception("No response received from AI model")
            
            # Debug: Check what's in response.text
            if hasattr(response, 'text'):
                logging.info(f"response.text type: {type(response.text)}, value: '{response.text[:100] if response.text else 'None or empty'}'")
        
            if not hasattr(response, 'text'):
                logging.error(f"Response object has no 'text' attribute")
                logging.error(f"Response attributes: {dir(response)}")
                raise Exception(f"Response object missing 'text' attribute. Type: {type(response)}")
            
            if not response.text:
                logging.error(f"response.text is empty or None")
                logging.error(f"response.text value: {repr(response.text)}")
                # Try output_text as fallback
                if hasattr(response, 'output_text') and response.output_text:
                    logging.info("Using output_text as fallback")
                    handbook_content = response.output_text
                else:
                    raise Exception("AI response text is empty. response.text=" + repr(response.text))
            else:
                handbook_content = response.text
            
            logging.info(f"Handbook content length: {len(handbook_content)} characters")
        except Exception as api_error:
            logging.error(f"API call failed: {str(api_error)}", exc_info=True)
            raise
        
        logging.info("Recruiter handbook generated successfully")
        
        # Calculate time taken
        handbook_time_taken = round(time.time() - handbook_start_time, 2)  # Round to 2 decimal places
        logging.info(f"Handbook generation took {handbook_time_taken} seconds")
        
        # Save handbook to database
        try:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            
            # Get user email from session
            user_email = session.get('user', {}).get('email') if 'user' in session else None
            
            cursor.execute('''
                INSERT INTO recruiter_handbooks (
                    oorwin_job_id, job_title, job_description, 
                    additional_context, markdown_content, time_taken, user_email, timestamp,
                    intake_json
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                oorwin_job_id if oorwin_job_id else None,
                job_title,
                job_description,
                additional_context,
                handbook_content,
                handbook_time_taken,
                user_email,
                datetime.now(),
                intake_to_json(intake),
            ))
            
            conn.commit()
            handbook_id = cursor.lastrowid
            conn.close()
            
            logging.info(f"Handbook saved to database with ID: {handbook_id}")
            
            # Return success with handbook_id
            return jsonify({
                'success': True,
                'markdown_content': handbook_content,
                'handbook_id': handbook_id
            })
        except Exception as e:
            logging.error(f"Error saving handbook to database: {str(e)}")
            # Return without handbook_id if save fails
            return jsonify({
                'success': True,
                'markdown_content': handbook_content
            })
        
    except Exception as e:
        logging.error(f"Error generating recruiter handbook: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500


@login_required
def generate_recruiter_handbook_stream():
    """
    Streaming version of recruiter handbook generation.

    Emits Server-Sent Events of three kinds:
      - {"type": "meta",  "handbook_id": <int|null>, "existing": <bool>, ...}
      - {"type": "delta", "content": "<markdown chunk>"}
      - {"type": "done",  "handbook_id": <int>, "time_taken": <float>}
      - {"type": "error", "message": "..."}

    Drops perceived latency from 25–50s (full completion) to 1–2s (first token).
    """
    try:
        data = request.get_json() or {}
        (
            intake,
            errors,
            job_title,
            job_description,
            additional_context,
            oorwin_job_id,
            client_call_transcript,
        ) = _parse_handbook_generation_request(data)

        if errors:
            return jsonify({'success': False, 'message': errors[0]}), 400

        handbook_prompt = _build_recruiter_handbook_prompt(
            job_description=job_description,
            additional_context=additional_context,
            intake=intake,
            client_call_transcript=client_call_transcript,
        )

        # Capture session info BEFORE entering the generator (Flask session is
        # not safe to access from inside a streamed response on every server).
        user_email = session.get('user', {}).get('email') if 'user' in session else None

        # Cache hit: serve the existing handbook in one event and stop. The
        # frontend treats this identically to the old "existing" flag.
        existing_meta = None
        if oorwin_job_id:
            try:
                conn = sqlite3.connect(DATABASE_NAME)
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT id, markdown_content, job_title, timestamp, user_email
                    FROM recruiter_handbooks
                    WHERE oorwin_job_id = ?
                    ORDER BY timestamp DESC
                    LIMIT 1
                ''', (oorwin_job_id,))
                row = cursor.fetchone()
                conn.close()
                if row:
                    hb_id, md, ej_title, ts, creator_email = row
                    creator_name = creator_email
                    try:
                        conn = sqlite3.connect(DATABASE_NAME)
                        cursor = conn.cursor()
                        cursor.execute('SELECT name FROM users WHERE email = ?', (creator_email,))
                        u = cursor.fetchone()
                        conn.close()
                        if u and u[0]:
                            creator_name = u[0]
                    except Exception:
                        pass
                    existing_meta = {
                        'handbook_id': hb_id,
                        'markdown_content': md,
                        'job_title': ej_title or job_title,
                        'oorwin_job_id': oorwin_job_id,
                        'created_at': ts,
                        'created_by': creator_name or creator_email or 'Unknown',
                        'created_by_email': creator_email,
                    }
            except Exception as check_error:
                logging.error(f"Error checking for existing handbook (stream): {check_error}")

        def generate():
            start_ts = time.time()
            try:
                if existing_meta:
                    yield "data: " + json.dumps({
                        'type': 'existing',
                        **existing_meta,
                    }) + "\n\n"
                    yield "data: " + json.dumps({
                        'type': 'delta',
                        'content': existing_meta['markdown_content']
                    }) + "\n\n"
                    yield "data: " + json.dumps({
                        'type': 'done',
                        'handbook_id': existing_meta['handbook_id'],
                        'existing': True,
                        'time_taken': 0.0,
                    }) + "\n\n"
                    return

                yield "data: " + json.dumps({'type': 'meta', 'message': 'Starting handbook generation...'}) + "\n\n"

                handbook_provider, handbook_groq_model, handbook_groq_reasoning_effort = get_handbook_provider_settings()
                handbook_model_override = (
                    handbook_groq_model if handbook_provider == 'groq'
                    else OPENAI_MODEL if handbook_provider == 'openai'
                    else None
                )

                # Stream tokens from the LLM.
                stream = generate_content_unified(
                    handbook_prompt,
                    stream=True,
                    provider_override=handbook_provider,
                    model_override=handbook_model_override,
                    groq_reasoning_effort_override=handbook_groq_reasoning_effort,
                    max_completion_tokens_override=HANDBOOK_MAX_COMPLETION_TOKENS,
                )

                buffer_parts = []
                chunk_count = 0
                last_section_idx = -1

                def _yield_delta(text_chunk):
                    nonlocal chunk_count, last_section_idx
                    if not text_chunk:
                        return None, None
                    buffer_parts.append(text_chunk)
                    chunk_count += 1
                    section_msg = None
                    section_idx = handbook_section_index_from_markdown("".join(buffer_parts))
                    if section_idx > last_section_idx:
                        last_section_idx = section_idx
                        section_msg = "data: " + json.dumps(
                            {"type": "section", "index": section_idx}
                        ) + "\n\n"
                    delta_msg = "data: " + json.dumps(
                        {"type": "delta", "content": text_chunk}
                    ) + "\n\n"
                    return section_msg, delta_msg

                # generate_content_unified returns a generator for stream=True.
                # For Gemini that yields response objects with .text;
                # for OpenAI/Groq it yields raw content strings.
                for piece in stream:
                    text_chunk = getattr(piece, 'text', piece) if not isinstance(piece, str) else piece
                    section_msg, delta_msg = _yield_delta(text_chunk)
                    if section_msg:
                        yield section_msg
                    if delta_msg:
                        yield delta_msg

                handbook_content = ''.join(buffer_parts).strip()
                if not handbook_content:
                    raise RuntimeError("AI model returned empty handbook content")

                # Persist to DB AFTER the stream completes (cheap insert; user
                # already saw the markdown).
                handbook_id = None
                try:
                    conn = sqlite3.connect(DATABASE_NAME)
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO recruiter_handbooks (
                            oorwin_job_id, job_title, job_description,
                            additional_context, markdown_content, time_taken, user_email, timestamp,
                            intake_json
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        oorwin_job_id if oorwin_job_id else None,
                        job_title,
                        job_description,
                        additional_context,
                        handbook_content,
                        round(time.time() - start_ts, 2),
                        user_email,
                        datetime.now(),
                        intake_to_json(intake),
                    ))
                    conn.commit()
                    handbook_id = cursor.lastrowid
                    conn.close()
                except Exception as save_err:
                    logging.error(f"Error saving streamed handbook: {save_err}")

                yield "data: " + json.dumps({
                    'type': 'done',
                    'handbook_id': handbook_id,
                    'time_taken': round(time.time() - start_ts, 2),
                    'chunks': chunk_count,
                    'length': len(handbook_content),
                }) + "\n\n"

            except Exception as e:
                logging.error(f"Streaming handbook error: {e}", exc_info=True)
                yield "data: " + json.dumps({'type': 'error', 'message': str(e)[:500]}) + "\n\n"

        return Response(
            stream_with_context(generate()),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache, no-transform',
                'X-Accel-Buffering': 'no',  # disable nginx proxy buffering
            },
        )

    except Exception as e:
        logging.error(f"Error in generate_recruiter_handbook_stream: {e}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500


def _parse_handbook_generation_request(data):
    """Normalize intake payload and validate required fields."""
    data = data or {}
    intake = normalize_intake_payload(data)
    errors = validate_intake(intake)
    job_title = intake.get("job_title", "")
    job_description = intake.get("job_description", "")
    additional_context = intake.get("additional_context", "")
    oorwin_job_id = (data.get("oorwin_job_id") or intake.get("req_id") or "").strip()
    if oorwin_job_id:
        intake["req_id"] = oorwin_job_id

    client_call_transcript = None
    selected_transcript_path = (data.get("selected_transcript_path") or "").strip()
    if selected_transcript_path:
        try:
            from pluto.gcs_transcripts import (
                gcs_configured,
                get_transcript_by_path,
                validate_transcript_path,
            )

            if gcs_configured() and validate_transcript_path(selected_transcript_path):
                client_call_transcript = get_transcript_by_path(selected_transcript_path)
                intake["selected_transcript_path"] = selected_transcript_path
        except Exception as exc:
            logging.warning(
                "Could not load client call transcript %s: %s",
                selected_transcript_path,
                exc,
            )

    return (
        intake,
        errors,
        job_title,
        job_description,
        additional_context,
        oorwin_job_id,
        client_call_transcript,
    )


def _build_recruiter_handbook_prompt(
    job_description: str,
    additional_context: str = "",
    intake: dict | None = None,
    client_call_transcript: dict | None = None,
) -> str:
    """Backward-compatible wrapper — see pluto.handbook_intake."""
    return build_recruiter_handbook_prompt(
        job_description=job_description,
        additional_context=additional_context,
        intake=intake,
        client_call_transcript=client_call_transcript,
    )


@login_required
def generate_handbook_summary():
    """Generate concise AI summary for handbook context."""
    try:
        data = request.get_json() or {}
        job_title = (data.get('job_title') or '').strip()
        job_description = (data.get('job_description') or '').strip()
        handbook_content = (data.get('handbook_content') or '').strip()

        if not handbook_content and not job_description:
            return jsonify({'success': False, 'message': 'Missing handbook content'}), 400

        summary_prompt = f"""You are an expert technical recruiting strategist.
Write exactly 4-5 complete sentences summarizing this job for recruiters.

Requirements:
- Plain text only
- No markdown, no bullets, no table syntax, no headings
- Keep all sentences complete and concise
- Include: role objective, must-have themes, sourcing direction, one screening caution

Job Title: {job_title}
Job Description:
{job_description}

Handbook Content:
{handbook_content[:7000]}
"""

        handbook_provider, handbook_groq_model, handbook_groq_reasoning_effort = get_handbook_provider_settings()
        handbook_model_override = (
            handbook_groq_model if handbook_provider == "groq"
            else OPENAI_MODEL if handbook_provider == "openai"
            else None
        )
        response = generate_content_unified(
            summary_prompt,
            provider_override=handbook_provider,
            model_override=handbook_model_override,
            groq_reasoning_effort_override=handbook_groq_reasoning_effort,
            max_completion_tokens_override=min(HANDBOOK_MAX_COMPLETION_TOKENS, 1024)
        )

        summary = (response.text or '').strip()
        summary = re.sub(r'^\s*#+\s*', '', summary)
        summary = re.sub(r'^\s*[-*\d.]+\s*', '', summary)
        summary = re.sub(r'\s+', ' ', summary).strip()
        if not summary:
            return jsonify({'success': False, 'message': 'Empty summary from model'}), 500

        return jsonify({'success': True, 'summary': summary})
    except Exception as e:
        logging.error(f"Error generating handbook summary: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


def _parse_jd_quality_json_from_model(raw_text):
    """Extract a JSON object from model output (handles fences and trailing prose)."""
    text = (raw_text or '').strip()
    if not text:
        return None
    if text.startswith('```'):
        text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s*```\s*$', '', text)
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    m = re.search(r'\{[\s\S]*\}', text)
    if m:
        try:
            return json.loads(m.group(0))
        except json.JSONDecodeError:
            return None
    return None


@limiter.limit("20 per minute")
@login_required
def jd_quality_score():
    """
    JD Quality & Ghost Role Scorer — runs before handbook generation.
    Returns score 0-100 plus dimension notes; if score < 70, include 3 improvement suggestions.
    """
    try:
        data = request.get_json() or {}
        job_title = (data.get('job_title') or '').strip()
        job_description = (data.get('job_description') or '').strip()

        if len(job_description) < 50:
            return jsonify({
                'success': False,
                'message': 'Job description is too short to score meaningfully (minimum ~50 characters).'
            }), 400

        jd_excerpt = job_description[:12000]
        scorer_prompt = f"""You are an expert talent-acquisition analyst and labor-market realist. Your job is to evaluate a job description (JD) for recruiting quality, realism, and inclusive language before sourcing begins.

Evaluate these dimensions:

1) **Completeness & Clarity** — Are responsibilities, must-haves, and success signals concrete (measurable skills, scope, stack, outcomes)? Flag vague filler, buzzwords without evidence, missing level/scope, or contradictory asks.

2) **Realism & "Ghost role" risk** — Do years-of-experience and seniority match the title and tech stack? Call out impossible combinations (e.g., requiring more years on a framework/library than it has existed in stable form), title vs IC expectations, or JDs that read like a unicorn wishlist unlikely to exist in one hire.

3) **Bias & exclusion** — Flag coded, exclusionary, or discriminatory language (age, gender, nationality, "native speaker", "digital native", "young and energetic", "culture fit" as a proxy, etc.). Note severity briefly.

**Overall score (0-100):** Weight clarity ~35%, realism ~40%, bias/inclusion ~25%. Be opinionated: mediocre JDs should land roughly 55–72; strong JDs 80+; seriously flawed or ghost-role JDs can be below 45.

Respond with **ONLY** a single JSON object (no markdown fences, no commentary before or after). Use exactly this shape:
{{
  "score": <integer 0-100>,
  "completeness_summary": "<one sentence>",
  "realism_summary": "<one sentence>",
  "bias_summary": "<one sentence or 'None significant' if clean>",
  "improvement_suggestions": ["<specific edit 1>", "<specific edit 2>", "<specific edit 3>"]
}}

The three strings in improvement_suggestions must be **actionable** (what to add, remove, or rephrase). If the score would be 70 or above, still provide three small polish suggestions (not warnings).

Job title (may be empty): {job_title}

Job description:
{jd_excerpt}
"""

        handbook_provider, handbook_groq_model, handbook_groq_reasoning_effort = get_handbook_provider_settings()
        handbook_model_override = (
            handbook_groq_model if handbook_provider == 'groq'
            else OPENAI_MODEL if handbook_provider == 'openai'
            else None
        )
        response = generate_content_unified(
            scorer_prompt,
            provider_override=handbook_provider,
            model_override=handbook_model_override,
            groq_reasoning_effort_override=handbook_groq_reasoning_effort,
            max_completion_tokens_override=min(HANDBOOK_MAX_COMPLETION_TOKENS, 1536)
        )

        raw = (response.text or '').strip()
        parsed = _parse_jd_quality_json_from_model(raw)
        if not isinstance(parsed, dict):
            logging.warning('JD quality scorer: failed to parse JSON from model output')
            return jsonify({
                'success': False,
                'message': 'Could not parse scorer response. Try again or shorten the JD.'
            }), 502

        score = parsed.get('score')
        try:
            score = int(score)
        except (TypeError, ValueError):
            score = 65
        score = max(0, min(100, score))

        suggestions = parsed.get('improvement_suggestions') or []
        if not isinstance(suggestions, list):
            suggestions = []
        suggestions = [str(s).strip() for s in suggestions if str(s).strip()]
        while len(suggestions) < 3:
            suggestions.append('Add concrete outcomes for the first 90 days and separate must-haves from nice-to-haves.')
        suggestions = suggestions[:3]

        return jsonify({
            'success': True,
            'score': score,
            'completeness_summary': (parsed.get('completeness_summary') or '').strip(),
            'realism_summary': (parsed.get('realism_summary') or '').strip(),
            'bias_summary': (parsed.get('bias_summary') or '').strip(),
            'improvement_suggestions': suggestions,
            'show_warning': score < 70
        })
    except Exception as e:
        logging.error(f"Error in jd_quality_score: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500


@login_required
def get_job_ids():
    """API endpoint to get all unique JobIDs for auto-suggest"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get unique JobIDs from evaluations
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM evaluations 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        eval_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Get unique JobIDs from handbooks
        cursor.execute('''
            SELECT DISTINCT oorwin_job_id 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id IS NOT NULL AND oorwin_job_id != ''
            ORDER BY oorwin_job_id
        ''')
        handbook_job_ids = [row[0] for row in cursor.fetchall()]
        
        # Combine and remove duplicates
        all_job_ids = list(set(eval_job_ids + handbook_job_ids))
        all_job_ids.sort()
        
        conn.close()
        
        return jsonify({
            'success': True,
            'job_ids': all_job_ids
        })
        
    except Exception as e:
        logging.error(f"Error fetching JobIDs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_job_data(job_id):
    """API endpoint to get job description for auto-fill based on JobID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # First try to get from handbooks (they always have full JD)
        cursor.execute('''
            SELECT job_title, job_description, additional_context, intake_json 
            FROM recruiter_handbooks 
            WHERE oorwin_job_id = ? 
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        
        if result and result[1]:  # Check if job_description is not empty
            conn.close()
            intake = intake_from_json(result[3]) if len(result) > 3 else None
            payload = {
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'additional_context': result[2] if result[2] else '',
                'source': 'handbook',
            }
            if intake:
                payload['intake'] = intake
            return jsonify(payload)
        
        # If not found in handbooks or JD is empty, try evaluations
        cursor.execute('''
            SELECT job_title, job_description 
            FROM evaluations 
            WHERE oorwin_job_id = ? AND job_description IS NOT NULL AND job_description != ''
            ORDER BY timestamp DESC 
            LIMIT 1
        ''', (job_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return jsonify({
                'success': True,
                'job_title': result[0],
                'job_description': result[1],
                'source': 'evaluation'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'No data found for this JobID'
            }), 404
        
    except Exception as e:
        logging.error(f"Error fetching job data for {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_handbooks():
    """API endpoint to get all recruiter handbooks (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'handbooks': []
            })
        
        # Build WHERE clause for filtering
        placeholders = ','.join(['?'] * len(accessible_emails))
        where_clause = f"WHERE user_email IN ({placeholders})"
        
        cursor.execute(f'''
            SELECT id, oorwin_job_id, job_title, job_description, 
                   additional_context, markdown_content, timestamp
            FROM recruiter_handbooks
            {where_clause}
            ORDER BY timestamp DESC
        ''', accessible_emails)
        
        rows = cursor.fetchall()
        conn.close()
        
        handbooks = []
        for row in rows:
            handbooks.append({
                'id': row[0],
                'oorwin_job_id': row[1],
                'job_title': row[2],
                'job_description': row[3],
                'additional_context': row[4],
                'markdown_content': row[5],
                'timestamp': row[6]
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

# ===== DASHBOARD & ANALYTICS ROUTES =====

@login_required
def get_analytics_overview():
    """API endpoint for dashboard overview metrics with user, team, and date filtering"""
    try:
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        # Get accessible user emails based on role
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"Analytics overview - Current user: {user_email}, Accessible emails: {accessible_emails}")
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            # Intersect with accessible emails
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Ensure we always include the current user's email if accessible_emails is empty
        if not accessible_emails:
            accessible_emails = [user_email]
            logging.info(f"Analytics overview - No accessible emails, using current user: {user_email}")
        
        logging.info(f"Analytics overview - Using accessible emails: {accessible_emails}")
        
        # Build WHERE clauses first, then build params to match
        eval_where_parts = []
        handbook_where_parts = []
        
        # Build user_email filter
        if filter_user and (filter_user in accessible_emails or filter_user == user_email):
            # When filtering by specific user, use exact match (don't include NULL records)
            eval_where_parts.append("user_email = ?")
            handbook_where_parts.append("user_email = ?")
        else:
            # Include accessible emails and NULL records
            placeholders = ','.join(['?'] * len(accessible_emails))
            eval_where_parts.append(f"(user_email IN ({placeholders}) OR user_email IS NULL)")
            handbook_where_parts.append(f"(user_email IN ({placeholders}) OR user_email IS NULL)")
        
        # Add date filters
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts) if eval_where_parts else ""
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts) if handbook_where_parts else ""
        
        logging.info(f"Analytics overview - Eval WHERE clause: {eval_where_clause}")
        logging.info(f"Analytics overview - Handbook WHERE clause: {handbook_where_clause}")
        
        # Build params to match WHERE clause placeholders (in order)
        eval_params = []
        handbook_params = []
        
        # Add user_email params (either filter_user or accessible_emails)
        if filter_user and (filter_user in accessible_emails or filter_user == user_email):
            eval_params.append(filter_user)
            handbook_params.append(filter_user)
        else:
            eval_params.extend(accessible_emails)
            handbook_params.extend(accessible_emails)
        
        # Add date params
        if date_from:
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        logging.info(f"Analytics overview - Eval params: {eval_params}")
        logging.info(f"Analytics overview - Handbook params: {handbook_params}")
        
        core = fetch_overview_core_metrics(
            cursor,
            eval_where_clause,
            eval_params,
            handbook_where_clause,
            handbook_params,
        )
        total_evaluations = core["total_evaluations"]
        total_handbooks = core["total_handbooks"]
        total_jobs = max(core["unique_jobs_evals"], core["unique_jobs_handbooks"])
        avg_match_score = core["avg_match_score"]
        
        # Calculate previous period for comparison (if date range is specified)
        prev_total_evaluations = None
        prev_total_handbooks = None
        prev_total_jobs = None
        prev_avg_match_score = None
        
        if date_from and date_to:
            try:
                from datetime import datetime, timedelta
                # Calculate previous period (same duration before the selected period)
                from_date_obj = datetime.strptime(date_from, '%Y-%m-%d')
                to_date_obj = datetime.strptime(date_to, '%Y-%m-%d')
                period_days = (to_date_obj - from_date_obj).days + 1
                
                prev_to_date = from_date_obj - timedelta(days=1)
                prev_from_date = prev_to_date - timedelta(days=period_days - 1)
                
                prev_date_from = prev_from_date.strftime('%Y-%m-%d')
                prev_date_to = prev_to_date.strftime('%Y-%m-%d')
                
                # Build previous period WHERE clauses
                prev_eval_where_parts = []
                prev_handbook_where_parts = []
                prev_eval_params = []
                prev_handbook_params = []
                
                if accessible_emails:
                    placeholders = ','.join(['?'] * len(accessible_emails))
                    prev_eval_where_parts.append(f"user_email IN ({placeholders})")
                    prev_handbook_where_parts.append(f"user_email IN ({placeholders})")
                    prev_eval_params.extend(accessible_emails)
                    prev_handbook_params.extend(accessible_emails)
                
                if filter_user and filter_user in accessible_emails:
                    prev_eval_where_parts.append("user_email = ?")
                    prev_handbook_where_parts.append("user_email = ?")
                    prev_eval_params.append(filter_user)
                    prev_handbook_params.append(filter_user)
                
                prev_eval_where_parts.append("DATE(timestamp) >= ?")
                prev_eval_where_parts.append("DATE(timestamp) <= ?")
                prev_handbook_where_parts.append("DATE(timestamp) >= ?")
                prev_handbook_where_parts.append("DATE(timestamp) <= ?")
                prev_eval_params.extend([prev_date_from, prev_date_to])
                prev_handbook_params.extend([prev_date_from, prev_date_to])
                
                prev_eval_where_clause = " WHERE " + " AND ".join(prev_eval_where_parts) if prev_eval_where_parts else ""
                prev_handbook_where_clause = " WHERE " + " AND ".join(prev_handbook_where_parts) if prev_handbook_where_parts else ""
                
                cursor.execute(f'SELECT COUNT(*) FROM evaluations{prev_eval_where_clause}', prev_eval_params)
                prev_total_evaluations = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(*) FROM recruiter_handbooks{prev_handbook_where_clause}', prev_handbook_params)
                prev_total_handbooks = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM evaluations{prev_eval_where_clause} AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""', prev_eval_params)
                prev_unique_jobs_evals = cursor.fetchone()[0]
                
                cursor.execute(f'SELECT COUNT(DISTINCT oorwin_job_id) FROM recruiter_handbooks{prev_handbook_where_clause} AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""', prev_handbook_params)
                prev_unique_jobs_handbooks = cursor.fetchone()[0]
                
                prev_total_jobs = max(prev_unique_jobs_evals, prev_unique_jobs_handbooks)
                
                cursor.execute(f'SELECT AVG(match_percentage) FROM evaluations{prev_eval_where_clause} AND match_percentage IS NOT NULL', prev_eval_params)
                prev_avg_match_score = cursor.fetchone()[0] or 0
                prev_avg_match_score = round(prev_avg_match_score, 1)
            except Exception as e:
                logging.error(f"Error calculating previous period: {str(e)}")
        
        # Get conversion rate (jobs with handbook that also have evaluations) with filters
        # Build join query with filters
        join_where_parts = []
        join_params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            join_where_parts.append(f"h.user_email IN ({placeholders})")
            join_where_parts.append(f"e.user_email IN ({placeholders})")
            join_params.extend(accessible_emails)
            join_params.extend(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            join_where_parts.append("h.user_email = ?")
            join_where_parts.append("e.user_email = ?")
            join_params.append(filter_user)
            join_params.append(filter_user)
        
        if date_from:
            join_where_parts.append("DATE(h.timestamp) >= ?")
            join_where_parts.append("DATE(e.timestamp) >= ?")
            join_params.append(date_from)
            join_params.append(date_from)
        
        if date_to:
            join_where_parts.append("DATE(h.timestamp) <= ?")
            join_where_parts.append("DATE(e.timestamp) <= ?")
            join_params.append(date_to)
            join_params.append(date_to)
        
        join_where_clause = " WHERE " + " AND ".join(join_where_parts) if join_where_parts else ""
        
        cursor.execute(f'''
            SELECT COUNT(DISTINCT h.oorwin_job_id) 
            FROM recruiter_handbooks h
            INNER JOIN evaluations e ON h.oorwin_job_id = e.oorwin_job_id
            {join_where_clause}
            AND h.oorwin_job_id IS NOT NULL AND h.oorwin_job_id != ""
        ''', join_params)
        jobs_with_both = cursor.fetchone()[0]
        conversion_rate = (jobs_with_both / total_handbooks * 100) if total_handbooks > 0 else 0
        
        # Get average evaluations per job with filters
        cursor.execute(f'''
            SELECT AVG(eval_count) FROM (
                SELECT oorwin_job_id, COUNT(*) as eval_count 
                FROM evaluations 
                {eval_where_clause}
                AND oorwin_job_id IS NOT NULL AND oorwin_job_id != ""
                GROUP BY oorwin_job_id
            )
        ''', eval_params)
        avg_evals_per_job = cursor.fetchone()[0] or 0
        
        # Get active jobs (activity in last 7 days) with filters
        active_where_parts = list(eval_where_parts)
        active_params = list(eval_params)
        active_where_parts.append("datetime(timestamp) >= datetime('now', '-7 days')")
        
        active_where_clause = " WHERE " + " AND ".join(active_where_parts) if active_where_parts else ""
        cursor.execute(f'''
            SELECT COUNT(DISTINCT oorwin_job_id) 
            FROM evaluations 
            {active_where_clause}
            AND oorwin_job_id IS NOT NULL 
            AND oorwin_job_id != ""
        ''', active_params)
        active_jobs_7d = cursor.fetchone()[0]
        
        # Get average eval time (calculate from time_taken column)
        eval_time_where = eval_where_clause + (" AND " if eval_where_clause else " WHERE ") + "time_taken IS NOT NULL AND time_taken > 0"
        cursor.execute(f'SELECT AVG(time_taken) FROM evaluations{eval_time_where}', eval_params)
        avg_eval_time_result = cursor.fetchone()[0]
        avg_eval_time = round(avg_eval_time_result, 1) if avg_eval_time_result else 0
        
        # Get average handbook generation time
        handbook_time_where = handbook_where_clause + (" AND " if handbook_where_clause else " WHERE ") + "time_taken IS NOT NULL AND time_taken > 0"
        cursor.execute(f'SELECT AVG(time_taken) FROM recruiter_handbooks{handbook_time_where}', handbook_params)
        avg_handbook_time_result = cursor.fetchone()[0]
        avg_handbook_time = round(avg_handbook_time_result, 1) if avg_handbook_time_result else 0
        
        conn.close()
        
        # Calculate trends
        def calculate_trend(current, previous):
            if previous is None or previous == 0:
                return None, None
            change = current - previous
            percent_change = (change / previous) * 100 if previous > 0 else 0
            return change, round(percent_change, 1)
        
        eval_change, eval_percent = calculate_trend(total_evaluations, prev_total_evaluations)
        handbook_change, handbook_percent = calculate_trend(total_handbooks, prev_total_handbooks)
        jobs_change, jobs_percent = calculate_trend(total_jobs, prev_total_jobs)
        score_change, score_percent = calculate_trend(avg_match_score, prev_avg_match_score)
        
        return jsonify({
            'success': True,
            'metrics': {
                'total_evaluations': total_evaluations,
                'total_handbooks': total_handbooks,
                'total_jobs': total_jobs,
                'active_jobs': active_jobs_7d,
                'avg_match_score': round(avg_match_score, 1),
                'conversion_rate': round(conversion_rate, 1),
                'avg_evals_per_job': round(avg_evals_per_job, 1),
                'avg_eval_time': avg_eval_time,
                'avg_handbook_time': avg_handbook_time,
                'trends': {
                    'evaluations': {
                        'change': eval_change,
                        'percent': eval_percent
                    },
                    'handbooks': {
                        'change': handbook_change,
                        'percent': handbook_percent
                    },
                    'jobs': {
                        'change': jobs_change,
                        'percent': jobs_percent
                    },
                    'match_score': {
                        'change': score_change,
                        'percent': score_percent
                    }
                }
            }
        })
        
    except Exception as e:
        logging.error(f"Error fetching analytics overview: {str(e)}", exc_info=True)
        import traceback
        logging.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'message': str(e),
            'error': traceback.format_exc()
        }), 500

@login_required
def export_analytics_csv():
    """Export team/user data to CSV with evaluations and handbooks"""
    try:
        import csv
        from io import StringIO
        
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        # Get accessible user emails based on role
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"CSV Export - User: {user_email}, Accessible emails: {len(accessible_emails) if accessible_emails else 0}")
        logging.info(f"CSV Export - Filters: team={filter_team}, user={filter_user}, date_from={date_from}, date_to={date_to}")
        
        # Filter by team if specified
        if filter_team:
            conn_temp = sqlite3.connect(DATABASE_NAME)
            cursor_temp = conn_temp.cursor()
            cursor_temp.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor_temp.fetchall()]
            conn_temp.close()
            logging.info(f"CSV Export - Team emails: {len(team_emails)}")
            if accessible_emails:
                accessible_emails = [e for e in accessible_emails if e in team_emails]
                logging.info(f"CSV Export - After team filter: {len(accessible_emails)} accessible emails")
            else:
                # If no accessible users but team filter specified, use team emails
                accessible_emails = team_emails
                logging.info(f"CSV Export - Using team emails directly: {len(accessible_emails)}")
        
        # Filter by specific user if selected
        if filter_user:
            if accessible_emails and filter_user in accessible_emails:
                accessible_emails = [filter_user]
                logging.info(f"CSV Export - Filtered to specific user: {filter_user}")
            elif not accessible_emails:
                # If no accessible users restriction, allow if user exists
                conn_temp = sqlite3.connect(DATABASE_NAME)
                cursor_temp = conn_temp.cursor()
                cursor_temp.execute('SELECT email FROM users WHERE email = ?', (filter_user,))
                if cursor_temp.fetchone():
                    accessible_emails = [filter_user]
                    logging.info(f"CSV Export - Using filter_user directly: {filter_user}")
                conn_temp.close()
        
        logging.info(f"CSV Export - Final accessible emails count: {len(accessible_emails) if accessible_emails else 0}")
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get user info for all accessible users (or all users if no restriction)
        if accessible_emails:
            user_placeholders = ','.join(['?'] * len(accessible_emails))
            cursor.execute(f'''
                SELECT email, name, role, team FROM users 
                WHERE email IN ({user_placeholders})
            ''', accessible_emails)
        else:
            # If no accessible users restriction, get all users (for Admin)
            cursor.execute('SELECT email, name, role, team FROM users')
        users_dict = {row[0]: {'name': row[1], 'role': row[2], 'team': row[3]} for row in cursor.fetchall()}
        
        # Build WHERE clause for evaluations and handbooks
        where_parts = []
        params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            where_parts.append(f"user_email IN ({placeholders})")
            params.extend(accessible_emails)
        else:
            # If no accessible users, still allow CSV export but with no data
            # This handles cases where user has no accessible data (e.g., new Recruiter)
            where_parts.append("1=0")  # Always false condition - returns no rows
        
        if date_from:
            where_parts.append("DATE(timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts) if where_parts else ""
        
        # Get evaluations
        cursor.execute(f'''
            SELECT user_email, filename, job_title, oorwin_job_id, 
                   match_percentage, timestamp
            FROM evaluations
            {where_clause}
            ORDER BY timestamp DESC, user_email
        ''', params)
        evaluations = cursor.fetchall()
        
        # Get handbooks
        cursor.execute(f'''
            SELECT user_email, job_title, oorwin_job_id, timestamp
            FROM recruiter_handbooks
            {where_clause}
            ORDER BY timestamp DESC, user_email
        ''', params)
        handbooks = cursor.fetchall()
        
        conn.close()
        
        # Create CSV content
        output = StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow([
            'User Email', 'User Name', 'Role', 'Team', 'Date', 
            'Type', 'Job Title', 'Job ID', 'Match Score', 'Filename'
        ])
        
        # Write evaluation rows
        for eval_row in evaluations:
            user_email_eval, filename, job_title, job_id, match_score, timestamp = eval_row
            user_info = users_dict.get(user_email_eval, {'name': '', 'role': '', 'team': ''})
            writer.writerow([
                user_email_eval,
                user_info['name'],
                user_info['role'],
                user_info['team'],
                timestamp,
                'Evaluation',
                job_title or '',
                job_id or '',
                match_score or '',
                filename or ''
            ])
        
        # Write handbook rows
        for hb_row in handbooks:
            user_email_hb, job_title, job_id, timestamp = hb_row
            user_info = users_dict.get(user_email_hb, {'name': '', 'role': '', 'team': ''})
            writer.writerow([
                user_email_hb,
                user_info['name'],
                user_info['role'],
                user_info['team'],
                timestamp,
                'Handbook',
                job_title or '',
                job_id or '',
                '',
                ''
            ])
        
        # Create response
        output.seek(0)
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'text/csv'
        response.headers['Content-Disposition'] = f'attachment; filename=team_data_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        
        return response
        
    except Exception as e:
        logging.error(f"Error exporting CSV: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_analytics_timeline():
    """API endpoint for activity timeline chart data"""
    try:
        days = int(request.args.get('days', 30))
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get daily evaluation counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM evaluations
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        eval_data = cursor.fetchall()
        
        # Get daily handbook counts
        cursor.execute(f'''
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM recruiter_handbooks
            WHERE datetime(timestamp) >= datetime('now', '-{days} days')
            GROUP BY DATE(timestamp)
            ORDER BY date
        ''')
        handbook_data = cursor.fetchall()
        
        conn.close()
        
        # Convert to dict for easy merging
        eval_dict = {row[0]: row[1] for row in eval_data}
        handbook_dict = {row[0]: row[1] for row in handbook_data}
        
        # Merge and create timeline
        from datetime import datetime, timedelta
        timeline = []
        start_date = datetime.now() - timedelta(days=days)
        
        for i in range(days + 1):
            current_date = (start_date + timedelta(days=i)).strftime('%Y-%m-%d')
            timeline.append({
                'date': current_date,
                'evaluations': eval_dict.get(current_date, 0),
                'handbooks': handbook_dict.get(current_date, 0)
            })
        
        return jsonify({
            'success': True,
            'timeline': timeline
        })
        
    except Exception as e:
        logging.error(f"Error fetching timeline data: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_team_performance():
    """API endpoint for team performance comparison"""
    try:
        user_email = session['user'].get('email')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get all teams from accessible users
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            cursor.execute(f'''
                SELECT DISTINCT team FROM users 
                WHERE email IN ({placeholders}) AND team IS NOT NULL AND team != ""
            ''', accessible_emails)
            teams = [row[0] for row in cursor.fetchall()]
        else:
            teams = []
        
        team_performance = []
        for team in teams:
            # Get team members
            cursor.execute('SELECT email FROM users WHERE team = ?', (team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            
            if not team_emails:
                continue
            
            # Build WHERE clause
            team_placeholders = ','.join(['?'] * len(team_emails))
            eval_where = [f"user_email IN ({team_placeholders})"]
            handbook_where = [f"user_email IN ({team_placeholders})"]
            eval_params = list(team_emails)
            handbook_params = list(team_emails)
            
            if date_from:
                eval_where.append("DATE(timestamp) >= ?")
                handbook_where.append("DATE(timestamp) >= ?")
                eval_params.append(date_from)
                handbook_params.append(date_from)
            
            if date_to:
                eval_where.append("DATE(timestamp) <= ?")
                handbook_where.append("DATE(timestamp) <= ?")
                eval_params.append(date_to)
                handbook_params.append(date_to)
            
            eval_where_clause = " WHERE " + " AND ".join(eval_where)
            handbook_where_clause = " WHERE " + " AND ".join(handbook_where)
            
            # Get counts
            cursor.execute(f'SELECT COUNT(*) FROM evaluations{eval_where_clause}', eval_params)
            eval_count = cursor.fetchone()[0]
            
            cursor.execute(f'SELECT COUNT(*) FROM recruiter_handbooks{handbook_where_clause}', handbook_params)
            handbook_count = cursor.fetchone()[0]
            
            cursor.execute(f'SELECT AVG(match_percentage) FROM evaluations{eval_where_clause} AND match_percentage IS NOT NULL', eval_params)
            avg_score = cursor.fetchone()[0] or 0
            
            team_performance.append({
                'team': team,
                'evaluations': eval_count,
                'handbooks': handbook_count,
                'avg_score': round(avg_score, 1)
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'teams': team_performance
        })
        
    except Exception as e:
        logging.error(f"Error fetching team performance: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_match_score_distribution():
    """API endpoint for match score distribution"""
    try:
        user_email = session['user'].get('email')
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Build WHERE clause
        where_parts = []
        params = []
        
        if accessible_emails:
            placeholders = ','.join(['?'] * len(accessible_emails))
            where_parts.append(f"user_email IN ({placeholders})")
            params.extend(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            where_parts.append("user_email = ?")
            params.append(filter_user)
        
        if date_from:
            where_parts.append("DATE(timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts) if where_parts else ""
        
        # Get score distribution (buckets: 0-20, 21-40, 41-60, 61-80, 81-100)
        cursor.execute(f'''
            SELECT 
                CASE 
                    WHEN match_percentage <= 20 THEN '0-20'
                    WHEN match_percentage <= 40 THEN '21-40'
                    WHEN match_percentage <= 60 THEN '41-60'
                    WHEN match_percentage <= 80 THEN '61-80'
                    ELSE '81-100'
                END as score_range,
                COUNT(*) as count
            FROM evaluations
            {where_clause}
            AND match_percentage IS NOT NULL
            GROUP BY score_range
            ORDER BY score_range
        ''', params)
        
        distribution = cursor.fetchall()
        conn.close()
        
        # Format for chart
        score_ranges = ['0-20', '21-40', '41-60', '61-80', '81-100']
        distribution_dict = {row[0]: row[1] for row in distribution}
        
        result = [{'range': r, 'count': distribution_dict.get(r, 0)} for r in score_ranges]
        
        return jsonify({
            'success': True,
            'distribution': result
        })
        
    except Exception as e:
        logging.error(f"Error fetching match score distribution: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_user_activity():
    """API endpoint for user activity breakdown"""
    try:
        user_email = session['user'].get('email')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'users': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Get user names
        placeholders = ','.join(['?'] * len(accessible_emails))
        cursor.execute(f'''
            SELECT email, name FROM users
            WHERE email IN ({placeholders})
        ''', accessible_emails)
        users_dict = {row[0]: row[1] or row[0] for row in cursor.fetchall()}
        conn.close()
        
        # Get activity counts
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        eval_where_parts = [f"user_email IN ({placeholders})"]
        handbook_where_parts = [f"user_email IN ({placeholders})"]
        eval_params = list(accessible_emails)
        handbook_params = list(accessible_emails)
        
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts)
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts)
        
        # Get evaluation counts per user
        cursor.execute(f'''
            SELECT user_email, COUNT(*) as count
            FROM evaluations
            {eval_where_clause}
            GROUP BY user_email
        ''', eval_params)
        eval_counts = {row[0]: row[1] for row in cursor.fetchall()}
        
        # Get handbook counts per user
        cursor.execute(f'''
            SELECT user_email, COUNT(*) as count
            FROM recruiter_handbooks
            {handbook_where_clause}
            GROUP BY user_email
        ''', handbook_params)
        handbook_counts = {row[0]: row[1] for row in cursor.fetchall()}
        
        conn.close()
        
        # Combine data
        user_activity = []
        for email in accessible_emails:
            user_activity.append({
                'email': email,
                'name': users_dict.get(email, email),
                'evaluations': eval_counts.get(email, 0),
                'handbooks': handbook_counts.get(email, 0),
                'total': eval_counts.get(email, 0) + handbook_counts.get(email, 0)
            })
        
        # Sort by total activity
        user_activity.sort(key=lambda x: x['total'], reverse=True)
        
        return jsonify({
            'success': True,
            'users': user_activity
        })
        
    except Exception as e:
        logging.error(f"Error fetching user activity: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_recent_activity():
    """API endpoint for recent activity feed"""
    try:
        user_email = session['user'].get('email')
        limit = int(request.args.get('limit', 20))
        filter_team = request.args.get('team', '')
        filter_user = request.args.get('user_email', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'activities': []
            })
        
        # Get user names
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        placeholders = ','.join(['?'] * len(accessible_emails))
        cursor.execute(f'''
            SELECT email, name FROM users
            WHERE email IN ({placeholders})
        ''', accessible_emails)
        users_dict = {row[0]: row[1] or row[0] for row in cursor.fetchall()}
        conn.close()
        
        # Build WHERE clauses
        eval_where_parts = [f"user_email IN ({placeholders})"]
        handbook_where_parts = [f"user_email IN ({placeholders})"]
        eval_params = list(accessible_emails)
        handbook_params = list(accessible_emails)
        
        if filter_user and filter_user in accessible_emails:
            eval_where_parts.append("user_email = ?")
            handbook_where_parts.append("user_email = ?")
            eval_params.append(filter_user)
            handbook_params.append(filter_user)
        
        if date_from:
            eval_where_parts.append("DATE(timestamp) >= ?")
            handbook_where_parts.append("DATE(timestamp) >= ?")
            eval_params.append(date_from)
            handbook_params.append(date_from)
        
        if date_to:
            eval_where_parts.append("DATE(timestamp) <= ?")
            handbook_where_parts.append("DATE(timestamp) <= ?")
            eval_params.append(date_to)
            handbook_params.append(date_to)
        
        eval_where_clause = " WHERE " + " AND ".join(eval_where_parts)
        handbook_where_clause = " WHERE " + " AND ".join(handbook_where_parts)
        
        # Get recent evaluations
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute(f'''
            SELECT 
                'evaluation' as type,
                user_email,
                filename,
                job_title,
                oorwin_job_id,
                match_percentage,
                timestamp
            FROM evaluations
            {eval_where_clause}
            ORDER BY timestamp DESC
            LIMIT {limit}
        ''', eval_params)
        
        evaluations = cursor.fetchall()
        
        # Get recent handbooks
        cursor.execute(f'''
            SELECT 
                'handbook' as type,
                user_email,
                NULL as filename,
                job_title,
                oorwin_job_id,
                NULL as match_percentage,
                timestamp
            FROM recruiter_handbooks
            {handbook_where_clause}
            ORDER BY timestamp DESC
            LIMIT {limit}
        ''', handbook_params)
        
        handbooks = cursor.fetchall()
        conn.close()
        
        # Combine and sort by timestamp
        activities = []
        for row in evaluations:
            activities.append({
                'type': 'evaluation',
                'user_email': row[1],
                'user_name': users_dict.get(row[1], row[1]),
                'filename': row[2],
                'job_title': row[3],
                'job_id': row[4],
                'match_percentage': row[5],
                'timestamp': row[6]
            })
        
        for row in handbooks:
            activities.append({
                'type': 'handbook',
                'user_email': row[1],
                'user_name': users_dict.get(row[1], row[1]),
                'filename': None,
                'job_title': row[3],
                'job_id': row[4],
                'match_percentage': None,
                'timestamp': row[6]
            })
        
        # Sort by timestamp descending
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        # Return top N
        activities = activities[:limit]
        
        return jsonify({
            'success': True,
            'activities': activities
        })
        
    except Exception as e:
        logging.error(f"Error fetching recent activity: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_top_jobs():
    """API endpoint for top jobs by activity (filtered by team)"""
    try:
        user_email = session['user'].get('email')
        limit = int(request.args.get('limit', 10))
        filter_user = request.args.get('user_email', '')
        filter_team = request.args.get('team', '')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Filter by team if specified
        if filter_team:
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM users WHERE team = ?', (filter_team,))
            team_emails = [row[0] for row in cursor.fetchall()]
            conn.close()
            accessible_emails = [e for e in accessible_emails if e in team_emails]
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'jobs': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Build WHERE clause
        where_parts = ["e.oorwin_job_id IS NOT NULL AND e.oorwin_job_id != ''"]
        params = []
        
        # Filter by accessible users
        placeholders = ','.join(['?'] * len(accessible_emails))
        where_parts.append(f"e.user_email IN ({placeholders})")
        params.extend(accessible_emails)
        
        # Filter by specific user if selected
        if filter_user and filter_user in accessible_emails:
            where_parts.append("e.user_email = ?")
            params.append(filter_user)
        
        # Filter by date range
        if date_from:
            where_parts.append("DATE(e.timestamp) >= ?")
            params.append(date_from)
        
        if date_to:
            where_parts.append("DATE(e.timestamp) <= ?")
            params.append(date_to)
        
        where_clause = " WHERE " + " AND ".join(where_parts)
        
        cursor.execute(f'''
            SELECT 
                e.oorwin_job_id,
                e.job_title,
                COUNT(*) as eval_count,
                AVG(e.match_percentage) as avg_score,
                MAX(e.timestamp) as last_active
            FROM evaluations e
            {where_clause}
            GROUP BY e.oorwin_job_id, e.job_title
            ORDER BY eval_count DESC
            LIMIT {limit}
        ''', params)
        
        rows = cursor.fetchall()
        conn.close()
        
        jobs = []
        for row in rows:
            jobs.append({
                'job_id': row[0],
                'job_title': row[1],
                'eval_count': row[2],
                'avg_score': round(row[3], 1) if row[3] else 0,
                'last_active': row[4]
            })
        
        return jsonify({
            'success': True,
            'jobs': jobs
        })
        
    except Exception as e:
        logging.error(f"Error fetching top jobs: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_single_handbook(handbook_id):
    """API endpoint to get a single handbook by ID"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, job_description, timestamp, markdown_content
            FROM recruiter_handbooks
            WHERE id = ?
        ''', (handbook_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            return jsonify({
                'success': False,
                'message': 'Handbook not found'
            }), 404
        
        handbook = {
            'id': row[0],
            'job_title': row[1],
            'oorwin_job_id': row[2],
            'job_description': row[3],
            'timestamp': row[4],
            'markdown_content': row[5]
        }
        
        return jsonify({
            'success': True,
            'handbook': handbook
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbook {handbook_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_handbooks_only():
    """API endpoint for handbooks-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, job_title, oorwin_job_id, timestamp, markdown_content
            FROM recruiter_handbooks
            ORDER BY timestamp DESC
        ''')
        
        handbooks = []
        for row in cursor.fetchall():
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'oorwin_job_id': row[2],
                'timestamp': row[3],
                'markdown_content': row[4][:500] + '...' if row[4] and len(row[4]) > 500 else row[4]  # Truncate for preview
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_evaluations_only():
    """API endpoint for evaluations-only history"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, job_title, oorwin_job_id, match_percentage, timestamp
            FROM evaluations
            ORDER BY timestamp DESC
        ''')
        
        evaluations = []
        for row in cursor.fetchall():
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'job_title': row[2],
                'oorwin_job_id': row[3] if row[3] else 'N/A',
                'match_percentage': row[4],
                'timestamp': row[5]
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations-only history: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_job_centric_history():
    """API endpoint for job-centric grouped history (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        # Debug logging
        logging.info(f"Job-centric history API - Current user: {user_email}, Accessible emails: {accessible_emails}")
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'jobs': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        result = fetch_job_centric_history(cursor, accessible_emails)
        conn.close()
        
        return jsonify({
            'success': True,
            'jobs': result
        })
        
    except Exception as e:
        logging.error(f"Error fetching job-centric history: {str(e)}")
        return jsonify({
            'success': False,
            'message': safe_api_error(e, fallback='Failed to load job history'),
        }), 500

@login_required
def get_handbooks_by_job(job_id):
    """API endpoint to get all handbooks for a specific job ID (filtered by role)"""
    try:
        # Get accessible user emails based on role
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'handbooks': []
            })
        
        # Build placeholders for SQL query
        placeholders = ','.join(['?'] * len(accessible_emails))
        
        cursor.execute(f'''
            SELECT id, job_title, job_description, additional_context, markdown_content, timestamp, user_email
            FROM recruiter_handbooks
            WHERE oorwin_job_id = ? AND user_email IN ({placeholders})
            ORDER BY timestamp DESC
        ''', (job_id,) + tuple(accessible_emails))
        
        rows = cursor.fetchall()
        
        # Get unique user emails from handbooks
        user_emails = set()
        for row in rows:
            if row[6]:  # user_email
                user_emails.add(row[6])
        
        # Fetch user names
        user_names = {}
        if user_emails:
            email_list = list(user_emails)
            user_placeholders = ','.join(['?'] * len(email_list))
            cursor.execute(f'''
                SELECT email, name FROM users
                WHERE email IN ({user_placeholders})
            ''', email_list)
            for email, name in cursor.fetchall():
                user_names[email] = name or email
        
        conn.close()
        
        handbooks = []
        for row in rows:
            user_email = row[6]
            user_name = user_names.get(user_email, user_email) if user_email else 'Unknown'
            
            handbooks.append({
                'id': row[0],
                'job_title': row[1],
                'job_description': row[2],
                'additional_context': row[3],
                'markdown_content': row[4],
                'timestamp': row[5],
                'user_email': user_email,
                'user_name': user_name
            })
        
        return jsonify({
            'success': True,
            'handbooks': handbooks
        })
        
    except Exception as e:
        logging.error(f"Error fetching handbooks for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def download_evaluation_pdf():
    """API endpoint to download resume evaluation as PDF"""
    try:
        data = request.get_json()
        evaluation_data = data.get('evaluation_data', {})
        
        if not evaluation_data:
            return jsonify({
                'success': False,
                'message': 'No evaluation data to download'
            }), 400
        
        logging.info("Generating PDF from evaluation data...")
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        MAIN_SECTION_COLOR = colors.HexColor('#0d6fae')

        def _escape_pdf_text(value) -> str:
            """Escape text for ReportLab Paragraph + normalize problematic unicode."""
            text = unescape(str(value or ''))
            # Normalize hyphen/dash variants that can render/extract badly in PDFs
            text = (text
                    .replace('\u2011', '-')  # non-breaking hyphen
                    .replace('\u2010', '-')  # hyphen
                    .replace('\u2012', '-')  # figure dash
                    .replace('\u2013', '-')  # en dash
                    .replace('\u2014', '-')  # em dash
                    .replace('\u2212', '-')  # minus
                    .replace('\u00a0', ' ')  # nbsp
                    # Remove common emoji markers used in UI labels
                    .replace('✅', '')
                    .replace('❌', '')
                    .replace('⚠️', '')
                    .replace('⚠', '')
                    )
            return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a2e',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#1a1a2e',
            spaceAfter=12,
            spaceBefore=12
        )

        # Main section header style (consistent color, bold)
        main_heading1_style = ParagraphStyle(
            'MainHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=MAIN_SECTION_COLOR,
            spaceAfter=0,
            spaceBefore=12,
            fontName='Helvetica-Bold',
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#6b7280',
            spaceAfter=10,
            spaceBefore=10
        )

        table_cell_style = ParagraphStyle(
            'EvalPdfTableCell',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            spaceAfter=0,
            spaceBefore=0,
        )

        table_header_style = ParagraphStyle(
            'EvalPdfTableHeader',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            spaceAfter=0,
            spaceBefore=0,
            fontName='Helvetica-Bold',
            textColor=colors.whitesmoke,
        )

        def _make_pdf_table(rows, col_widths, header_bg=MAIN_SECTION_COLOR):
            """Create a ReportLab table with consistent styling and wrapped cells."""
            processed = []
            for r_idx, row in enumerate(rows):
                out_row = []
                for cell in row:
                    style = table_header_style if r_idx == 0 else table_cell_style
                    out_row.append(Paragraph(_escape_pdf_text(cell), style))
                processed.append(out_row)

            tbl = Table(processed, colWidths=col_widths, repeatRows=1, hAlign='LEFT')
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), header_bg),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#6b7280')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ]))
            return tbl

        def _add_main_section_heading(text: str):
            """Add a main section heading with consistent color + underline separator."""
            clean = _escape_pdf_text(text)
            # Draw underline with the heading in the SAME flowable so it visually touches the text.
            heading_tbl = Table(
                [[Paragraph(f"<b>{clean}</b>", main_heading1_style)]],
                colWidths=[A4[0] - 100],
                hAlign='LEFT'
            )
            heading_tbl.setStyle(TableStyle([
                ('LINEBELOW', (0, 0), (-1, -1), 1.1, MAIN_SECTION_COLOR),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            elements.append(heading_tbl)
            # Tiny gap below underline (keeps content readable without wasting space)
            elements.append(Spacer(1, 0.04 * inch))
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_LEFT,
            spaceAfter=12
        )
        
        # Title
        job_title = evaluation_data.get('job_title', 'Resume Evaluation')
        filename = evaluation_data.get('filename', 'Unknown')
        elements.append(Paragraph("Resume Evaluation Report", title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Job Information
        elements.append(Paragraph(f"<b>Job Title:</b> {_escape_pdf_text(job_title)}", body_style))
        elements.append(Paragraph(f"<b>Candidate Resume:</b> {_escape_pdf_text(filename)}", body_style))
        if evaluation_data.get('oorwin_job_id'):
            elements.append(Paragraph(f"<b>Job ID:</b> {_escape_pdf_text(evaluation_data.get('oorwin_job_id'))}", body_style))
        # Reduce whitespace before the first section header
        elements.append(Spacer(1, 0.10 * inch))
        
        # Match Score
        match_percentage = evaluation_data.get('match_percentage', 0)
        match_percentage_str = evaluation_data.get('match_percentage_str', f"{match_percentage}%")
        _add_main_section_heading(f"Match Score: {match_percentage_str}")
        elements.append(Spacer(1, 0.2*inch))
        
        # Match Factors
        match_factors = evaluation_data.get('match_factors', {})
        if match_factors:
            _add_main_section_heading("Match Factors Breakdown")
            available_width = A4[0] - 100

            tile_style = ParagraphStyle(
                'MatchFactorTile',
                parent=styles['Normal'],
                fontSize=9,
                leading=11,
                spaceAfter=0,
                spaceBefore=0,
                alignment=TA_CENTER,
            )

            def _get_factor_value(mf: dict, keys: list[str]):
                if not isinstance(mf, dict):
                    return None
                # direct keys
                for k in keys:
                    if k in mf:
                        return mf.get(k)
                # case-insensitive fallback
                lower_map = {str(k).lower(): v for k, v in mf.items()}
                for k in keys:
                    lk = str(k).lower()
                    if lk in lower_map:
                        return lower_map.get(lk)
                return None

            tiles = [
                ("Skills", _get_factor_value(match_factors, ["Skills Match", "Skills"])),
                ("Experience", _get_factor_value(match_factors, ["Experience Match", "Experience"])),
                ("Education", _get_factor_value(match_factors, ["Education Match", "Education"])),
                ("Industry", _get_factor_value(match_factors, ["Industry Knowledge", "Industry Match", "Industry"])),
                ("Certification", _get_factor_value(match_factors, ["Certification Match", "Certification"])),
            ]

            def _fmt_percent(val):
                if val is None:
                    return "N/A"
                try:
                    # val may be "85" or 85 or "85%"
                    s = str(val).strip()
                    if s.endswith('%'):
                        s = s[:-1].strip()
                    num = float(s)
                    return f"{int(round(num))}%"
                except Exception:
                    s = str(val).strip()
                    return s if s else "N/A"

            row = []
            for label, val in tiles:
                row.append(Paragraph(
                    f"<font size='12' color='#0d6fae'><b>{_escape_pdf_text(_fmt_percent(val))}</b></font><br/>"
                    f"<font size='8' color='#6b7280'>{_escape_pdf_text(label)}</font>",
                    tile_style
                ))

            col_widths = [available_width / len(row)] * len(row)
            tile_table = Table([row], colWidths=col_widths, hAlign='LEFT')
            tile_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
                ('BOX', (0, 0), (-1, -1), 0.8, MAIN_SECTION_COLOR),
                ('INNERGRID', (0, 0), (-1, -1), 0.8, MAIN_SECTION_COLOR),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            elements.append(tile_table)
            elements.append(Spacer(1, 0.2 * inch))
        
        # Profile Summary
        profile_summary = evaluation_data.get('profile_summary', '')
        if profile_summary:
            _add_main_section_heading("Profile Summary")
            elements.append(Paragraph(_escape_pdf_text(profile_summary), body_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Missing Keywords (removed from PDF as requested)

        # Candidate Fit Analysis (as shown in UI)
        candidate_fit = evaluation_data.get('candidate_fit_analysis') or {}
        if isinstance(candidate_fit, str):
            try:
                candidate_fit = json.loads(candidate_fit)
            except Exception:
                candidate_fit = {}

        if isinstance(candidate_fit, dict) and candidate_fit:
            available_width = A4[0] - 100
            _add_main_section_heading("Candidate Fit Analysis against Job JD")

            # 1) Dimension Evaluation table
            dims = candidate_fit.get('Dimension Evaluation') or []
            if isinstance(dims, list) and dims:
                rows = [["Dimension", "Evaluation", "Recruiter Comments"]]
                for dim in dims:
                    if isinstance(dim, dict):
                        rows.append([
                            dim.get('Dimension', ''),
                            dim.get('Evaluation', ''),
                            dim.get('Recruiter Comments', ''),
                        ])
                    else:
                        rows.append([str(dim), '', ''])

                col_widths = [available_width * 0.24, available_width * 0.18, available_width * 0.58]
                elements.append(_make_pdf_table(rows, col_widths))
                elements.append(Spacer(1, 0.2 * inch))
            else:
                elements.append(Paragraph("No dimension evaluation available.", body_style))
                elements.append(Spacer(1, 0.2 * inch))

            # 2) Risk & Gaps to Probe table
            risks = candidate_fit.get('Risk and Gaps')
            elements.append(Paragraph("<b>Risk & Gaps to Probe</b>", heading2_style))
            if isinstance(risks, list) and risks:
                rows = [["Area", "Risk", "Recruiter Strategy"]]
                for risk in risks:
                    if isinstance(risk, dict):
                        rows.append([
                            risk.get('Area', ''),
                            risk.get('Risk', ''),
                            risk.get('Recruiter Strategy', ''),
                        ])
                    else:
                        rows.append([str(risk), '', ''])
                col_widths = [available_width * 0.22, available_width * 0.36, available_width * 0.42]
                elements.append(_make_pdf_table(rows, col_widths))
            else:
                elements.append(Paragraph("No major risks identified.", body_style))
            elements.append(Spacer(1, 0.2 * inch))

            # 3) Recruiter Recommendation
            rec = candidate_fit.get('Recommendation') or {}
            elements.append(Paragraph("<b>Recruiter Recommendation</b>", heading2_style))
            if isinstance(rec, dict) and rec:
                rows = [["Verdict", "Fit Level", "Rationale"]]
                rows.append([
                    rec.get('Verdict', 'N/A'),
                    rec.get('Fit Level', 'N/A'),
                    rec.get('Rationale', 'No rationale provided'),
                ])
                col_widths = [available_width * 0.22, available_width * 0.14, available_width * 0.64]
                elements.append(_make_pdf_table(rows, col_widths))
            else:
                elements.append(Paragraph("No recommendation available.", body_style))
            elements.append(Spacer(1, 0.2 * inch))

            # 4) Recruiter Narrative
            narrative = candidate_fit.get('Recruiter Narrative') or ''
            elements.append(Paragraph("<b>Recruiter Narrative for Submission</b>", heading2_style))
            if narrative:
                elements.append(Paragraph(f"<i>{_escape_pdf_text(narrative)}</i>", body_style))
            else:
                elements.append(Paragraph("No recruiter narrative available.", body_style))
            elements.append(Spacer(1, 0.2 * inch))
        
        # Job Stability Analysis (removed from PDF as requested)
        
        # Career Progression
        career_progression = evaluation_data.get('career_progression', {})
        if isinstance(career_progression, str):
            try:
                career_progression = json.loads(career_progression)
            except Exception:
                career_progression = {}
        if career_progression:
            _add_main_section_heading("Career Progression Analysis")
            if career_progression.get('progression_score'):
                elements.append(Paragraph(f"Progression Score: {_escape_pdf_text(career_progression.get('progression_score'))}", body_style))
            if career_progression.get('key_observations'):
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph("<b>Key Observations:</b>", heading2_style))
                for obs in career_progression.get('key_observations', []):
                    elements.append(Paragraph(f"• {_escape_pdf_text(obs)}", body_style))
            elements.append(Spacer(1, 0.2*inch))

            # Red Flags & Concerns (removed from PDF as requested)
        
        # Qualification Fit Assessment (removed from PDF as requested)
        
        # Interview Questions (removed from PDF as requested)
        
        # Build PDF with metadata + logo + watermark (watermark drawn BEFORE content for readability)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        logo_candidates = (
            'peoplelogic-logo-300x77.png',  # NEW preferred logo
            'peoplelogic-logo.png',
            'logo.png',
            'peoplelogic.png',
        )
        logo_path = next(
            (os.path.join(base_dir, 'static', n) for n in logo_candidates if os.path.exists(os.path.join(base_dir, 'static', n))),
            None
        )

        watermark_text = "PeopleLogic Confidential"
        watermark_alpha = 0.04  # 96% transparent (very light so it doesn't block reading)
        watermark_color = colors.HexColor('#f2f2f2')  # near-white gray

        def _draw_branding(canv, _doc):
            # Metadata
            try:
                title_parts = [p for p in ["Resume Evaluation", job_title, filename] if p]
                canv.setTitle(" - ".join(title_parts))
                canv.setAuthor("PeopleLogic PeopleBot")
            except Exception:
                pass

            # Watermark behind content
            try:
                canv.saveState()
                try:
                    canv.setFillAlpha(watermark_alpha)
                except Exception:
                    # If alpha isn't supported, rely on the near-white color
                    pass
                canv.setFillColor(watermark_color)
                canv.setFont('Helvetica-Bold', 32)  # +2 compared to previous size (30)
                canv.translate(A4[0] / 2.0, A4[1] / 2.0)
                canv.rotate(35)
                canv.drawCentredString(0, 0, watermark_text)
                canv.restoreState()
            except Exception:
                pass

            # Logo in the header (top-right)
            try:
                if logo_path:
                    img = ImageReader(logo_path)
                    iw, ih = img.getSize()
                    max_w = 130
                    max_h = 36
                    scale = min(max_w / float(iw), max_h / float(ih))
                    w = iw * scale
                    h = ih * scale
                    x = A4[0] - 50 - w
                    y = A4[1] - 28 - h
                    canv.drawImage(img, x, y, width=w, height=h, mask='auto', preserveAspectRatio=True)
            except Exception:
                pass

        doc.build(elements, onFirstPage=_draw_branding, onLaterPages=_draw_branding)
        
        # Get PDF data
        pdf_data = buffer.getvalue()
        buffer.close()
        
        logging.info("PDF generated successfully")
        
        # Return PDF as response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=Resume_Evaluation_{filename}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating evaluation PDF: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Failed to generate PDF: {str(e)}'
        }), 500

@login_required
def get_evaluation_full(eval_id):
    """API endpoint to get full evaluation data for viewing"""
    try:
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Fetch evaluation data
        cursor.execute('''
            SELECT 
                e.id, e.filename, e.job_title, e.job_description,
                e.match_percentage, e.match_factors, e.profile_summary,
                e.missing_keywords, e.job_stability, e.career_progression,
                e.oorwin_job_id, e.timestamp, e.candidate_fit_analysis, e.over_under_qualification
            FROM evaluations e
            WHERE e.id = ?
        ''', (eval_id,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({
                'success': False,
                'message': 'Evaluation not found'
            }), 404
        
        # Fetch interview questions
        cursor.execute('''
            SELECT technical_questions, nontechnical_questions, behavioral_questions
            FROM interview_questions
            WHERE evaluation_id = ?
        ''', (eval_id,))
        
        questions_row = cursor.fetchone()
        conn.close()
        
        # Parse JSON fields
        import json
        evaluation = {
            'id': row[0],
            'filename': row[1],
            'job_title': row[2],
            'job_description': row[3],
            'match_percentage': row[4],
            'match_percentage_str': str(int(row[4])) + '%' if row[4] else '0%',
            'match_factors': json.loads(row[5]) if row[5] else {},
            'profile_summary': row[6],
            'missing_keywords': json.loads(row[7]) if row[7] else [],
            'job_stability': json.loads(row[8]) if row[8] else {},
            'career_progression': json.loads(row[9]) if row[9] else {},
            'oorwin_job_id': row[10],
            'timestamp': row[11],
            # Parse new fields from database (or use empty if not present)
            'candidate_fit_analysis': json.loads(row[12]) if (len(row) > 12 and row[12]) else {},
            'over_under_qualification': row[13] if (len(row) > 13 and row[13]) else ''
        }
        
        # Helper function to normalize questions (convert objects to strings)
        def normalize_questions(questions_list):
            if not questions_list:
                return []
            normalized = []
            for q in questions_list:
                if isinstance(q, str):
                    normalized.append(q)
                elif isinstance(q, dict):
                    # Extract question text from common property names
                    normalized.append(q.get('question') or q.get('text') or q.get('content') or q.get('value') or str(q))
                else:
                    normalized.append(str(q))
            return normalized
        
        if questions_row:
            tech_raw = json.loads(questions_row[0]) if questions_row[0] else []
            nontech_raw = json.loads(questions_row[1]) if questions_row[1] else []
            behavioral_raw = json.loads(questions_row[2]) if questions_row[2] else []
            
            evaluation['technical_questions'] = normalize_questions(tech_raw)
            evaluation['nontechnical_questions'] = normalize_questions(nontech_raw)
            evaluation['behavioral_questions'] = normalize_questions(behavioral_raw)
        else:
            evaluation['technical_questions'] = []
            evaluation['nontechnical_questions'] = []
            evaluation['behavioral_questions'] = []
        
        return jsonify({
            'success': True,
            'evaluation': evaluation
        })
        
    except Exception as e:
        logging.error(f"Error fetching full evaluation {eval_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@login_required
def get_evaluations_by_job(job_id):
    """API endpoint to get all evaluations for a specific job ID (filtered by team membership)"""
    try:
        # Get accessible user emails based on team membership
        user_email = session['user'].get('email')
        accessible_emails = get_accessible_user_emails(user_email)
        
        if not accessible_emails:
            return jsonify({
                'success': True,
                'evaluations': []
            })
        
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # Build placeholders for SQL query
        placeholders = ','.join(['?'] * len(accessible_emails))
        
        # Get evaluations for this job (with evaluator email)
        cursor.execute(f'''
            SELECT id, filename, match_percentage, timestamp, user_email,
                   evaluation_mode, batch_group_id
            FROM evaluations
            WHERE oorwin_job_id = ? AND user_email IN ({placeholders})
            ORDER BY timestamp DESC
        ''', (job_id,) + tuple(accessible_emails))
        rows = cursor.fetchall()

        # Collect unique evaluator emails
        evaluator_emails = {row[4] for row in rows if row[4]}

        # Map emails to user names
        evaluator_names = {}
        if evaluator_emails:
            email_list = list(evaluator_emails)
            email_placeholders = ','.join(['?'] * len(email_list))
            cursor.execute(f'''
                SELECT email, name FROM users
                WHERE email IN ({email_placeholders})
            ''', email_list)
            for email, name in cursor.fetchall():
                evaluator_names[email] = name or email

        conn.close()
        
        evaluations = []
        for row in rows:
            email = row[4]
            eval_mode = (row[5] or 'single').strip().lower() if len(row) > 5 else 'single'
            evaluations.append({
                'id': row[0],
                'filename': row[1],
                'display_filename': display_upload_filename(row[1]),
                'candidate_name': candidate_display_name(row[1]),
                'match_percentage': row[2],
                'timestamp': row[3],
                'user_email': email,
                'evaluator_name': evaluator_names.get(email, email) if email else 'Unknown',
                'evaluation_mode': eval_mode,
                'batch_group_id': row[6] if len(row) > 6 else None,
            })
        
        return jsonify({
            'success': True,
            'evaluations': evaluations
        })
        
    except Exception as e:
        logging.error(f"Error fetching evaluations for job {job_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@limiter.limit("10 per minute")
@login_required
async def download_handbook_pdf():
    """API endpoint to download recruiter handbook as PDF"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown_content', '').strip()
        job_title = (data.get('job_title') or '').strip()
        oorwin_job_id = (data.get('oorwin_job_id') or '').strip()

        # Some stored/transported handbook content can become double-escaped (literal "\n", "\t").
        # If that happens, PDF formatting collapses. Normalize to real newlines/tabs.
        if '\\n' in markdown_content and '\n' not in markdown_content:
            markdown_content = (
                markdown_content
                .replace('\\r\\n', '\n')
                .replace('\\n', '\n')
                .replace('\\t', '\t')
            )
        
        if not markdown_content:
            return jsonify({
                'success': False,
                'message': 'No content to download'
            }), 400
        
        logging.info("Generating PDF from handbook content...")
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,  # Reduced margins for better table fit
            leftMargin=50,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        MAIN_SECTION_COLOR = colors.HexColor('#0d6fae')
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a2e',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#1a1a2e',
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Blue style for main section titles
        heading1_blue_style = ParagraphStyle(
            'CustomHeading1Blue',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=MAIN_SECTION_COLOR,
            spaceAfter=12,
            spaceBefore=12
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#6b7280',
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Blue style for main section titles (H2)
        heading2_blue_style = ParagraphStyle(
            'CustomHeading2Blue',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=MAIN_SECTION_COLOR,
            spaceAfter=10,
            spaceBefore=10
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['BodyText'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        )

        code_style = ParagraphStyle(
            'CodeBlock',
            parent=styles['BodyText'],
            fontName='Courier',
            fontSize=9,
            leading=11,
            spaceAfter=0,
            spaceBefore=0,
        )

        main_section_heading_style = ParagraphStyle(
            'MainSectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=MAIN_SECTION_COLOR,
            spaceAfter=0,
            spaceBefore=10,
            fontName='Helvetica-Bold',
        )

        def _extract_leading_ui_icon(text: str):
            """
            UI uses emojis like 📖 ✨ ⚠️ ✅ as leading markers.
            ReportLab fonts/Paragraph don't reliably support all emojis (esp. non-BMP like 📖),
            so we render them as small vector icons in the PDF.
            """
            raw = str(text or '')
            t = raw.lstrip()
            if not t:
                return None, raw
            candidates = ('📖', '✨', '⚠️', '⚠', '✅')
            for c in candidates:
                if t.startswith(c):
                    rest = t[len(c):].lstrip()
                    icon = '⚠️' if c in ('⚠️', '⚠') else c
                    return icon, rest
            return None, raw

        class UiIcon(Flowable):
            """Small vector icon used to mimic UI emojis in PDFs."""
            def __init__(self, icon: str, size: float = 12):
                super().__init__()
                self.icon = icon
                self.size = float(size)
                self.width = self.size
                self.height = self.size

            def wrap(self, availWidth, availHeight):
                return self.width, self.height

            def draw(self):
                c = self.canv
                s = self.size

                if self.icon == '✅':
                    # green circle + white check
                    c.saveState()
                    c.setFillColor(colors.HexColor('#7d8e2c'))
                    c.setStrokeColor(colors.HexColor('#647220'))
                    c.circle(s / 2, s / 2, (s / 2) - 0.8, stroke=1, fill=1)
                    c.setStrokeColor(colors.white)
                    c.setLineWidth(max(1.2, s * 0.12))
                    c.line(s * 0.28, s * 0.52, s * 0.43, s * 0.35)
                    c.line(s * 0.43, s * 0.35, s * 0.74, s * 0.68)
                    c.restoreState()
                    return

                if self.icon == '⚠️':
                    # yellow/orange triangle + exclamation
                    c.saveState()
                    c.setFillColor(colors.HexColor('#f6c206'))
                    c.setStrokeColor(colors.HexColor('#d4a605'))
                    p = c.beginPath()
                    p.moveTo(s * 0.50, s * 0.92)
                    p.lineTo(s * 0.08, s * 0.10)
                    p.lineTo(s * 0.92, s * 0.10)
                    p.close()
                    c.drawPath(p, stroke=1, fill=1)
                    c.setStrokeColor(colors.white)
                    c.setLineWidth(max(1.0, s * 0.10))
                    c.line(s * 0.50, s * 0.30, s * 0.50, s * 0.62)
                    c.setFillColor(colors.white)
                    c.circle(s * 0.50, s * 0.20, max(0.8, s * 0.06), stroke=0, fill=1)
                    c.restoreState()
                    return

                if self.icon == '✨':
                    # gold sparkle star
                    c.saveState()
                    c.setStrokeColor(colors.HexColor('#f6c206'))
                    c.setLineWidth(max(1.0, s * 0.10))
                    cx, cy = s / 2, s / 2
                    c.line(cx, s * 0.10, cx, s * 0.90)
                    c.line(s * 0.10, cy, s * 0.90, cy)
                    c.line(s * 0.22, s * 0.22, s * 0.78, s * 0.78)
                    c.line(s * 0.22, s * 0.78, s * 0.78, s * 0.22)
                    c.restoreState()
                    return

                if self.icon == '📖':
                    # simple blue book
                    c.saveState()
                    c.setFillColor(colors.HexColor('#0d6fae'))
                    c.setStrokeColor(colors.HexColor('#0a5a8e'))
                    c.roundRect(s * 0.06, s * 0.12, s * 0.88, s * 0.76, s * 0.10, stroke=1, fill=1)
                    c.setStrokeColor(colors.HexColor('#e8f4fc'))
                    c.setLineWidth(max(0.8, s * 0.06))
                    c.line(s * 0.50, s * 0.14, s * 0.50, s * 0.86)  # spine
                    c.line(s * 0.16, s * 0.72, s * 0.46, s * 0.72)
                    c.line(s * 0.54, s * 0.72, s * 0.84, s * 0.72)
                    c.restoreState()
                    return

        def _icon_paragraph(icon: str, paragraph: Paragraph, icon_size: float = 12):
            """Return a flowable that shows an optional icon before a Paragraph."""
            if not icon:
                return paragraph
            available_width = A4[0] - 100  # matches doc margins (50 each side)
            icon_w = icon_size + 4
            tbl = Table(
                [[UiIcon(icon, size=icon_size), paragraph]],
                colWidths=[icon_w, max(available_width - icon_w, 50)],
                hAlign='LEFT'
            )
            tbl.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            return tbl
        
        # Helper function to check if heading is a main section title (should be blue)
        def is_main_section_title(text):
            """Check if heading text matches one of the 7 main section titles (EXCLUDE Introduction)"""
            # Explicitly exclude "Introduction"
            _icon, text_clean = _extract_leading_ui_icon(text)
            text_clean = str(text_clean or '').strip()
            if re.match(r'^Introduction\s*:?\s*$', text_clean, re.IGNORECASE):
                return False
            
            main_titles = [
                r'^\s*\d+\.?\s*Primary\s+Sourcing\s+Parameters\s*\(Must-Have\)\s*:?',
                r'^\s*\d+\.?\s*Screening\s+Framework\s*:?',
                r'^\s*\d+\.?\s*Target\s+Talent\s+Pools\s*:?',
                r'^\s*\d+\.?\s*Red\s+Flags\s+to\s+Watch\s*:?',
                r'^\s*\d+\.?\s*Recruiter\s+Checklist\s*\(Pre-call\)\s*:?',
                r'^\s*\d+\.?\s*Recruiter\s+Sales\s+Pitch\s*\(to\s+candidates\)\s*:?',
                r'^\s*\d+\.?\s*Overqualification/Overkill\s+Risk\s+Assessment\s*:?',
                # Also match without numbers (in case LLM doesn't include them)
                r'^\s*Primary\s+Sourcing\s+Parameters\s*\(Must-Have\)\s*:?',
                r'^\s*Screening\s+Framework\s*:?',
                r'^\s*Target\s+Talent\s+Pools\s*:?',
                r'^\s*Red\s+Flags\s+to\s+Watch\s*:?',
                r'^\s*Recruiter\s+Checklist\s*\(Pre-call\)\s*:?',
                r'^\s*Recruiter\s+Sales\s+Pitch\s*\(to\s+candidates\)\s*:?',
                r'^\s*Overqualification/Overkill\s+Risk\s+Assessment\s*:?'
            ]
            return any(re.match(pattern, text_clean, re.IGNORECASE) for pattern in main_titles)

        def _add_main_section_heading(text_raw: str):
            """Add one of the 7 main section headers with consistent color + underline separator."""
            _icon, text_clean = _extract_leading_ui_icon(text_raw)
            text_clean = str(text_clean or '').strip()
            if not text_clean:
                return

            elements.append(Spacer(1, 0.02 * inch))

            # Draw underline with the heading in the SAME flowable so it visually touches the text.
            heading_tbl = Table(
                [[Paragraph(_inline_to_rl_markup(text_clean), main_section_heading_style)]],
                colWidths=[A4[0] - 100],
                hAlign='LEFT'
            )
            heading_tbl.setStyle(TableStyle([
                ('LINEBELOW', (0, 0), (-1, -1), 1.1, MAIN_SECTION_COLOR),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            elements.append(heading_tbl)
            # Tiny gap below underline (keeps content readable without wasting space)
            elements.append(Spacer(1, 0.04 * inch))

        def _inline_to_rl_markup(text: str) -> str:
            """Convert a subset of markdown inline syntax to ReportLab Paragraph markup."""
            # Must escape first (ReportLab Paragraph uses an HTML-like mini language)
            text = unescape(text)
            # Normalize common Unicode hyphen/dash variants that can render/extract badly in PDFs.
            # U+2011 (non-breaking hyphen) is the main culprit behind "rednflag", "endntonend", etc.
            text = (text
                    .replace('\u2011', '-')  # non-breaking hyphen
                    .replace('\u2010', '-')  # hyphen
                    .replace('\u2012', '-')  # figure dash
                    .replace('\u2013', '-')  # en dash
                    .replace('\u2014', '-')  # em dash
                    .replace('\u2212', '-')  # minus sign
                    .replace('\u00a0', ' ')  # non-breaking space
                    )
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            # Inline code first to avoid being affected by italic/bold conversions
            text = re.sub(r'`([^`]+)`', r'<font face="Courier">\1</font>', text)
            # Bold + italic
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            return text

        def _is_table_separator(table_line: str) -> bool:
            """Detect markdown table separator rows like |---|---| or ---|---."""
            if '|' not in table_line:
                return False
            stripped = table_line.strip()
            # Remove pipes; remaining should be only -, :, and whitespace
            core = stripped.replace('|', '').strip()
            return bool(core) and all(ch in '-: ' for ch in core)

        def _parse_table_cells(table_line: str) -> list:
            """Parse a markdown table row into cells, supporting optional leading/trailing pipes."""
            parts = [p.strip() for p in table_line.split('|')]
            if parts and parts[0] == '':
                parts = parts[1:]
            if parts and parts[-1] == '':
                parts = parts[:-1]
            return parts

        def _get_list_style(level: int) -> ParagraphStyle:
            level = max(0, min(level, 6))
            return ParagraphStyle(
                f'ListLevel{level}',
                parent=styles['BodyText'],
                fontSize=11,
                leftIndent=20 + (level * 14),
                spaceAfter=6,
            )

        def _flush_table(table_rows_local: list):
            """Convert collected markdown table rows into a ReportLab Table and append to elements."""
            if not table_rows_local:
                return

            # Normalize rows to same column count
            num_cols = max((len(r) for r in table_rows_local), default=0)
            if num_cols <= 0:
                return
            normalized = [r + [''] * (num_cols - len(r)) for r in table_rows_local]

            header_text = ' '.join([str(c) for c in normalized[0]])
            is_sourcing_table = (
                'Skill' in header_text or
                'Experience' in header_text or
                'Recruiter Cue' in header_text or
                'Why It Matters' in header_text
            )

            # Available width (A4 width - margins)
            available_width = A4[0] - 100  # 50pt margin each side (matches doc margins above)

            # Column widths tuned for the 4-col sourcing table
            if is_sourcing_table and num_cols == 4:
                col0 = 25
                col1 = available_width * 0.34
                col2 = available_width * 0.26
                col3 = max(available_width - col0 - col1 - col2, 120)
                col_widths = [col0, col1, col2, col3]
            elif is_sourcing_table and num_cols >= 4:
                col0 = 25
                remaining = max(available_width - col0, 200)
                col_widths = [col0] + [remaining / (num_cols - 1)] * (num_cols - 1)
            else:
                col_widths = [available_width / num_cols] * num_cols

            table_cell_style = ParagraphStyle(
                'TableCell',
                parent=styles['Normal'],
                fontSize=9,
                leading=11,
                spaceAfter=0,
                spaceBefore=0,
            )
            table_header_style = ParagraphStyle(
                'TableHeader',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                spaceAfter=0,
                spaceBefore=0,
                fontName='Helvetica-Bold',
            )

            processed = []
            for r_idx, row in enumerate(normalized):
                row_out = []
                for c in row:
                    cell_style = table_header_style if r_idx == 0 else table_cell_style
                    row_out.append(Paragraph(_inline_to_rl_markup(str(c)), cell_style))
                processed.append(row_out)

            pdf_table = Table(processed, colWidths=col_widths, repeatRows=1, hAlign='LEFT')
            table_style = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0d6fae')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('TOPPADDING', (0, 0), (-1, 0), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#6b7280')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
            ]

            # Blue/bold first column for the sourcing table numbering
            if is_sourcing_table and num_cols >= 4:
                table_style.extend([
                    ('TEXTCOLOR', (0, 1), (0, -1), colors.HexColor('#0d6fae')),
                    ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
                    ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ])

            pdf_table.setStyle(TableStyle(table_style))
            elements.append(Spacer(1, 0.1 * inch))
            elements.append(pdf_table)
            elements.append(Spacer(1, 0.2 * inch))
        
        # Parse markdown content and convert to PDF elements
        # PDF-only filtering requested:
        # - Remove: "2. Screening Framework"
        # - Remove: "6. Recruiter Sales Pitch (to candidates)" (or any numbered variant)
        # - Renumber remaining main sections sequentially
        def _split_heading_prefix(raw: str) -> tuple[str, str]:
            s = str(raw or '').rstrip('\r')
            m = re.match(r'^(#{1,6}\s+)(.*)$', s.strip())
            if m:
                return m.group(1), m.group(2).strip()
            return '', s.strip()

        def _strip_wrapping_bold(text: str) -> str:
            t = str(text or '').strip()
            if t.startswith('**') and t.endswith('**') and len(t) >= 4:
                return t[2:-2].strip()
            return t

        def _is_removed_handbook_section(title: str) -> bool:
            base = re.sub(r'^\s*\d+\.?\s*', '', str(title or '')).strip()
            base = re.sub(r'\s+', ' ', base).strip().lower()
            return (
                base.startswith('screening framework') or
                base.startswith('recruiter sales pitch')
            )

        filtered_lines = []
        skipping_section = False
        new_num = 0
        for raw in markdown_content.splitlines():
            raw_line = str(raw).rstrip('\r')
            # Remove the "Mini Table of Contents" line from the handbook PDF output (robust to markdown/bullets/punctuation)
            cleaned = re.sub(r'[*_`#>\-\u2022:]', ' ', str(raw_line))
            cleaned = re.sub(r'\s+', ' ', cleaned).strip().lower()
            if re.search(r'\bmini\s+tab(?:le)?\s+of\s+contents\b', cleaned):
                continue
            prefix, content = _split_heading_prefix(raw_line)
            content = _strip_wrapping_bold(content)
            _icon, content_no_icon = _extract_leading_ui_icon(content)
            content_no_icon = str(content_no_icon or '').strip()

            if content_no_icon and is_main_section_title(content_no_icon):
                base_title = re.sub(r'^\s*\d+\.?\s*', '', content_no_icon).strip()
                if _is_removed_handbook_section(base_title):
                    skipping_section = True
                    continue

                skipping_section = False
                new_num += 1
                new_heading = f"{new_num}. {base_title}"
                filtered_lines.append((prefix + new_heading).strip())
                continue

            if skipping_section:
                continue

            filtered_lines.append(raw_line)

        markdown_content = "\n".join(filtered_lines)
        lines = markdown_content.splitlines()
        i = 0
        in_table = False
        table_rows = []
        seen_intro = False  # Track if we've seen Introduction section
        
        while i < len(lines):
            raw_line = lines[i].rstrip('\r')
            indent_spaces = len(raw_line) - len(raw_line.lstrip(' '))
            list_level = max(0, min(indent_spaces // 4, 6))  # markdown usually indents by 4 spaces
            line = raw_line.strip()
            
            # Skip empty lines (but add small spacing)
            if not line:
                if not in_table:
                    elements.append(Spacer(1, 0.08 * inch))
                i += 1
                continue

            # Handle fenced code blocks ```...```
            if line.startswith('```'):
                # Flush any pending table before code block
                if in_table and table_rows:
                    _flush_table(table_rows)
                    table_rows = []
                    in_table = False

                code_lines = []
                i += 1
                while i < len(lines) and not str(lines[i]).strip().startswith('```'):
                    code_lines.append(str(lines[i]).rstrip('\r'))
                    i += 1
                # Skip closing fence if present
                if i < len(lines) and str(lines[i]).strip().startswith('```'):
                    i += 1

                code_text = "\n".join(code_lines).rstrip()
                if code_text:
                    code_box = Table(
                        [[Preformatted(code_text, code_style)]],
                        colWidths=[A4[0] - 100],
                        hAlign='LEFT'
                    )
                    code_box.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f1f3f5')),
                        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    elements.append(code_box)
                    elements.append(Spacer(1, 0.15 * inch))
                continue
            
            # Skip duplicate "Introduction:" lines (blue colored duplicates that appear after title)
            if re.match(r'^Introduction:?\s*$', line, re.IGNORECASE):
                if seen_intro:
                    # Skip this duplicate intro line and any following empty lines
                    i += 1
                    while i < len(lines) and not lines[i].strip():
                        i += 1
                    continue
                seen_intro = True
            
            # Skip TOC markdown links (they don't work in PDF anyway)
            if re.match(r'^-\s*\[.*\]\(#.*\)', line):
                i += 1
                continue

            # Treat the 7 main numbered section titles as proper headings (even if markdown lacks ##)
            if is_main_section_title(line):
                _add_main_section_heading(line)
                i += 1
                continue
            
            # Handle markdown tables (supports optional leading/trailing pipes)
            if '|' in line:
                # If we're not already in a table, only start one if the next non-empty line is a separator row
                if not in_table:
                    lookahead = i + 1
                    while lookahead < len(lines) and not str(lines[lookahead]).strip():
                        lookahead += 1
                    if lookahead < len(lines) and _is_table_separator(str(lines[lookahead]).strip()):
                        cells = _parse_table_cells(line)
                        if cells:
                            table_rows.append(cells)
                            in_table = True
                            i = lookahead + 1  # skip separator row
                            continue
                else:
                    # In table: add row or skip separator
                    if _is_table_separator(line):
                        i += 1
                        continue
                    cells = _parse_table_cells(line)
                    if cells:
                        table_rows.append(cells)
                        i += 1
                        continue

            # If we were building a table, finish it when we hit non-table content
            if in_table and table_rows:
                _flush_table(table_rows)
                table_rows = []
                in_table = False
            
            # Keep `line` as raw markdown text for control-flow checks.
            # Apply `_inline_to_rl_markup()` only when creating Paragraphs to avoid double-escaping.
            
            # Handle headers
            if line.startswith('# '):
                text = line[2:].strip()
                icon, text = _extract_leading_ui_icon(text)
                elements.append(Spacer(1, 0.2*inch))
                elements.append(_icon_paragraph(icon, Paragraph(_inline_to_rl_markup(text), title_style), icon_size=14))
            elif line.startswith('## '):
                text = line[3:].strip()
                icon, text_clean = _extract_leading_ui_icon(text)
                text_clean = str(text_clean or '').strip()
                elements.append(Spacer(1, 0.15*inch))
                # Use blue style if it's a main section title
                if is_main_section_title(text_clean):
                    _add_main_section_heading(text_clean)
                else:
                    elements.append(_icon_paragraph(icon, Paragraph(_inline_to_rl_markup(text_clean), heading1_style), icon_size=12))
            elif line.startswith('### '):
                text = line[4:].strip()
                icon, text_clean = _extract_leading_ui_icon(text)
                text_clean = str(text_clean or '').strip()
                elements.append(Spacer(1, 0.1*inch))
                # Use blue style if it's a main section title
                if is_main_section_title(text_clean):
                    _add_main_section_heading(text_clean)
                else:
                    elements.append(_icon_paragraph(icon, Paragraph(_inline_to_rl_markup(text_clean), heading2_style), icon_size=12))
            elif line.startswith('#### '):
                text = line[5:].strip()
                icon, text_clean = _extract_leading_ui_icon(text)
                text_clean = str(text_clean or '').strip()
                elements.append(Spacer(1, 0.08*inch))
                # Use blue style if it's a main section title
                if is_main_section_title(text_clean):
                    _add_main_section_heading(text_clean)
                else:
                    elements.append(_icon_paragraph(icon, Paragraph(_inline_to_rl_markup(text_clean), heading2_style), icon_size=11))
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                text = line[2:].strip() if line.startswith('- ') or line.startswith('* ') else line[2:].strip()
                elements.append(Paragraph('• ' + _inline_to_rl_markup(text), _get_list_style(list_level)))
            elif line.startswith('o '):
                text = line[2:].strip()
                elements.append(Paragraph('○ ' + _inline_to_rl_markup(text), _get_list_style(list_level)))
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line):
                m = re.match(r'^(\d+)\.\s+(.*)$', line)
                num = m.group(1) if m else ''
                body = m.group(2) if m else line
                elements.append(Paragraph(f'{num}. ' + _inline_to_rl_markup(body), _get_list_style(list_level)))
            # Handle bold text (standalone)
            elif line.startswith('**') and line.endswith('**'):
                text = '<b>' + _inline_to_rl_markup(line[2:-2].strip()) + '</b>'
                elements.append(Paragraph(text, body_style))
            # Handle horizontal rules
            elif line.startswith('---') or line.startswith('___'):
                elements.append(Spacer(1, 0.3*inch))
            # Regular text
            else:
                if line.strip():
                    icon, text_clean = _extract_leading_ui_icon(line)
                    if icon:
                        elements.append(_icon_paragraph(icon, Paragraph(_inline_to_rl_markup(text_clean), body_style), icon_size=11))
                    else:
                        elements.append(Paragraph(_inline_to_rl_markup(line), body_style))
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_rows:
            _flush_table(table_rows)
        
        # Build PDF with metadata + logo + watermark (watermark drawn BEFORE content for readability)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        logo_candidates = (
            'peoplelogic-logo-300x77.png',  # NEW preferred logo
            'peoplelogic-logo.png',
            'logo.png',
            'peoplelogic.png',
        )
        logo_path = next(
            (os.path.join(base_dir, 'static', n) for n in logo_candidates if os.path.exists(os.path.join(base_dir, 'static', n))),
            None
        )

        watermark_text = "PeopleLogic Confidential"
        watermark_alpha = 0.04  # 96% transparent (very light so it doesn't block reading)
        watermark_color = colors.HexColor('#f2f2f2')  # near-white gray

        def _draw_branding(canv, _doc):
            # Metadata
            try:
                title_parts = [p for p in ["Recruiter Handbook", job_title, oorwin_job_id] if p]
                canv.setTitle(" - ".join(title_parts))
                canv.setAuthor("PeopleLogic PeopleBot")
            except Exception:
                pass

            # Watermark behind content
            try:
                canv.saveState()
                try:
                    canv.setFillAlpha(watermark_alpha)
                except Exception:
                    pass
                canv.setFillColor(watermark_color)
                canv.setFont('Helvetica-Bold', 32)  # +2 compared to previous size (30)
                canv.translate(A4[0] / 2.0, A4[1] / 2.0)
                canv.rotate(35)
                canv.drawCentredString(0, 0, watermark_text)
                canv.restoreState()
            except Exception:
                pass

            # Logo in the header (top-right)
            try:
                if logo_path:
                    img = ImageReader(logo_path)
                    iw, ih = img.getSize()
                    max_w = 130
                    max_h = 36
                    scale = min(max_w / float(iw), max_h / float(ih))
                    w = iw * scale
                    h = ih * scale
                    x = A4[0] - 50 - w
                    y = A4[1] - 28 - h
                    canv.drawImage(img, x, y, width=w, height=h, mask='auto', preserveAspectRatio=True)
            except Exception:
                pass

        doc.build(elements, onFirstPage=_draw_branding, onLaterPages=_draw_branding)
        
        # Get PDF data
        pdf_data = buffer.getvalue()
        buffer.close()
        
        logging.info("PDF generated successfully")
        
        # Return PDF as response
        response = make_response(pdf_data)
        response.headers['Content-Type'] = 'application/pdf'
        
        # Filename format: Recruiter_Handbook_{Job_Title}.pdf
        safe_title = re.sub(r'[^A-Za-z0-9 _\\-]+', '', job_title).strip() or 'Handbook'
        safe_title = re.sub(r'\\s+', '_', safe_title).strip('_')
        filename = f"Recruiter_Handbook_{safe_title}.pdf"
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500


def get_default_interview_questions(job_title):
    """Generate default interview questions based on job title"""
    # Default technical questions based on common job titles
    technical_questions = {
        "software": [
            "Describe your experience with different programming languages and frameworks.",
            "How do you approach debugging a complex issue in your code?",
            "Explain your understanding of object-oriented programming principles.",
            "How do you ensure code quality and maintainability?",
            "Describe a challenging technical problem you solved recently."
        ],
        "data": [
            "Explain the difference between supervised and unsupervised learning.",
            "How do you handle missing or inconsistent data in your analysis?",
            "Describe your experience with SQL and database optimization.",
            "What tools and libraries do you use for data visualization?",
            "How do you validate the results of your data analysis?"
        ],
        "manager": [
            "How do you approach resource allocation in a project?",
            "Describe your experience with agile methodologies.",
            "How do you handle conflicts within your team?",
            "What metrics do you use to measure project success?",
            "How do you ensure your team meets deadlines and quality standards?"
        ],
        "analyst": [
            "Describe your approach to gathering requirements from stakeholders.",
            "How do you prioritize features or improvements?",
            "What tools do you use for data analysis and reporting?",
            "How do you communicate complex findings to non-technical stakeholders?",
            "Describe a situation where your analysis led to a significant business decision."
        ],
        "designer": [
            "How do you approach the design process for a new project?",
            "Describe your experience with different design tools and software.",
            "How do you incorporate user feedback into your designs?",
            "How do you balance aesthetics with functionality?",
            "Describe a design challenge you faced and how you overcame it."
        ]
    }
    
    # Default non-technical questions
    nontechnical_questions = [
        "How do you prioritize your work when dealing with multiple deadlines?",
        "Describe a situation where you had to collaborate with a difficult team member.",
        "How do you stay updated with the latest trends and developments in your field?",
        "Describe your ideal work environment and company culture.",
        "How do you handle feedback and criticism?"
    ]
    
    # Determine which set of technical questions to use based on job title
    job_title_lower = job_title.lower()
    selected_technical_questions = []
    
    if any(keyword in job_title_lower for keyword in ["developer", "engineer", "programmer", "software", "code", "web"]):
        selected_technical_questions = technical_questions["software"]
    elif any(keyword in job_title_lower for keyword in ["data", "analytics", "scientist", "ml", "ai"]):
        selected_technical_questions = technical_questions["data"]
    elif any(keyword in job_title_lower for keyword in ["manager", "director", "lead", "head"]):
        selected_technical_questions = technical_questions["manager"]
    elif any(keyword in job_title_lower for keyword in ["analyst", "business", "product"]):
        selected_technical_questions = technical_questions["analyst"]
    elif any(keyword in job_title_lower for keyword in ["designer", "ux", "ui", "graphic"]):
        selected_technical_questions = technical_questions["designer"]
    else:
        # If no match, use a mix of questions
        selected_technical_questions = [
            technical_questions["software"][0],
            technical_questions["analyst"][0],
            technical_questions["manager"][0],
            "Describe your technical skills that are most relevant to this position.",
            "What technical challenges are you looking forward to tackling in this role?"
        ]
    
    return selected_technical_questions, nontechnical_questions

# Batch evaluate multiple resumes against the same JD
@limiter.limit("3 per minute")
@login_required
def evaluate_batch():
    """Evaluate multiple resumes against the same JD and return a comparison ranking."""
    try:
        if 'resumes' not in request.files:
            return jsonify({'success': False, 'error': 'No resumes provided'}), 400

        files = request.files.getlist('resumes')
        if not files:
            return jsonify({'success': False, 'error': 'No files received'}), 400

        job_title = request.form.get('job_title')
        job_description = request.form.get('job_description')
        if not job_title or not job_description:
            return jsonify({'success': False, 'error': 'Missing job title or description'}), 400

        additional_context = request.form.get('additional_context', '').strip()
        additional_context_block = f"**Additional Context (client constraints/preference):** {additional_context}" if additional_context else ""

        stored_files = []
        for f in files:
            if f.filename == '' or not allowed_file(f.filename):
                continue
            try:
                file_path, filename = store_uploaded_file(
                    f, app.config['UPLOAD_FOLDER'], ALLOWED_EXTENSIONS
                )
            except ValueError:
                continue
            stored_files.append((file_path, filename))

        eval_overrides = get_evaluation_generation_overrides()

        async def _batch_eval_one(file_path, filename):
            loop = asyncio.get_event_loop()
            resume_text = await loop.run_in_executor(
                _IO_EXECUTOR, extract_text_from_file, file_path
            )
            if not resume_text:
                return None

            formatted_prompt = input_prompt_template.format(
                resume_text=resume_text,
                job_description=job_description,
                additional_context_block=additional_context_block,
            )
            main_response = await async_gemini_generate(formatted_prompt, **eval_overrides)
            if not main_response:
                return None

            match_percentage_str = main_response.get('JD Match', '0%')
            match_percentage = int(str(match_percentage_str).strip('%') or 0)
            _generic_mf_labels = frozenset({
                'skills match', 'experience match', 'education match',
                'industry knowledge', 'certification match', 'jd match',
                'match factors', 'overall match',
            })
            top_strengths = []
            ps = (main_response.get('Profile Summary') or '').strip()
            if ps:
                snippet = ps.split('.')[0].strip()
                if len(snippet) > 140:
                    snippet = snippet[:137] + '...'
                if snippet:
                    top_strengths.append(snippet)
            mf = main_response.get('Match Factors', {})
            if isinstance(mf, dict) and mf and len(top_strengths) < 5:
                try:
                    sorted_items = sorted(
                        mf.items(),
                        key=lambda kv: float(str(kv[1]).split('%')[0])
                        if isinstance(kv[1], str) and '%' in str(kv[1])
                        else float(kv[1]) if kv[1] is not None else 0,
                        reverse=True,
                    )
                except Exception:
                    sorted_items = list(mf.items())
                for k, _v in sorted_items:
                    if len(top_strengths) >= 5:
                        break
                    label = str(k).strip()
                    if label.lower() in _generic_mf_labels:
                        continue
                    if label.lower().endswith(' match'):
                        continue
                    top_strengths.append(label)
            if not top_strengths:
                words = [w.strip('.,;:()').title() for w in ps.split() if len(w) > 3]
                uniq = []
                for w in words:
                    if w not in uniq:
                        uniq.append(w)
                top_strengths = uniq[:5]

            key_gaps = (
                list(main_response.get('MissingKeywords', []))
                if isinstance(main_response.get('MissingKeywords', []), list)
                else []
            )

            return {
                'filename': filename,
                'match_percentage': match_percentage,
                'top_strengths': top_strengths,
                'key_gaps': key_gaps,
                'main_response': main_response,
                'resume_text': resume_text,
            }

        async def _run_batch_parallel():
            sem = asyncio.Semaphore(3)

            async def _limited(file_path, filename):
                async with sem:
                    return await _batch_eval_one(file_path, filename)

            return await asyncio.gather(
                *[_limited(fp, fn) for fp, fn in stored_files],
                return_exceptions=True,
            )

        raw_results = asyncio.run(_run_batch_parallel())
        results = [
            r for r in raw_results
            if isinstance(r, dict) and r.get('filename')
        ]
        for r in raw_results:
            if isinstance(r, Exception):
                logging.error(f"Batch evaluation task failed: {r}")

        results.sort(key=lambda x: x['match_percentage'], reverse=True)
        if not results:
            return jsonify({'success': False, 'error': 'Failed to evaluate uploaded resumes'}), 500

        for r in results:
            r['display_filename'] = display_upload_filename(r.get('filename'))
            r['candidate_name'] = candidate_display_name(r.get('filename'))

        batch_group_id = str(uuid.uuid4())
        user_email = session.get('user', {}).get('email') if 'user' in session else None
        oorwin_job_id = (request.form.get('oorwin_job_id') or '').strip() or None
        for r in results:
            main_response = r.pop('main_response', {}) or {}
            resume_text = r.pop('resume_text', '') or ''
            match_percentage = int(r.get('match_percentage') or 0)
            missing_keywords = (
                list(main_response.get('MissingKeywords', []))
                if isinstance(main_response.get('MissingKeywords', []), list)
                else []
            )
            profile_summary = main_response.get('Profile Summary', '') or 'No summary provided.'
            match_factors = main_response.get('Match Factors', {}) if isinstance(
                main_response.get('Match Factors', {}), dict
            ) else {}
            candidate_fit_analysis = main_response.get('Candidate Fit Analysis', {}) or {}
            over_under_qualification = main_response.get(
                'Over/UnderQualification Analysis',
                'Evaluated as part of a multi-resume batch comparison.',
            )
            stability_data = get_fast_stability_estimate(resume_text) if resume_text else {}
            career_data = get_fast_career_estimate(resume_text) if resume_text else {
                'progression_score': 50,
                'key_observations': [],
                'career_path': [],
                'red_flags': [],
                'reasoning': 'Batch evaluation — limited progression detail',
            }
            additional_info = {
                'job_stability': stability_data,
                'career_progression': career_data,
                'reasoning': main_response.get('Reasoning', ''),
                'batch_comparison': True,
            }
            eval_uuid = str(uuid.uuid4())
            db_id = save_evaluation(
                eval_uuid,
                r['filename'],
                job_title,
                match_percentage,
                missing_keywords,
                profile_summary,
                match_factors,
                stability_data,
                additional_info,
                oorwin_job_id,
                candidate_fit_analysis,
                over_under_qualification,
                user_email,
                None,
                evaluation_mode='batch',
                batch_group_id=batch_group_id,
                job_description=job_description,
            )
            if db_id:
                technical_questions, nontechnical_questions = get_default_interview_questions(job_title)
                save_interview_questions(
                    db_id,
                    json.dumps((technical_questions or [])[:5]),
                    json.dumps((nontechnical_questions or [])[:5]),
                    json.dumps((QUICK_CHECKS or [])[:5]),
                )
                r['evaluation_id'] = db_id
            else:
                r['evaluation_id'] = None
                logging.warning('Batch eval: failed to save evaluation for %s', r.get('filename'))
            r['evaluation_mode'] = 'batch'
            r['batch_group_id'] = batch_group_id

        # Build a recruiter-style markdown comparison report (compact)
        def eval_mark(mark_score):
            if mark_score >= 70:
                return '✅'
            if mark_score >= 40:
                return '⚠️'
            return '❌'

        # JD Summary placeholder (kept short)
        md_lines = []
        md_lines.append('# 🧭 JD Summary')
        md_lines.append('(Concise summary of the role. Auto-generated placeholders — edit as needed.)')
        md_lines.append('')
        md_lines.append('| JD Pillar | Key Expectations |')
        md_lines.append('|------------|------------------|')
        md_lines.append('| Role Objective | Define and deliver measurable impact for the business |')
        md_lines.append('| Core Focus Areas | Execution, stakeholder alignment, metrics |')
        md_lines.append('| Key Competencies | Problem solving, delivery, collaboration |')
        md_lines.append('| Consulting & Client Engagement | Discovery, advisory, influence |')
        md_lines.append('| AI / Analytics / Domain | Practical awareness and usage |')
        md_lines.append('| Cultural Fit | Ownership, clarity, bias for action |')
        md_lines.append('\n---\n')

        # Per-candidate sections
        for r in results:
            cname = r['candidate_name']
            display_file = r['display_filename']
            pct = r.get('match_percentage', 0)
            verdict = (
                'Strong fit' if pct >= 70 else
                ('Moderate fit' if pct >= 40 else 'Weak fit')
            )
            md_lines.append(f'## {cname}')
            md_lines.append('')
            md_lines.append('### Candidate summary')
            md_lines.append('')
            md_lines.append('| Field | Details |')
            md_lines.append('|-------|---------|')
            md_lines.append(f'| **Candidate** | {cname} |')
            md_lines.append(f'| **Resume file** | {display_file} |')
            md_lines.append(f'| **Match score** | {pct}% |')
            md_lines.append(f'| **Verdict** | {verdict} |')
            md_lines.append(f'| **Current role** | — |')
            md_lines.append(f'| **Experience** | — |')
            md_lines.append(f'| **Industry / domain** | — |')
            md_lines.append(f'| **Education** | — |')
            md_lines.append(f'| **Location** | — |')
            themes = ', '.join(r.get('top_strengths', [])[:5]) or '—'
            md_lines.append(f'| **Key themes** | {themes} |')
            md_lines.append('')
            md_lines.append('\n---\n')

            # Ensure non-empty strengths/gaps
            if not r.get('top_strengths'):
                r['top_strengths'] = ['Relevant experience signals', 'Stakeholder collaboration']
            if not r.get('key_gaps'):
                r['key_gaps'] = ['No critical gap surfaced']

            md_lines.append('# 📊 Comparative Fit Analysis (JD vs Resume)')
            md_lines.append('| **Dimension** | **Evaluation** | **Commentary** |')
            md_lines.append('|----------------|----------------|----------------|')
            mark = eval_mark(r['match_percentage'])
            sig = ", ".join(r.get("top_strengths", [])[:3]) or '—'
            gap_one = (r.get('key_gaps') or ['—'])[0]
            md_lines.append(f'| Domain Expertise | {mark} | Signals: {sig} |')
            md_lines.append(f'| Consulting & Advisory Orientation | {mark} | Based on profile narrative |')
            md_lines.append(f'| AI / Analytics Awareness | {mark} | Tooling/awareness inferred |')
            md_lines.append(f'| Account Growth / Leadership | {mark} | Team/initiative ownership |')
            md_lines.append(f'| Client Gravitas (C-suite Influence) | {mark} | Stakeholder influence indicators |')
            md_lines.append(f'| Communication & Storytelling | {mark} | Clarity of outcomes |')
            md_lines.append(f'| Technical or Delivery Depth | {mark} | Depth vs breadth balance |')
            md_lines.append(f'| Cultural Fit (Consulting + Innovation) | {mark} | Bias for action, collaboration |')
            md_lines.append('')

            md_lines.append('# 💪 Key Strengths')
            strengths = r.get('top_strengths', [])[:5] or ['General delivery', 'Collaboration']
            for s in strengths:
                md_lines.append(f'- {s}')
            md_lines.append('')

            md_lines.append('# ⚠️ Gaps / Risks')
            md_lines.append('| Gap | Explanation | Impact |')
            md_lines.append('|------|-------------|---------|')
            gaps = r.get('key_gaps', [])[:3] or ['No critical gap surfaced']
            for g in gaps:
                md_lines.append(f'| {g} | — | Medium |')
            md_lines.append('')

            md_lines.append('# 🧾 Scorecard Summary')
            def to_star(score):
                # Map 0-100 to 1-5
                return max(1, min(5, round(score/20)))
            star = to_star(r['match_percentage'])
            md_lines.append('| Category | Rating (1–5) | Comment |')
            md_lines.append('|-----------|--------------|----------|')
            for cat in ['Domain Fit','Consulting Gravitas','AI / Analytics Awareness','Account Growth Leadership','Client Relationship / Communication','Cultural Fit']:
                md_lines.append(f'| {cat} | {star} | Derived from resume signals |')
            overall10 = round(r['match_percentage']/10, 1)
            verdict = '✅ Strong Fit' if r['match_percentage'] >= 70 else (
                '⚠️ Moderate Fit' if r['match_percentage'] >= 40 else '❌ Weak Fit'
            )
            md_lines.append('')
            md_lines.append(f'**Overall Fit Score:** {overall10} / 10  ')
            md_lines.append(f'**Verdict:** {verdict}')
            md_lines.append('\n---\n')

            md_lines.append('# ✅ Final Recruiter Verdict')
            md_lines.append('> Candidate shows relevant capability signals with room to validate consulting gravitas and delivery depth. Recommend next-step screening focused on stakeholder influence, structured problem solving, and measurable impact.')
            md_lines.append('')

        # Summary comparison table (candidate names as column headers)
        md_lines.append('# Multi-candidate comparison')
        md_lines.append('')
        col_names = [r['candidate_name'] for r in results]
        header = '| **Criteria** | ' + ' | '.join(col_names) + ' |'
        sep = '|:---|' + '|'.join([':---:' for _ in results]) + '|'
        md_lines.append(header)
        md_lines.append(sep)

        def row_line(label, value_fn):
            vals = [str(value_fn(r)) for r in results]
            return '| ' + label + ' | ' + ' | '.join(vals) + ' |'

        md_lines.append(row_line('Match score', lambda r: f"{r.get('match_percentage', 0)}%"))
        md_lines.append(row_line(
            'Verdict',
            lambda r: 'Strong' if r.get('match_percentage', 0) >= 70 else (
                'Moderate' if r.get('match_percentage', 0) >= 40 else 'Weak'
            ),
        ))
        md_lines.append(row_line(
            'Fit (1–10)',
            lambda r: f"{round(r.get('match_percentage', 0) / 10, 1)}",
        ))
        md_lines.append(row_line(
            'Top strength',
            lambda r: (r.get('top_strengths') or ['—'])[0],
        ))
        md_lines.append(row_line(
            'Primary gap',
            lambda r: (r.get('key_gaps') or ['—'])[0],
        ))
        md_lines.append('')
        md_lines.append(
            '*Resume files: ' +
            '; '.join(f"{r['candidate_name']} ({r['display_filename']})" for r in results) +
            '*'
        )

        report_markdown = '\n'.join(md_lines)

        return jsonify({
            'success': True,
            'results': results,
            'report_markdown': report_markdown,
            'batch_group_id': batch_group_id,
            'saved_count': sum(1 for r in results if r.get('evaluation_id')),
        })
    except Exception as e:
        logging.error(f"Error in evaluate_batch: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# Legacy page routes moved off decorators
from pluto.blueprints.pages_legacy import register_legacy_pages_blueprint  # noqa: E402

register_legacy_pages_blueprint()

# API / eval / handbook / analytics / PDF routes (see pluto.blueprints.pluto_api)
from pluto.blueprints.pluto_api import register_pluto_api_blueprint  # noqa: E402

register_pluto_api_blueprint()


def warm_llm_providers():
    """One-token warm-up so the first user request avoids cold SDK latency."""
    try:
        overrides = get_evaluation_generation_overrides()
        generate_content_unified(
            "Reply with OK only.",
            max_completion_tokens_override=8,
            provider_override=overrides.get("provider_override"),
            model_override=overrides.get("model_override"),
        )
        logging.info("LLM provider warm-up completed")
    except Exception as exc:
        logging.debug("LLM warm-up skipped: %s", exc)


if __name__ == "__main__":
    # Initialize logging
    logging.basicConfig(level=logging.INFO,
                       format='%(asctime)s - %(levelname)s - %(message)s')

    # Initialize database
    init_db()
    
    # Update database schema
    update_db_schema()
    
    try:
        # Initialize Pinecone safely
        vectorstore = initialize_pinecone()
        
        # Build BM25 index
        logging.info("🔍 Building BM25 index...")
        build_bm25_index(POLICIES_FOLDER)
        
        # Set up LLM and QA chain
        logging.info("🤖 Setting up LLM and QA chain...")
        llm, qa_chain, retriever = setup_llm_chain()
        
        # Start Flask server with ASGI support using hypercorn
        logging.info("🌐 Starting server...")
        from hypercorn.config import Config
        from hypercorn.asyncio import serve

        config = Config()
        config.bind = ["localhost:5000"]
        config.use_reloader = True
        
        asyncio.run(serve(asgi_app, config))
        
    except Exception as e:
        logging.error(f"❌ Startup error: {str(e)}")
        raise
